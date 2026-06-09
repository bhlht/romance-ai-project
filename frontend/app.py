# -*- coding: utf-8 -*-
"""Romance AI Streamlit Frontend

This file implements the Streamlit UI, user authentication, per‑user project
persistence, and all interactive features (idea generator, story writer, etc.).
The changes focus on three user‑experience issues:
1. **Login persistence** – after a successful login the URL now contains
   `?user=<username>` using `st.experimental_set_query_params`. The query
   parameters survive a page refresh.
2. **Automatic project loading** – when the app starts (or after a refresh) it
   automatically opens the most‑recently‑modified project for the logged‑in user.
   The project can also be forced via a `project` query parameter.
3. **Visibility of storage location** – the absolute path of the user’s
   `story_data/<username>` directory is shown in the sidebar, and a button
   allows copying it to the clipboard.
"""

import os
import json
import glob
import hashlib
import base64
from io import BytesIO
from datetime import datetime
import re

import streamlit as st
import requests
from dotenv import load_dotenv
from docx import Document
import pandas as pd

# ---------------------------------------------------------------------------
# Environment
# ---------------------------------------------------------------------------
load_dotenv(override=True)
# Environment Configuration
ENV = os.getenv("ENV", "dev") # 'dev' or 'prod'

if ENV == "prod":
    BACKEND_URL = os.getenv("BACKEND_URL", "https://romance-ai-backend-46410417920.asia-southeast1.run.app")
    print(f"Server starting in [PRODUCTION] mode. Backend: {BACKEND_URL}")
else:
    BACKEND_URL = os.getenv("BACKEND_URL", "http://localhost:8001")
    print(f"Server starting in [DEVELOPMENT] mode. Backend: {BACKEND_URL}")

def log_error_to_backend(error_type: str, message: str, detail: str = None, context: dict = None) -> None:
    """프론트엔드에서 발생한 예외를 백엔드에 안전하게 로깅합니다."""
    try:
        combined_context = {
            "username": st.session_state.get("username", "anonymous"),
            "current_project": st.session_state.get("current_project", ""),
            "current_chapter_idx": st.session_state.get("current_chapter_idx", 1),
            **(context or {})
        }
        
        import traceback
        tb = detail
        if not tb:
            tb = traceback.format_exc()
            if "NoneType: None" in tb:
                tb = "No frontend traceback captured."
                
        payload = {
            "error_type": error_type,
            "message": message,
            "detail": tb,
            "context": combined_context
        }
        # 백엔드 호출 (타임아웃 3초로 짧게 설정하여 UI 블로킹 방지)
        requests.post(f"{BACKEND_URL}/log_error", json=payload, timeout=3)
    except Exception as e:
        print(f"[FrontendLogger] Failed to log error to backend: {e}")

# ---------------------------------------------------------------------------
# Authentication helpers
# ---------------------------------------------------------------------------
USER_DB_FILE = "users.json"
BASE_DATA_DIR = "story_data"


def hash_password(password: str) -> str:
    """Return a SHA256 hash of *password* (hex string)."""
    return hashlib.sha256(password.encode("utf-8")).hexdigest()

def generate_session_token(username: str, password_hash: str) -> str:
    """Generate a simple stateless session token."""
    return hashlib.sha256(f"{username}{password_hash}".encode("utf-8")).hexdigest()


def load_users() -> dict:
    if os.path.exists(USER_DB_FILE):
        try:
            with open(USER_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_users(users: dict) -> None:
    try:
        with open(USER_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(users, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Save User Error: {e}")


def login_system() -> None:
    """Render the login / registration UI.

    Successful login stores the username in ``st.session_state.user`` and also
    updates the URL query parameters so that a page reload keeps the user logged
    in.
    """
    st.title("🔒 Romance AI: Login / Register")
    col1, col2 = st.columns([1, 1])
    users = load_users()

    # ---------- Login ----------
    with col1:
        st.subheader("Login")
        with st.form("login_form"):
            username = st.text_input("Username", key="login_user")
            password = st.text_input("Password", type="password", key="login_pass")
            submitted = st.form_submit_button("Login")

            if submitted:
                if not re.match(r'^[a-zA-Z0-9]+$', username):
                    st.error("Access denied: Please use an alphanumeric ID (English/Numbers only).")
                elif username in users and users[username] == hash_password(password):
                    st.session_state.user = username
                    st.session_state.logged_out = False
                    # Persist via query params with validation token
                    token = generate_session_token(username, users[username])
                    st.query_params["user"] = username
                    st.query_params["token"] = token
                    st.success(f"Welcome back, {username}!")
                    st.rerun()
                else:
                    st.error("Invalid username or password.")

    # ---------- Register ----------
    with col2:
        st.subheader("Register")
        with st.form("register_form"):
            reg_user = st.text_input("New Username", key="reg_user")
            reg_pass = st.text_input("New Password", type="password", key="reg_pass")
            submitted = st.form_submit_button("Register")
            if submitted:
                if not re.match(r'^[a-zA-Z0-9]+$', reg_user):
                     st.error("Username must contain only English letters and numbers.")
                elif reg_user in users:
                    st.error("Username already exists.")
                elif not reg_user or not reg_pass:
                    st.error("Please fill in both fields.")
                else:
                    users[reg_user] = hash_password(reg_pass)
                    save_users(users)
                    st.success("Registration successful! Please login.")

# ---------------------------------------------------------------------------
# Project persistence helpers
# ---------------------------------------------------------------------------

def get_user_project_dir(username: str) -> str:
    user_dir = os.path.join(BASE_DATA_DIR, username)
    if not os.path.exists(user_dir):
        os.makedirs(user_dir)
    return user_dir


def list_projects(username: str) -> list:
    """Return a list of project names (without extension)."""
    user_dir = get_user_project_dir(username)
    files = glob.glob(os.path.join(user_dir, "*.json"))
    return [os.path.basename(f).replace(".json", "") for f in files]


def list_projects_sorted(username: str) -> list:
    """Return projects sorted by modification time (newest first)."""
    user_dir = get_user_project_dir(username)
    files = glob.glob(os.path.join(user_dir, "*.json"))
    files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return [os.path.basename(p).replace(".json", "") for p in files]


def load_project(username: str, project_name: str) -> dict:
    user_dir = get_user_project_dir(username)
    path = os.path.join(user_dir, f"{project_name}.json")
    if os.path.exists(path):
        try:
            os.utime(path, None)  # Touch file to update mtime for "last active" sorting
        except Exception:
            pass
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_project(username: str, project_name: str, data: dict) -> None:
    user_dir = get_user_project_dir(username)
    path = os.path.join(user_dir, f"{project_name}.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Failed to auto-save project: {e}")


# ── Publisher Hub Storage Helpers ──

def get_user_publisher_dir(username: str) -> str:
    pub_dir = os.path.join(BASE_DATA_DIR, username, "publisher_hub")
    if not os.path.exists(pub_dir):
        os.makedirs(pub_dir)
    return pub_dir

def list_publisher_documents(username: str) -> list:
    pub_dir = get_user_publisher_dir(username)
    files = glob.glob(os.path.join(pub_dir, "*.json"))
    files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return [os.path.basename(f).replace(".json", "") for f in files]

def load_publisher_document(username: str, doc_title: str) -> dict:
    pub_dir = get_user_publisher_dir(username)
    path = os.path.join(pub_dir, f"{doc_title}.json")
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_publisher_document(username: str, doc_title: str, data: dict) -> None:
    pub_dir = get_user_publisher_dir(username)
    path = os.path.join(pub_dir, f"{doc_title}.json")
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"원고 저장 실패: {e}")

def delete_publisher_document(username: str, doc_title: str) -> None:
    pub_dir = get_user_publisher_dir(username)
    path = os.path.join(pub_dir, f"{doc_title}.json")
    if os.path.exists(path):
        try:
            os.remove(path)
        except Exception as e:
            st.error(f"원고 삭제 실패: {e}")

def auto_save_publisher_doc():
    if st.session_state.get("pub_current_doc") and st.session_state.get("user"):
        save_publisher_document(
            st.session_state.user,
            st.session_state.pub_current_doc,
            {
                "raw_text": st.session_state.get("pub_raw_text", ""),
                "filename": st.session_state.get("pub_editor_filename", ""),
                "episodes": st.session_state.get("pub_episodes", []),
                "edits": st.session_state.get("pub_editor_edits", {}),
                "char_sheet": st.session_state.get("pub_local_char_sheet", ""),
                "world_setting": st.session_state.get("pub_local_world_setting", "")
            }
        )


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def create_download_file(text: str, format: str = "txt"):
    if format == "txt":
        return text.encode("utf-8")
    elif format == "docx":
        doc = Document()
        doc.add_heading('Romance Novel', 0)
        for para in text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()
    else:
        return text.encode("utf-8")



def get_clean_rag_params():
    """
    Sanitizes RAG session state variables before sending to backend to avoid 422 validation errors.
    Returns: (enabled, category_id, series_id, keyword)
    """
    enabled = st.session_state.get("rag_enabled", True)
    
    cat_id = st.session_state.get("rag_category_id")
    if cat_id in ("", "None", None):
        cat_id = None
    else:
        try:
            cat_id = int(cat_id)
        except:
            cat_id = None
            
    series_id = st.session_state.get("rag_series_id")
    if series_id in ("", "None", None):
        series_id = None
    else:
        try:
            series_id = int(series_id)
        except:
            series_id = None
            
    keyword = st.session_state.get("rag_keyword", "")
    if keyword in (None, "None"):
        keyword = ""
        
    return enabled, cat_id, series_id, keyword


def fetch_rag_categories():
    try:
        res = requests.get(f"{BACKEND_URL}/rag/categories", timeout=10)
        if res.status_code == 200:
            return res.json().get("categories", [])
    except Exception as e:
        print(f"Failed to fetch RAG categories: {e}")
    return []

def fetch_rag_series(category_id=None, series_id=None, search_query=None):
    # Clean inputs
    if category_id in ("", "None", None):
        category_id = None
    else:
        try: category_id = int(category_id)
        except: category_id = None

    if series_id in ("", "None", None):
        series_id = None
    else:
        try: series_id = int(series_id)
        except: series_id = None

    if category_id is None and series_id is None and search_query is None:
        return []
    try:
        params = {}
        if category_id is not None:
            params["category_id"] = category_id
        if series_id is not None:
            params["series_id"] = series_id
        if search_query is not None:
            params["search_query"] = search_query
        res = requests.get(f"{BACKEND_URL}/rag/series", params=params, timeout=10)
        if res.status_code == 200:
            return res.json().get("series", [])
    except Exception as e:
        print(f"Failed to fetch RAG series: {e}")
    return []



def generate_story(
    prompt: str, 
    temperature: float, 
    model: str, 
    style: str, 
    persona: str, 
    humor_level: int = 0,
    chapter_num: int = 1,
    ch_focus: str = "",
    plot_summary: str = "",
    writer_memo: str = "",
    chapter_brief: str = "",        # [Proactive STEP A] 집필 지침
    continuity_ledger: list = None  # [Proactive STEP C] 연속성 원장
) -> str:
    current_story = st.session_state.get("current_story", "")
    memory_chain = st.session_state.get("memory_chain", [])
    chars = st.session_state.get("char_sheet", "")
    world = st.session_state.get("world_setting", "")
    
    # Fetch previous chapter text if chapter_num > 1 to maintain continuity/bridge
    prev_chapter_text = ""
    if chapter_num > 1 and "chapters" in st.session_state:
        prev_chapter_text = st.session_state.chapters.get(str(chapter_num - 1), "")
    
    # RAG settings from state
    rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
    
    payload = {
        "prompt": prompt,
        "context": current_story,
        "previous_chapter_context": prev_chapter_text,
        "chars": chars,
        "world": world,
        "style": style,
        "persona": persona,
        "humor_level": humor_level,
        "chapter_num": chapter_num,
        "ch_focus": ch_focus,
        "memory_chain": memory_chain,
        "plot_summary": plot_summary,
        "writer_memo": writer_memo,
        "temperature": temperature,
        "model": model,
        "max_length": 2000,
        "rag_enabled": rag_enabled,
        "rag_category_id": rag_category_id,
        "rag_series_id": rag_series_id,
        "rag_keyword": rag_keyword,
        "style_guide": st.session_state.get("ig_style_guide", ""),
        # [Proactive] STEP A Brief 및 STEP C Ledger 주입
        "chapter_brief": chapter_brief or "",
        "continuity_ledger": continuity_ledger if continuity_ledger is not None else st.session_state.get("continuity_ledger", []),
    }
    # Clean payload by removing None values to avoid Pydantic v2 validation errors on old backend versions
    payload = {k: v for k, v in payload.items() if v is not None}
    
    try:
        response = requests.post(
            f"{BACKEND_URL}/generate/romance",
            json=payload, 
            timeout=600 # Increased for V3 pipeline
        )
        if response.status_code == 200:
            new_text = response.json().get("generated_text", "")
            
            # [Auto-Memory] Automatically summarize and save to Long-Term Memory
            try:
                requests.post(
                    f"{BACKEND_URL}/summarize",
                    json={
                        "text": new_text,
                        "chapter_num": chapter_num,
                        "chars": chars,
                        "model": "models/gemini-3-flash-preview"
                    },
                    timeout=60
                )
            except:
                pass # Silently fail for autum-memory as it's not critical for the draft return
                
            return new_text
        else:
            return f"[Error: 서버 오류 ({response.status_code}): {response.text}]"
            
    except Exception as e:
        return f"[Error: 연결 오류 - {str(e)}]"

# ---------------------------------------------------------------------------
# Main application
# ---------------------------------------------------------------------------

def main() -> None:
    st.set_page_config(page_title="Romance AI Creator", page_icon="💖", layout="wide")
    
    # Hide default Streamlit Deploy button, Main Menu (three dots), Header, and Footer
    st.markdown(
        """
        <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        .stDeployButton {display:none !important;}
        header {visibility: hidden;}
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # Model Options Constants
    writer_model_options = [
        "gemini-2.5-pro",
        "gemini-3.5-flash",
        "gemini-2.5-flash",
        "gemini-2.0-flash",
    ]
    assistant_model_options = [
        "gemini-2.5-flash",
        "gemini-3.5-flash",
        "gemini-2.0-flash",
        "gemini-2.5-pro",
    ]

    def handle_logout() -> None:
        st.session_state.user = None
        st.session_state.logged_out = True
        st.query_params.clear()

    # ---------- Auto‑login via query params ----------
    if "user" not in st.session_state:
        st.session_state.user = None
    if "logged_out" not in st.session_state:
        st.session_state.logged_out = False
    
    # Initialize Session State Variables
    if "clear_last_prompt_flag" not in st.session_state:
        st.session_state.clear_last_prompt_flag = False
    if st.session_state.clear_last_prompt_flag:
        st.session_state.last_prompt = ""
        st.session_state.clear_last_prompt_flag = False

    if "setting_target_vols" not in st.session_state:
        st.session_state.setting_target_vols = 1
    if "setting_target_chapters" not in st.session_state:
        st.session_state.setting_target_chapters = 50
    if "chapters_settings" not in st.session_state:
        st.session_state.chapters_settings = {}
    if "apply_plot_style" not in st.session_state:
        st.session_state.apply_plot_style = True
    if "batch_job_id" not in st.session_state:
        st.session_state.batch_job_id = None
    if "auto_merge_trigger" not in st.session_state:
        st.session_state.auto_merge_trigger = None
    if "auto_merge_enabled" not in st.session_state:
        st.session_state.auto_merge_enabled = True
    if "self_healing_enabled" not in st.session_state:
        st.session_state.self_healing_enabled = True

    if "current_page" not in st.session_state:
        st.session_state.current_page = "✍️ Story Engine"
        
    # RAG state variables
    if "rag_enabled" not in st.session_state:
        st.session_state.rag_enabled = True
    if "rag_categories" not in st.session_state:
        st.session_state.rag_categories = fetch_rag_categories()
    if "rag_category_id" not in st.session_state:
        st.session_state.rag_category_id = None
    if "rag_series" not in st.session_state:
        st.session_state.rag_series = fetch_rag_series()
    if "rag_series_id" not in st.session_state:
        st.session_state.rag_series_id = None
    if "rag_keyword" not in st.session_state:
        st.session_state.rag_keyword = ""
    if "rag_series_search_val" not in st.session_state:
        st.session_state.rag_series_search_val = ""
        
    if not st.session_state.user and not st.session_state.logged_out:
        qp = st.query_params
        if "user" in qp and "token" in qp:
            saved_user = str(qp["user"]).strip()
            # Enforce clean usernames even for auto-login
            if not re.match(r'^[a-zA-Z0-9_\-]+$', saved_user):
                st.session_state.user = None
            else:
                saved_token = qp["token"]
                users = load_users()
                if saved_user in users:
                    expected_token = generate_session_token(saved_user, users[saved_user])
                    if saved_token == expected_token:
                        st.session_state.user = saved_user
                    else:
                        st.warning("Previous session expired or invalid. Please login again.")
                else:
                     # Only show error if we expected to find the user
                     st.warning("User record not found. Please login again.")
    if not st.session_state.user:
        login_system()
        return

    username = st.session_state.user

    # [Persistence Fix] Ensure URL always has the token while logged in
    # This prevents the token from disappearing and breaking the session on refresh
    users_lookup = load_users()
    if username in users_lookup:
        valid_token = generate_session_token(username, users_lookup[username])
        qp = st.query_params
        if qp.get("user") != username or qp.get("token") != valid_token:
            st.query_params["user"] = username
            st.query_params["token"] = valid_token
            # No rerun needed here, it updates the URL for the next reload

    # ---------- Sidebar – logout & project manager ----------
    with st.sidebar:
        if st.session_state.current_page == "📦 Publisher Hub":
            st.write(f"👤 **{username} (출판 허브)**")
            if st.button("Logout", key="logout_pub", on_click=handle_logout):
                st.rerun()
                
            st.title("📦 원고 관리자")
            
            pub_mode = st.radio("원고 관리", ["원고 불러오기", "새 원고 등록"], horizontal=True, key="pub_sidebar_mode")
            existing_docs = list_publisher_documents(username)
            
            if "pub_current_doc" not in st.session_state:
                st.session_state.pub_current_doc = existing_docs[0] if existing_docs else None
                
            if pub_mode == "원고 불러오기":
                if existing_docs:
                    items_per_page = 5
                    total_items = len(existing_docs)
                    total_pages = max(1, (total_items + items_per_page - 1) // items_per_page)
                    
                    if "pub_list_page" not in st.session_state:
                        st.session_state.pub_list_page = 0
                        
                    st.session_state.pub_list_page = max(0, min(st.session_state.pub_list_page, total_pages - 1))
                    
                    start_idx = st.session_state.pub_list_page * items_per_page
                    end_idx = min(start_idx + items_per_page, total_items)
                    
                    page_docs = existing_docs[start_idx:end_idx]
                    
                    st.markdown(f"**🗂️ 원고 목록 ({st.session_state.pub_list_page + 1}/{total_pages} Page)**")
                    
                    for doc in page_docs:
                        col_doc, col_del = st.columns([4, 1])
                        with col_doc:
                            active_style = "👉 " if st.session_state.pub_current_doc == doc else "📄 "
                            if st.button(f"{active_style}{doc}", key=f"pub_btn_load_{doc}"):
                                st.session_state.pub_current_doc = doc
                                doc_data = load_publisher_document(username, doc)
                                st.session_state.pub_raw_text = doc_data.get("raw_text", "")
                                st.session_state.pub_editor_filename = doc_data.get("filename", "")
                                st.session_state.pub_episodes = doc_data.get("episodes", [])
                                st.session_state.pub_editor_edits = doc_data.get("edits", {})
                                st.session_state.pub_local_char_sheet = doc_data.get("char_sheet", "")
                                st.session_state.pub_local_world_setting = doc_data.get("world_setting", "")
                                if "pub_editor_current_ep" not in st.session_state:
                                    st.session_state.pub_editor_current_ep = 0
                                st.rerun()
                        with col_del:
                            if st.button("🗑️", key=f"pub_btn_del_{doc}", help=f"'{doc}' 원고 삭제"):
                                delete_publisher_document(username, doc)
                                if st.session_state.pub_current_doc == doc:
                                    st.session_state.pub_current_doc = None
                                    st.session_state.pub_raw_text = ""
                                    st.session_state.pub_editor_filename = ""
                                    st.session_state.pub_episodes = []
                                    st.session_state.pub_editor_edits = {}
                                    st.session_state.pub_local_char_sheet = ""
                                    st.session_state.pub_local_world_setting = ""
                                st.rerun()
                                
                    p_col1, p_col2 = st.columns(2)
                    with p_col1:
                        if st.button("⬅️ 이전", key="pub_page_prev", disabled=(st.session_state.pub_list_page == 0)):
                            st.session_state.pub_list_page -= 1
                            st.rerun()
                    with p_col2:
                        if st.button("다음 ➡️", key="pub_page_next", disabled=(st.session_state.pub_list_page >= total_pages - 1)):
                            st.session_state.pub_list_page += 1
                            st.rerun()
                else:
                    st.warning("등록된 원고가 없습니다.")
            else:  # 새 원고 등록
                new_doc_title = st.text_input("새 작품 제목", placeholder="제목을 입력하세요")
                st.caption("비워둘 경우 파일 업로드 시 본문에서 자동 추출됩니다.")
                if st.button("새 작업 만들기", key="pub_btn_create_new"):
                    if new_doc_title:
                        clean_title = "".join(ch for ch in new_doc_title if ch.isalnum() or ch in "_- ").strip().replace(" ", "_")
                        save_publisher_document(username, clean_title, {
                            "raw_text": "",
                            "filename": f"{clean_title}.txt",
                            "episodes": [],
                            "edits": {},
                            "char_sheet": "",
                            "world_setting": ""
                        })
                        st.session_state.pub_current_doc = clean_title
                        st.session_state.pub_raw_text = ""
                        st.session_state.pub_editor_filename = f"{clean_title}.txt"
                        st.session_state.pub_episodes = []
                        st.session_state.pub_editor_edits = {}
                        st.session_state.pub_local_char_sheet = ""
                        st.session_state.pub_local_world_setting = ""
                        st.rerun()
                    else:
                        st.error("제목을 입력해 주세요.")
        else:
            st.write(f"👤 **{username}**")
            if st.button("Logout", key="logout_proj", on_click=handle_logout):
                st.rerun()
            st.title("🗂️ Project Manager")

            # Load projects sorted by newest first
            existing_projects = list_projects_sorted(username)

            # If a project is supplied via URL, honour it (and verify it exists)
            qp = st.query_params
            if "project" in qp and qp["project"] in existing_projects:
                st.session_state.current_project = qp["project"]
            elif "current_project" not in st.session_state:
                # default to most recent project if any
                st.session_state.current_project = existing_projects[0] if existing_projects else None

            # Project selection UI
            mode = st.radio("Project Action", ["Load Existing", "Create New"], horizontal=True)
            if mode == "Load Existing":
                if existing_projects:
                    sel = st.selectbox("Select Project", existing_projects, index=0 if not st.session_state.current_project else existing_projects.index(st.session_state.current_project))
                    if st.button("Load"):
                        st.session_state.current_project = sel
                        # update URL so refresh keeps this project
                        # Must preserve token
                        token = st.query_params.get("token", "")
                        st.query_params["user"] = username
                        st.query_params["project"] = sel
                        if token:
                            st.query_params["token"] = token
                        st.rerun()
                else:
                    st.warning("No projects found. Create one!")
            else:  # Create New
                new_name = st.text_input("New Project Name", placeholder="My_Romance_Novel_1")
                if new_name and st.button("Create & Switch"):
                    clean = "".join(ch for ch in new_name if ch.isalnum() or ch in "_- ").strip().replace(" ", "_")
                    # Create empty project file immediately so it persists
                    save_project(username, clean, {
                        "current_story": "",
                        "char_sheet": "",
                        "world_setting": "",
                        "memory_chain": [],
                        "critique": None,
                        "idea_suggestion": "",
                        "custom_persona_input": "",
                        "apply_plot_style": True
                    })
                    st.session_state.current_project = clean
                    st.query_params["user"] = username
                    st.query_params["project"] = clean
                    # Must preserve token
                    token = st.query_params.get("token", "")
                    if token:
                        st.query_params["token"] = token
                    st.rerun()

    # If no project selected, show welcome screen (except for Publisher Hub)
    if st.session_state.current_page != "📦 Publisher Hub" and not st.session_state.get("current_project"):
        st.title("Welcome! Please create or load a project to start.")
        return

    # ---------- Load project data (once per project change) ----------
    if st.session_state.current_page != "📦 Publisher Hub" and st.session_state.get("current_project"):
        if "loaded_project" not in st.session_state or st.session_state.loaded_project != st.session_state.current_project:
            data = load_project(username, st.session_state.current_project)
            st.session_state.custom_persona_input = data.get("custom_persona_input", "")
            st.session_state.plot_outline = data.get("plot_outline", "")
            st.session_state.char_sheet = data.get("char_sheet", "")
            st.session_state.world_setting = data.get("world_setting", "")
            st.session_state.memory_chain = data.get("memory_chain", [])
            st.session_state.critique = data.get("critique")
            st.session_state.idea_suggestion = data.get("idea_suggestion", "")
            
            # --- Multi-Chapter Support ---
            raw_chapters = data.get("chapters", {})
            # Migration: If it's an old project with only current_story, put it in Chapter 1
            curr_story = data.get("current_story", "")
            if not raw_chapters and curr_story:
                raw_chapters = {"1": curr_story}
            
            st.session_state.chapters = raw_chapters
            if "current_chapter_idx" not in st.session_state:
                # First priority: check saved active chapter
                saved_idx = data.get("current_chapter_idx")
                if saved_idx is not None:
                    st.session_state.current_chapter_idx = int(saved_idx)
                elif raw_chapters:
                    st.session_state.current_chapter_idx = max([int(k) for k in raw_chapters.keys()])
                else:
                    st.session_state.current_chapter_idx = 1
            
            # Sync current_story with the selected chapter
            idx_str = str(st.session_state.current_chapter_idx)
            st.session_state.current_story = st.session_state.chapters.get(idx_str, "")
            st.session_state.story_editor_ui = st.session_state.current_story
            # Persist Settings
            st.session_state.setting_temperature = data.get("setting_temperature", 0.7)
            st.session_state.widget_setting_temperature = st.session_state.setting_temperature
            st.session_state.setting_model_writer = data.get("setting_model_writer", data.get("setting_model", "gemini-2.5-pro"))
            st.session_state.widget_setting_model_writer = st.session_state.setting_model_writer
            st.session_state.setting_model_assistant = data.get("setting_model_assistant", "gemini-2.5-flash")
            st.session_state.widget_setting_model_assistant = st.session_state.setting_model_assistant
            st.session_state.setting_style = data.get("setting_style", "기본")
            st.session_state.widget_setting_style = st.session_state.setting_style
            st.session_state.setting_preset = data.get("setting_preset", "Direct Input")
            st.session_state.widget_setting_preset = st.session_state.setting_preset
            # Load structural settings directly from json without guards, as we want file updates to take precedence
            st.session_state.setting_target_vols = data.get("setting_target_vols", 1)
            st.session_state.widget_setting_target_vols = st.session_state.setting_target_vols
            st.session_state.setting_target_chapters = data.get("setting_target_chapters", 50)
            st.session_state.widget_setting_target_chapters = st.session_state.setting_target_chapters
            st.session_state.setting_humor = data.get("setting_humor", 5) # Default 5
            st.session_state.widget_setting_humor = st.session_state.setting_humor
            st.session_state.last_prompt = data.get("last_prompt", "")
    
            # Persist Idea Generator
            st.session_state.ig_genre = data.get("ig_genre", "전통로맨스")
            st.session_state.widget_ig_genre = st.session_state.ig_genre
            st.session_state.ig_spice = data.get("ig_spice", "19금(없음)")
            st.session_state.widget_ig_spice = st.session_state.ig_spice
            st.session_state.ig_style_guide = data.get("ig_style_guide", "빠른 전개,\n짧은 문장,\n짧은 문단,\n대화 중심,\n강한 훅(hook),\n대사 55%,\n행동 25%,\n내면 15%,\n배경 설명 5%")
            st.session_state.widget_ig_style_guide = st.session_state.ig_style_guide
            st.session_state.ig_moods = data.get("ig_moods", [])
            st.session_state.widget_ig_moods = st.session_state.ig_moods
            st.session_state.ig_male = data.get("ig_male", [])
            st.session_state.widget_ig_male = st.session_state.ig_male
            st.session_state.ig_female = data.get("ig_female", [])
            st.session_state.widget_ig_female = st.session_state.ig_female
            st.session_state.ig_arc = data.get("ig_arc", "")
            st.session_state.widget_ig_arc = st.session_state.ig_arc
            
            # Load RAG settings
            st.session_state.rag_enabled = data.get("rag_enabled", True)
            
            raw_cat_id = data.get("rag_category_id")
            if raw_cat_id in ("", "None", None):
                st.session_state.rag_category_id = None
            else:
                try: st.session_state.rag_category_id = int(raw_cat_id)
                except: st.session_state.rag_category_id = None
                
            raw_series_id = data.get("rag_series_id")
            if raw_series_id in ("", "None", None):
                st.session_state.rag_series_id = None
            else:
                try: st.session_state.rag_series_id = int(raw_series_id)
                except: st.session_state.rag_series_id = None
                
            st.session_state.rag_keyword = data.get("rag_keyword", "")
            if st.session_state.rag_keyword in (None, "None"):
                st.session_state.rag_keyword = ""
                
            st.session_state.rag_series_search_val = ""
            
            # Load additional UI states
            st.session_state.prediction_report = data.get("prediction_report")
            st.session_state.apply_plot_style = data.get("apply_plot_style", True)
            st.session_state.pkg_titles = data.get("pkg_titles", [])
            st.session_state.pkg_blurb = data.get("pkg_blurb", "")
            st.session_state.pkg_keywords = data.get("pkg_keywords", [])
            st.session_state.review_result = data.get("review_result")
            st.session_state.batch_job_id = data.get("batch_job_id", None)
            st.session_state.batch_status = data.get("batch_status", {})
            st.session_state.auto_merge_enabled = data.get("auto_merge_enabled", True)
            st.session_state.self_healing_enabled = data.get("self_healing_enabled", True)
            
            # Load RAG series cache
            series_cache = []
            if st.session_state.rag_category_id is not None:
                series_cache = fetch_rag_series(category_id=st.session_state.rag_category_id)
                
            if st.session_state.rag_series_id is not None:
                already_loaded = any(s["series_id"] == st.session_state.rag_series_id for s in series_cache)
                if not already_loaded:
                    saved_series_list = fetch_rag_series(series_id=st.session_state.rag_series_id)
                    series_cache.extend(saved_series_list)
                    
            st.session_state.rag_series = series_cache
            st.session_state.loaded_project = st.session_state.current_project
            # [Proactive] 연속성 원장 및 집필 브리프 세션 초기화 (프로젝트 로드 시)
            if "continuity_ledger" not in st.session_state:
                st.session_state.continuity_ledger = data.get("continuity_ledger", [])
            if "chapter_brief" not in st.session_state:
                st.session_state.chapter_brief = ""

    # ---------- Auto‑save helper ----------
    def auto_save() -> None:
        # Check load flag to prevent saving default values before file loading completes
        if st.session_state.get("loaded_project") != st.session_state.get("current_project"):
            return
            
        # Before saving, sync the current_story from UI to the chapters dict
        if "chapters" not in st.session_state:
            st.session_state.chapters = {}
        
        idx_str = str(st.session_state.get("current_chapter_idx", 1))
        st.session_state.chapters[idx_str] = st.session_state.get("current_story", "")

        payload = {
            "current_chapter_idx": st.session_state.get("current_chapter_idx", 1),
            "chapters": st.session_state.chapters, # New multi-chapter field
            "current_story": st.session_state.get("current_story", ""), # Keep for backward compatibility
            "char_sheet": st.session_state.get("char_sheet", ""),
            "world_setting": st.session_state.get("world_setting", ""),
            "memory_chain": st.session_state.get("memory_chain", []),
            "critique": st.session_state.get("critique"),
            "idea_suggestion": st.session_state.get("idea_suggestion", ""),
            "custom_persona_input": st.session_state.get("custom_persona_input", ""),
            # Persist Settings
            "setting_temperature": st.session_state.get("setting_temperature", 0.7),
            "setting_model_writer": st.session_state.get("setting_model_writer", "gemini-2.5-pro"),
            "setting_model_assistant": st.session_state.get("setting_model_assistant", "gemini-2.5-flash"),
            "setting_style": st.session_state.get("setting_style", "기본"),
            "setting_preset": st.session_state.get("setting_preset", "Direct Input"),
            "setting_target_vols": st.session_state.setting_target_vols,
            "setting_target_chapters": st.session_state.setting_target_chapters,
            "setting_humor": st.session_state.setting_humor,
            "last_prompt": st.session_state.get("last_prompt", ""),
            # Persist Idea Generator
            "ig_genre": st.session_state.get("ig_genre", "전통로맨스"),
            "ig_spice": st.session_state.get("ig_spice", "19금(없음)"),
            "ig_style_guide": st.session_state.get("ig_style_guide", "빠른 전개,\n짧은 문장,\n짧은 문단,\n대화 중심,\n강한 훅(hook),\n대사 55%,\n행동 25%,\n내면 15%,\n배경 설명 5%"),
            "ig_moods": st.session_state.get("ig_moods", []),
            "ig_male": st.session_state.get("ig_male", []),
            "ig_female": st.session_state.get("ig_female", []),
            "ig_arc": st.session_state.get("ig_arc", ""),
            "plot_outline": st.session_state.get("plot_outline", ""),
            # RAG Settings
            "rag_enabled": st.session_state.get("rag_enabled", True),
            "rag_category_id": st.session_state.get("rag_category_id"),
            "rag_series_id": st.session_state.get("rag_series_id"),
            "rag_keyword": st.session_state.get("rag_keyword", ""),
            # Additional UI states
            "prediction_report": st.session_state.get("prediction_report"),
            "apply_plot_style": st.session_state.get("apply_plot_style", True),
            "pkg_titles": st.session_state.get("pkg_titles", []),
            "pkg_blurb": st.session_state.get("pkg_blurb", ""),
            "pkg_keywords": st.session_state.get("pkg_keywords", []),
            "chapters_settings": st.session_state.get("chapters_settings", {}),
            "review_result": st.session_state.get("review_result"),
            "batch_job_id": st.session_state.get("batch_job_id", ""),
            "batch_status": st.session_state.get("batch_status", {}),
            "auto_merge_enabled": st.session_state.get("auto_merge_enabled", True),
            "self_healing_enabled": st.session_state.get("self_healing_enabled", True),
            # [Proactive STEP C] 연속성 원장 준영
            "continuity_ledger": st.session_state.get("continuity_ledger", []),
        }
        save_project(username, st.session_state.current_project, payload)

    def on_setting_change(key: str) -> None:
        widget_key = f"widget_{key}"
        if widget_key in st.session_state:
            st.session_state[key] = st.session_state[widget_key]
        auto_save()

    def on_style_change() -> None:
        if "widget_setting_style" in st.session_state:
            st.session_state.setting_style = st.session_state.widget_setting_style
            new_style = st.session_state.setting_style
            guide_mapping = {
                "기본": "빠른 전개,\n짧은 문장,\n짧은 문단,\n대화 중심,\n강한 훅(hook),\n대사 55%,\n행동 25%,\n내면 15%,\n배경 설명 5%",
                "웹소설체": "빠른 전개,\n짧은 문장,\n짧은 문단,\n대화 중심,\n강한 훅(hook),\n대사 55%,\n행동 25%,\n내면 15%,\n배경 설명 5%",
                "감성적": "서정적인 묘사,\n인물의 내밀한 심리 묘사 강화,\n아련하고 감성적인 분위기,\n풍부한 은유와 감각적 형용사 사용,\n내면 40%,\n배경 설명 30%,\n대사 20%,\n행동 10%",
                "담백한": "불필요한 미사여구 배제,\n간결하고 객관적인 3인칭 묘사,\n대사보다는 행동과 절제된 감정 표현 위주,\n행동 50%,\n내면 20%,\n대사 20%,\n배경 설명 10%",
                "고전": "중후하고 고전적인 문체,\n격식 있는 어조와 어휘 사용,\n세밀한 서사 및 역사/배경 설정 묘사,\n대화보다 깊이 있는 지문 묘사 중심,\n배경 설명 40%,\n내면 30%,\n대사 15%,\n행동 15%",
                "유머러스": "재치 있고 가벼운 대사 티키타카,\n코믹하고 과장된 상황 연출,\n웃음을 자아내는 행동 묘사,\n대사 50%,\n행동 30%,\n내면 10%,\n배경 설명 10%"
            }
            new_guide = guide_mapping.get(new_style, "")
            # 새 바인딩 키 방식으로 업데이트
            st.session_state.ig_style_guide = new_guide
            if "ig_style_guide_textarea" in st.session_state:
                st.session_state.ig_style_guide_textarea = new_guide
        auto_save()

    def on_preset_change() -> None:
        if "widget_setting_preset" in st.session_state:
            st.session_state.setting_preset = st.session_state.widget_setting_preset
            preset_val = st.session_state.setting_preset
            persona_presets = {
                "Direct Input": "",
                "김은숙 st": "김은숙 작가 스타일로, 인물 간의 대사가 빠르고 재치 있게...",
                "지브리 st": "지브리 애니메이션처럼, 서정적이고 아름답게...",
                "박찬욱 st": "박찬욱 감독 스타일로, 우아하고 잔혹하게...",
                "막장 드라마": "아침 드라마처럼, 자극적이고 빠르게...",
            }
            st.session_state.custom_persona_input = persona_presets.get(preset_val, "")
            st.session_state.custom_persona_ui = st.session_state.custom_persona_input
        auto_save()

    def init_widget_key(key: str, default_val) -> None:
        widget_key = f"widget_{key}"
        if widget_key not in st.session_state:
            st.session_state[widget_key] = st.session_state.get(key, default_val)

    # ---------- Settings sidebar ----------
    if st.session_state.current_page != "📦 Publisher Hub":
        with st.sidebar:
            # ── 근본 해결: init_widget_key + widget_* 패턴 완전 제거 ─────────────
            # widget_* key 방식은 Streamlit이 위젯 소멸 시 key를 삭제하고,
            # 동시에 value= 와 Session State 양쪽에서 초기화되면 충돌 에러 발생.
            # 모든 사이드바 설정 위젯을 value= 직접 바인딩으로 통일.

            # 기본값 보장 (최초 1회)
            if "setting_temperature" not in st.session_state:
                st.session_state.setting_temperature = 0.7
            if "setting_humor" not in st.session_state:
                st.session_state.setting_humor = 5
            if "setting_model_writer" not in st.session_state:
                st.session_state.setting_model_writer = "gemini-2.5-pro"
            if "setting_model_assistant" not in st.session_state:
                st.session_state.setting_model_assistant = "gemini-2.5-flash"
            if "setting_style" not in st.session_state:
                st.session_state.setting_style = "기본"
            if "setting_preset" not in st.session_state:
                st.session_state.setting_preset = "Direct Input"
            if "setting_target_vols" not in st.session_state:
                st.session_state.setting_target_vols = 1
            if "setting_target_chapters" not in st.session_state:
                st.session_state.setting_target_chapters = 50

            st.markdown("<style>.stProgress {max-width:50%;}</style>", unsafe_allow_html=True)
            st.title("Settings")
            st.caption("⚙️ **프로젝트 기본값 설정**")
            st.caption("이곳의 설정값은 대량 집필(Factory) 모드 및 새로 추가되는 화차의 기본 시작 값으로 적용됩니다. 개별 화차의 스타일 수치는 에디터 본문 하단에서 언제든지 변경할 수 있습니다.")

            temperature = st.slider(
                "기본 창의성 (Default Creativity)", 0.1, 1.0,
                value=float(st.session_state.setting_temperature),
                key="sb_temperature",
                on_change=lambda: [
                    st.session_state.update({"setting_temperature": st.session_state.sb_temperature}),
                    auto_save()
                ]
            )
            st.session_state.setting_temperature = temperature

            humor_level = st.slider(
                "기본 유머 감각 (Default Humor)", 0, 10,
                value=int(st.session_state.setting_humor),
                key="sb_humor",
                help="0: Serious, 10: Hilarious/Slapstick",
                on_change=lambda: [
                    st.session_state.update({"setting_humor": st.session_state.sb_humor}),
                    auto_save()
                ]
            )
            st.session_state.setting_humor = humor_level

            # 모델 선택 - 유효성 검증
            w_model = st.session_state.setting_model_writer
            if w_model not in writer_model_options:
                w_model = "gemini-2.5-pro"
                st.session_state.setting_model_writer = w_model
            default_w_idx = writer_model_options.index(w_model)

            a_model = st.session_state.setting_model_assistant
            if a_model not in assistant_model_options:
                a_model = "gemini-2.5-flash"
                st.session_state.setting_model_assistant = a_model
            default_a_idx = assistant_model_options.index(a_model)

            selected_model_writer = st.selectbox(
                "AI Writer (소설 집필용)",
                writer_model_options,
                index=default_w_idx,
                key="sb_model_writer",
                on_change=lambda: [
                    st.session_state.update({"setting_model_writer": st.session_state.sb_model_writer}),
                    auto_save()
                ]
            )
            st.session_state.setting_model_writer = selected_model_writer

            selected_model_assistant = st.selectbox(
                "AI Assistant (기획/분석/검수용)",
                assistant_model_options,
                index=default_a_idx,
                key="sb_model_assistant",
                on_change=lambda: [
                    st.session_state.update({"setting_model_assistant": st.session_state.sb_model_assistant}),
                    auto_save()
                ]
            )
            st.session_state.setting_model_assistant = selected_model_assistant

            st.markdown("---")
            st.markdown("✍️ **Writing Style**")
            style_options = ["기본", "웹소설체", "감성적", "담백한", "고전", "유머러스"]
            _style_idx = style_options.index(st.session_state.setting_style) if st.session_state.setting_style in style_options else 0
            selected_style = st.selectbox(
                "Style Preset", style_options,
                index=_style_idx,
                key="sb_style",
                on_change=lambda: [
                    st.session_state.update({"widget_setting_style": st.session_state.sb_style,
                                             "setting_style": st.session_state.sb_style}),
                    on_style_change()
                ]
            )
            st.session_state.setting_style = selected_style

            st.markdown("---")
            st.markdown("🧑‍💻 **Persona**")
            persona_presets = {
                "Direct Input": "",
                "김은숙 st": "김은숙 작가 스타일로, 인물 간의 대사가 빠르고 재치 있게...",
                "지브리 st": "지브리 애니메이션처럼, 서정적이고 아름답게...",
                "박찬욱 st": "박찬욱 감독 스타일로, 우아하고 잔혹하게...",
                "막장 드라마": "아침 드라마처럼, 자극적이고 빠르게...",
            }

            _preset_idx = list(persona_presets.keys()).index(st.session_state.setting_preset) if st.session_state.setting_preset in persona_presets else 0
            sel_preset = st.selectbox(
                "Persona Inspiration", list(persona_presets.keys()),
                index=_preset_idx,
                key="sb_preset",
                on_change=lambda: [
                    st.session_state.update({"widget_setting_preset": st.session_state.sb_preset,
                                             "setting_preset": st.session_state.sb_preset}),
                    on_preset_change()
                ]
            )
            st.session_state.setting_preset = sel_preset

            def update_persona():
                 st.session_state.custom_persona_input = st.session_state.custom_persona_ui
                 auto_save()

            custom_persona = st.text_input("Custom Persona", value=st.session_state.custom_persona_input, key="custom_persona_ui", on_change=update_persona)

            st.markdown("📅 **Structure**")
            _target_vols = st.number_input(
                "Target Volumes (목표 권수)",
                min_value=1,
                max_value=100,
                value=int(st.session_state.setting_target_vols),
                key="sb_target_vols",
                on_change=lambda: [
                    st.session_state.update({"setting_target_vols": st.session_state.sb_target_vols}),
                    auto_save()
                ]
            )
            st.session_state.setting_target_vols = _target_vols

            _target_chapters = st.number_input(
                "Chapters per Volume (1권당 목표 화수)",
                min_value=1,
                max_value=200,
                value=int(st.session_state.setting_target_chapters),
                key="sb_target_chapters",
                on_change=lambda: [
                    st.session_state.update({"setting_target_chapters": st.session_state.sb_target_chapters}),
                    auto_save()
                ]
            )
            st.session_state.setting_target_chapters = _target_chapters
            st.markdown("---")
            
            # RAG Settings Controls
            st.markdown("🔍 **RAG Settings (Supabase pgvector)**")
            with st.expander("RAG 설정 열기", expanded=True):
                rag_enabled = st.checkbox("RAG 참조 활성화", value=st.session_state.rag_enabled, key="rag_enabled_ui")
                if st.session_state.rag_enabled != rag_enabled:
                    st.session_state.rag_enabled = rag_enabled
                    auto_save()
                    
                # Categories (ebook_t)
                categories = st.session_state.get("rag_categories", [])
                if not categories:
                    categories = fetch_rag_categories()
                    st.session_state.rag_categories = categories
                cat_options = ["전체 (선택 안 함)"] + [c["cd_tname"] for c in categories]
                
                # Find index of current selected category
                selected_cat_name = "전체 (선택 안 함)"
                selected_cat_id = st.session_state.get("rag_category_id")
                if selected_cat_id is not None:
                    for c in categories:
                        if c["cd_t"] == selected_cat_id:
                            selected_cat_name = c["cd_tname"]
                            break
                
                cat_index = cat_options.index(selected_cat_name) if selected_cat_name in cat_options else 0
                
                # Category selectbox
                sel_cat_name = st.selectbox("장르/카테고리 선택", cat_options, index=cat_index, key="rag_cat_ui")
                
                # Handle category selection change
                new_cat_id = None
                if sel_cat_name != "전체 (선택 안 함)":
                    for c in categories:
                        if c["cd_tname"] == sel_cat_name:
                            new_cat_id = c["cd_t"]
                            break
                
                # If category changed, refresh the series list
                if new_cat_id != st.session_state.rag_category_id:
                    st.session_state.rag_category_id = new_cat_id
                    st.session_state.rag_series = fetch_rag_series(category_id=new_cat_id)
                    st.session_state.rag_series_id = None # Reset series selection
                    st.session_state.rag_series_search_val = "" # Reset search query
                    auto_save()
                    st.rerun()
                    
                # Series search text input
                search_query_val = st.text_input("도서 제목 검색 (엔터 키 입력)", value=st.session_state.get("rag_series_search_val", ""), key="rag_series_search_ui")
                
                # Detect search query change
                if search_query_val != st.session_state.get("rag_series_search_val", ""):
                    st.session_state.rag_series_search_val = search_query_val
                    
                    # Fetch matching series
                    matched_series = fetch_rag_series(
                        category_id=st.session_state.rag_category_id,
                        search_query=search_query_val
                    )
                    
                    # If specific series is currently selected, preserve it in the options
                    if st.session_state.rag_series_id is not None:
                        already_in = any(s["series_id"] == st.session_state.rag_series_id for s in matched_series)
                        if not already_in:
                            saved_series_list = fetch_rag_series(series_id=st.session_state.rag_series_id)
                            matched_series.extend(saved_series_list)
                            
                    st.session_state.rag_series = matched_series
                    st.rerun()

                # Series list for selectbox options
                series_list = st.session_state.get("rag_series", [])
                series_options = ["전체 (선택 안 함)"] + [s["display_title"] for s in series_list]
                
                selected_series_title = "전체 (선택 안 함)"
                selected_series_id = st.session_state.get("rag_series_id")
                if selected_series_id is not None:
                    for s in series_list:
                        if s["series_id"] == selected_series_id:
                            selected_series_title = s["display_title"]
                            break
                            
                series_index = series_options.index(selected_series_title) if selected_series_title in series_options else 0
                
                sel_series_title = st.selectbox("도서/시리즈 선택", series_options, index=series_index, key="rag_series_ui")
                
                new_series_id = None
                if sel_series_title != "전체 (선택 안 함)":
                    for s in series_list:
                        if s["display_title"] == sel_series_title:
                            new_series_id = s["series_id"]
                            break
                
                if new_series_id != st.session_state.rag_series_id:
                    st.session_state.rag_series_id = new_series_id
                    auto_save()
                    
                # Keyword filter (matches tags or content)
                kw_val = st.text_input("특정 키워드 필터 (태그)", value=st.session_state.rag_keyword, key="rag_keyword_ui")
                if kw_val != st.session_state.rag_keyword:
                    st.session_state.rag_keyword = kw_val
                    auto_save()
                    
            st.markdown("---")
            with st.expander("🧠 장기 기억 (이전 회차 요약)", expanded=False):
                if not st.session_state.get("memory_chain"):
                    st.info("아직 저장된 기억이 없어요.")
                else:
                    for m in st.session_state.memory_chain:
                        ch_num = m.get('chapter', '?')
                        summary_text = m.get('chunk_summary', m.get('summary', ''))
                        ch_updates = m.get('entity_changes', m.get('entity_updates', {}))
                        if not ch_updates:
                            ch_updates = {}
                        
                        char_val = ch_updates.get('characters', '')
                        if isinstance(char_val, dict):
                            char_desc = " / ".join([f"{k}: {v}" for k, v in char_val.items()]) if char_val else "변동 없음"
                        else:
                            char_desc = str(char_val) if char_val else "변동 없음"
                            
                        setting_val = ch_updates.get('settings', ch_updates.get('world', ''))
                        if isinstance(setting_val, dict):
                            setting_desc = " / ".join([f"{k}: {v}" for k, v in setting_val.items()]) if setting_val else "변동 없음"
                        else:
                            setting_desc = str(setting_val) if setting_val else "변동 없음"
                            
                        cliff_str = f"<br><b>🔗 클리프행어</b>: {m.get('cliffhanger_point')}" if m.get('cliffhanger_point') else ""
                        
                        details_html = f"""
                        <details style="margin-bottom: 10px; padding: 8px; border: 1px solid #dfe1e5; border-radius: 6px; background-color: rgba(255,255,255,0.05);">
                            <summary style="cursor: pointer; font-weight: 600; font-size: 0.95rem; list-style-type: none;">📖 제 {ch_num}화 요약 팩</summary>
                            <div style="margin-top: 8px; font-size: 0.88rem; line-height: 1.5; color: #d0d3d4;">
                                <b>📝 요약</b>: {summary_text}<br>
                                <b>👥 인물변화</b>: {char_desc}<br>
                                <b>🌍 설정변화</b>: {setting_desc}
                                {cliff_str}
                            </div>
                        </details>
                        """
                        st.markdown(details_html, unsafe_allow_html=True)
                
                if st.button("현재 본문 요약 저장", key="save_mem"):
                    with st.spinner("요약 중..."):
                        try:
                            res = requests.post(
                                f"{BACKEND_URL}/summarize",
                                json={
                                    "text": st.session_state.get("current_story", ""),
                                    "chapter_num": st.session_state.get("current_chapter_idx", 1),
                                    "model": st.session_state.get("setting_model_assistant", "gemini-2.5-flash")
                                },
                                timeout=60
                            )
                            if res.status_code == 200:
                                sum_data = res.json()
                                ch_num = sum_data.get("chapter_num", st.session_state.get("current_chapter_idx", 1))
                                chain = st.session_state.get("memory_chain", [])
                                chain = [m for m in chain if int(m.get("chapter", 0)) != int(ch_num)]
                                chain.append(sum_data)
                                chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                st.session_state.memory_chain = chain
                                auto_save()
                                st.rerun()
                            else:
                                st.error(res.text)
                        except Exception as e:
                            st.error(e)
                
                if st.button("기억 초기화", key="clear_mem"):
                    st.session_state.memory_chain = []
                    auto_save()
                    st.rerun()

            # --- 전체 화차 네비게이션 및 진척도 표시 (사이드바 하단 추가) ---
            st.markdown("---")
            st.markdown("### 📖 전체 화차 진행 현황")
            planned_chaps = int(st.session_state.get("setting_target_chapters", 50))
            available_chaps = sorted([int(k) for k in st.session_state.get("chapters", {}).keys()]) if st.session_state.get("chapters") else [1]
            total_chaps = max(planned_chaps, max(available_chaps) if available_chaps else 1)
            current_idx = st.session_state.get("current_chapter_idx", 1)
            
            # 1. 진척도 프로그레스바
            progress_val = float(current_idx) / float(total_chaps) if total_chaps > 0 else 0.0
            st.progress(min(progress_val, 1.0))
            st.caption(f"**현재 위치**: 제 {current_idx}화 / 전체 {total_chaps}화")
            
            # 2. 이전화/다음화 간편 이동 버튼
            col_nav_1, col_nav_2 = st.columns(2)
            with col_nav_1:
                if st.button("◀ 이전 화", use_container_width=True, disabled=(current_idx <= 1)):
                    auto_save()
                    st.session_state.current_chapter_idx = current_idx - 1
                    st.session_state.current_story = st.session_state.chapters.get(str(current_idx - 1), "")
                    st.session_state.story_editor_ui = st.session_state.current_story
                    st.session_state.last_prompt = ""
                    if "next_prompt_options" in st.session_state:
                        del st.session_state.next_prompt_options
                    auto_save()
                    st.rerun()
            with col_nav_2:
                if st.button("다음 화 ▶", use_container_width=True, disabled=(current_idx >= total_chaps)):
                    auto_save()
                    st.session_state.current_chapter_idx = current_idx + 1
                    st.session_state.current_story = st.session_state.chapters.get(str(current_idx + 1), "")
                    st.session_state.story_editor_ui = st.session_state.current_story
                    st.session_state.last_prompt = ""
                    if "next_prompt_options" in st.session_state:
                        del st.session_state.next_prompt_options
                    auto_save()
            # 3. 직접 회차 이동 입력 필드
            def jump_to_chapter():
                target_ch = st.session_state.jump_chapter_input
                if target_ch in available_chaps:
                    auto_save()
                    st.session_state.current_chapter_idx = target_ch
                    st.session_state.current_story = st.session_state.chapters.get(str(target_ch), "")
                    st.session_state.story_editor_ui = st.session_state.current_story
                    st.session_state.last_prompt = ""
                    if "next_prompt_options" in st.session_state:
                        del st.session_state.next_prompt_options
                    auto_save()
                else:
                    st.toast(f"제 {target_ch}화는 아직 구성되지 않은 회차입니다.", icon="⚠️")
                    
            st.number_input(
                "🎯 직접 회차 이동 (번호 입력 후 엔터)",
                min_value=1,
                max_value=total_chaps if total_chaps > 0 else 1,
                value=current_idx,
                step=1,
                key="jump_chapter_input",
                on_change=jump_to_chapter
            )
            st.markdown("<br>", unsafe_allow_html=True)

            if st.button("📖 초보자용 가이드 열기 (팝업)", use_container_width=True):
                st.session_state.show_guide = True
                st.rerun()

    # ---------- Main UI Styles & Divider ----------
    st.markdown("""
<style>
/* Import Outfit Google Font */
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');

/* Apply font globally */
html, body, [class*="css"], .stApp {
    font-family: 'Outfit', 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
}

/* Premium Gradient Title Style */
.main-title {
    font-size: 2.6rem !important;
    font-weight: 700 !important;
    background: linear-gradient(135deg, #FF4B4B 0%, #FF85A2 50%, #9F55FF 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 0.8rem !important;
    padding-bottom: 0.2rem !important;
    letter-spacing: -0.5px;
}

/* Visual separation line between header and content */
.header-divider {
    height: 1.5px;
    background: linear-gradient(90deg, transparent, rgba(255, 75, 75, 0.4) 15%, rgba(159, 85, 255, 0.4) 85%, transparent);
    margin: 15px 0 25px 0;
}

/* Distinct styling for sub-titles */
.section-title {
    font-size: 1.45rem !important;
    font-weight: 700 !important;
    color: #FF85A2 !important;
    margin-top: 10px !important;
    margin-bottom: 6px !important;
    border-left: 5px solid #FF4B4B;
    padding-left: 12px !important;
    letter-spacing: -0.3px;
}

.section-subtitle {
    font-size: 0.95rem !important;
    color: #B2B2B2 !important;
    margin-bottom: 22px !important;
    padding-left: 17px !important;
}

/* Premium Tab Navigation Styling */
div[data-baseweb="tab-list"] {
    gap: 12px !important;
    background-color: rgba(255, 255, 255, 0.02) !important;
    padding: 8px !important;
    border-radius: 12px !important;
    border: 1px solid rgba(255, 255, 255, 0.05) !important;
    margin-bottom: 20px !important;
}

button[data-baseweb="tab"] {
    border-radius: 8px !important;
    padding: 12px 22px !important;
    font-size: 1.05rem !important;
    transition: all 0.25s ease !important;
    background-color: rgba(0, 0, 0, 0.03) !important;
    border: 1px solid rgba(0, 0, 0, 0.05) !important;
}

/* Base style for all tab buttons and their inner text elements (Black & Thin) */
button[data-baseweb="tab"], 
button[data-baseweb="tab"] p, 
button[data-baseweb="tab"] span, 
button[data-baseweb="tab"] div {
    color: #000000 !important; /* Elegant Black */
    font-weight: 300 !important; /* Thin/Slender Font style */
    font-size: 1.05rem !important;
}

/* Hover style */
button[data-baseweb="tab"]:hover,
button[data-baseweb="tab"]:hover p,
button[data-baseweb="tab"]:hover span,
button[data-baseweb="tab"]:hover div {
    color: #FF4B4B !important; /* Highlight red on hover */
    background-color: rgba(0, 0, 0, 0.05) !important;
}

/* Active / Selected style */
button[aria-selected="true"] {
    background-color: rgba(255, 75, 75, 0.12) !important;
    border: 1px solid rgba(255, 75, 75, 0.25) !important;
    box-shadow: 0 4px 12px rgba(255, 75, 75, 0.08) !important;
}

button[aria-selected="true"],
button[aria-selected="true"] p,
button[aria-selected="true"] span,
button[aria-selected="true"] div {
    color: #FF4B4B !important; /* Distinct reddish-pink for active selection */
    font-weight: 400 !important;
}
</style>
""", unsafe_allow_html=True)

    if st.session_state.current_page == "📦 Publisher Hub":
        st.markdown('<div class="main-title">📦 Publisher Hub (출판 허브)</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="main-title">💖 Romance AI: {st.session_state.current_project}</div>', unsafe_allow_html=True)
    
    pages = ["✍️ Story Engine", "🏭 Novel Factory (Batch)", "🧐 Editor's Desk", "🎨 Art Studio", "📦 Publisher Hub"]
    selected_page = st.radio(
        "Navigation",
        pages,
        index=pages.index(st.session_state.current_page) if st.session_state.current_page in pages else 0,
        horizontal=True,
        label_visibility="collapsed",
        key="main_navigation_radio"
    )
    if selected_page != st.session_state.current_page:
        st.session_state.current_page = selected_page
        st.rerun()

    # Render Divider Line to separate top header from bottom content
    st.markdown('<div class="header-divider"></div>', unsafe_allow_html=True)

    # ---------- Global Beginner's Guide Modal (Popup Container) ----------
    if st.session_state.get("show_guide", False):
        st.markdown("""
        <div style="background-color: rgba(255, 75, 75, 0.08); border: 2px solid rgba(255, 75, 75, 0.3); padding: 25px; border-radius: 15px; margin-bottom: 25px; box-shadow: 0 8px 32px rgba(0,0,0,0.15); backdrop-filter: blur(8px);">
            <h2 style="color: #FF4B4B; margin-top: 0; display: flex; align-items: center; gap: 10px;">
                📖 Romance AI 초보자 가이드 & 사용 설명서
            </h2>
            <p style="color: #E2E2E2; font-size: 0.98rem; margin-bottom: 15px; line-height: 1.6;">
                처음 오신 작가님을 위한 종합 설명서입니다. 아래 가이드를 넓은 화면으로 천천히 읽어보신 후, 글쓰기를 시작해 보세요!
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        with st.container(border=True):
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("""
### 1️⃣ 좌측 사이드바 설정 (기본 세팅)
* **⚙️ 프로젝트 기본값 설정**: 
  - **창의성 (Creativity)**: 값이 높을수록 AI가 더 의외성 있고 새로운 전개로 소설을 쓰며, 낮을수록 정통 로맨스 정석 문장으로 씁니다. (기본 0.7 권장)
  - **유머 감각**: 티키타카 대사나 코믹 상황 강도를 0(진지함)부터 10(슬랩스틱)까지 조절합니다.
  - **AI Writer & Assistant**: 본문을 집필할 모델(Pro 등급 권장)과 기획/분석용 모델(Flash 등급 권장)을 선택합니다.
  - **Style Preset**: 기본, 웹소설체, 감성적, 담백한, 고전, 유머러스 중 원하는 문체를 선택하면 스타일 가이드가 자동으로 매핑됩니다.
  - **Custom Persona**: 특정 작가(김은숙, 지브리 등) 스타일이나 자신만의 페르소나 지침을 입력합니다.
  - **Structure**: 목표 권수와 권당 화수를 지정합니다.
* **🔍 RAG Settings (지식 데이터베이스)**:
  - Supabase 데이터베이스 내 기출판된 장르/시리즈/도서를 연계해 RAG 기반 참조를 켜거나 끌 수 있습니다.
* **🧠 장기 기억 (Memory Chain)**:
  - 이전 화차들의 요약 팩이 누적 보관되는 곳입니다. 장편 소설 집필 시 줄거리 설정 붕괴를 방지해 줍니다.
  - **[현재 본문 요약 저장]**을 누르면 방금 쓴 화차의 요약이 저장되어 다음 화차 집필 시 반영됩니다.
* **📖 전체 화차 진행 현황**:
  - 현재 집필 중인 화차 위치를 보여주며, 이전/다음 화 버튼을 통해 쉽게 회차를 변경할 수 있습니다.
""")
            with col2:
                st.markdown("""
### 2️⃣ 메인 메뉴 & 글 쓰는 순서 가이드
프로젝트 매니저에서 **Create New**로 새 프로젝트를 생성하고 로드한 뒤, 상단 탭을 왼쪽에서 오른쪽으로 넘어가며 사용하세요.

##### 1. [✍️ Story Engine] (스토리 기획 & 집필)
* **1단계: 📚 Story Bible & Inspiration (세계관 / 캐릭터 / 영감)**
  - `Character Sheet`에 인물 관계도와 캐릭터 성격을 입력합니다.
  - `World Setting`에 공간, 규칙, 마법 등 소설 배경을 적습니다.
  - *영감이 필요하다면?* 하단의 `✨ Need Inspiration?`에서 장르와 키워드를 골라 **[Generate Idea & Synopsis]**를 클릭해 시놉시스를 생성하세요.
* **2단계: 📝 Plot & Sync (전체 뼈대 기획)**
  - 시놉시스를 바탕으로 **[Generate Plot Outline]**을 클릭하면, 50화 분량의 각 화별 에피소드 아웃라인이 자동 설계됩니다.
* **3단계: ✍️ Chapter Editor (회차별 본문 집필)**
  - 현재 화차의 시작 문장을 입력하거나 빈 상태에서 **[Generate / Continue]**를 누르면 AI가 소설을 이어 씁니다.
  - 하단의 `✨ Next Action Suggestions`를 누르면 AI가 다음에 이어질 스토리 방향 3가지를 제안해 줍니다.
* **4단계: 📤 Export & Publish (원고 저장)**
  - 작성된 원고 전체를 `EPUB` 책 파일이나 `TXT`로 한 번에 다운로드합니다.

##### 2. [🏭 Novel Factory] (자동 소설 공장)
* 기획(Plot)된 아웃라인을 바탕으로 **50화 전체 분량을 클릭 한 번으로 일괄 자동 집필**하는 자동화 팩토리입니다.

##### 3. [🧐 Editor's Desk] (검토 및 인공지능 교정)
* **[Run Deep Analysis]**를 누른 뒤 비평을 확인하고, **[Generate Fix based on Critique]**를 누르면 피드백을 반영한 문장을 다시 써서 보여줍니다. 마음에 든다면 **[Apply Auto-Fix]**로 본문을 교정하세요.

##### 4. [🎨 Art Studio] (Imagen 3 표지 제작)
* **[소설 본문에서 이미지 프롬프트 자동 추출]**을 눌러 영어 묘사를 생성한 뒤, **[표지 이미지 생성하기]**를 클릭해 Imagen 3 엔진으로 표지를 만듭니다.

##### 5. [📦 Publisher Hub] (외부 원고 가져오기)
* 이미 보유 중인 대용량 텍스트 파일(TXT, EPUB)을 업로드하여 문장 끊김 없이 일정한 글자 수나 회차 수로 자동 분할(스마트 분할기)할 수 있습니다.
""")
            st.markdown("---")
            if st.button("❌ 가이드 닫기", type="primary", use_container_width=True):
                st.session_state.show_guide = False
                st.rerun()
            st.markdown("---")


    # ==========================================
    # TAB 1: STORY ENGINE (Classic)
    # ==========================================
    if st.session_state.current_page == "✍️ Story Engine":
        # Create sub-tabs to organize the workflow into a logical pipeline
        engine_tab_bible, engine_tab_plot, engine_tab_editor, engine_tab_export = st.tabs([
            "📚 Story Bible & Inspiration (세계관 / 아이디어 / 출판)",
            "📝 Plot & Sync (뼈대 기획 / 설정 동기화)",
            "✍️ Chapter Editor (회차별 집필)",
            "📤 Export & Publish (원고 내려받기)"
        ])

        # ==========================================
        # SUB-TAB 3: STORY BIBLE & INSPIRATION (세계관 / 아이디어)
        # ==========================================
        with engine_tab_bible:
            st.markdown('<div class="section-title">📚 소설 기본 설정 (Story Bible & Idea Generator)</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-subtitle">스토리의 핵심 뼈대를 구축하기 위한 인물 정보(Character Sheet) 및 세계관 배경(World Setting)을 등록하고 영감을 다듬습니다.</div>', unsafe_allow_html=True)


            # Story Bible inputs
            c1_bible, c2_bible = st.columns(2)
            with c1_bible:
                st.text_area(
                    "Character Sheet (인물 관계도 및 캐릭터 설명)",
                    value=st.session_state.get("char_sheet", ""),
                    key="char_sheet_input_tab",
                    height=350,
                    placeholder="예시:\n- 김철수: 24세, 까칠한 천재 해커. 과거의 상처가 있음.\n- 이영희: 28세, 열정적인 형사.",
                    on_change=lambda: [st.session_state.update({"char_sheet": st.session_state.char_sheet_input_tab}), auto_save()],
                )
            with c2_bible:
                st.text_area(
                    "World Setting (세계관 및 배경 설정)",
                    value=st.session_state.get("world_setting", ""),
                    key="world_setting_input_tab",
                    height=350,
                    placeholder="예시:\n- 플로팅 아일랜드: 하늘에 떠 있는 거대 섬.\n- 마법 설정: 왕족만 별의 마법을 씀.",
                    on_change=lambda: [st.session_state.update({"world_setting": st.session_state.world_setting_input_tab}), auto_save()],
                )

            # Inspiration (Idea Generator)
            st.markdown("---")
            st.markdown("### 🎭 Magic Idea Generator (아이디어 창고)")

            # ── 근본 해결: widget_* 간접 참조 제거 ──────────────────────────────
            # Streamlit은 위젯이 화면에 없으면 widget_* 키를 session_state에서 자동 삭제.
            # 따라서 widget_* key 방식 대신 value= 직접 바인딩 + on_change에서 ig_* 직접 저장.
            # 탭/페이지 이동 후 돌아와도 ig_* 세션 값은 절대 사라지지 않음.

            # 세션 기본값 보장 (최초 1회만)
            if "ig_genre" not in st.session_state:
                st.session_state.ig_genre = "전통로맨스"
            if "ig_spice" not in st.session_state:
                st.session_state.ig_spice = "19금(없음)"
            if "ig_moods" not in st.session_state:
                st.session_state.ig_moods = []
            if "ig_male" not in st.session_state:
                st.session_state.ig_male = []
            if "ig_female" not in st.session_state:
                st.session_state.ig_female = []
            if "ig_arc" not in st.session_state:
                st.session_state.ig_arc = ""
            if "ig_style_guide" not in st.session_state:
                st.session_state.ig_style_guide = "빠른 전개,\n짧은 문장,\n짧은 문단,\n대화 중심,\n강한 훅(hook),\n대사 55%,\n행동 25%,\n내면 15%,\n배경 설명 5%"

            c1_id, c2_id, c3_id = st.columns(3)
            with c1_id:
                genre_options = [
                    "전통로맨스", "사극(하)로맨스", "사극(중)로맨스", "사극(상)로맨스",
                    "현대로맨스", "판타지(약)로맨스", "판타지(중)로맨스", "판타지(강)로맨스",
                ]
                _genre_idx = genre_options.index(st.session_state.ig_genre) if st.session_state.ig_genre in genre_options else 0
                selected_genre = st.selectbox(
                    "Genre (장르)", genre_options,
                    index=_genre_idx,
                    key="ig_genre_selectbox",
                    on_change=lambda: [
                        st.session_state.update({"ig_genre": st.session_state.ig_genre_selectbox}),
                        auto_save()
                    ]
                )
                st.session_state.ig_genre = selected_genre
            with c2_id:
                spice_options = ["19금(없음)", "19금(하)", "19금(중)", "19금(상)"]
                _spice_idx = spice_options.index(st.session_state.ig_spice) if st.session_state.ig_spice in spice_options else 0
                selected_spice = st.selectbox(
                    "Spice Level (수위)", spice_options,
                    index=_spice_idx,
                    key="ig_spice_selectbox",
                    on_change=lambda: [
                        st.session_state.update({"ig_spice": st.session_state.ig_spice_selectbox}),
                        auto_save()
                    ]
                )
                st.session_state.ig_spice = selected_spice
            with c3_id:
                pass

            mood_tags = [
                "감성코드", "격정멜로", "금지된사랑", "달달물", "로맨틱", "막장드라마",
                "반전남녀", "순수남녀", "신파", "악녀시점", "애잔물", "위기탈출",
                "위험한사랑", "육아물", "잔잔물", "질투물", "케미커플", "티격태격",
                "피폐물", "하드코어", "힐링",
            ]
            selected_moods = st.multiselect(
                "Mood & Atmosphere (분위기)", mood_tags,
                default=st.session_state.ig_moods,
                key="ig_moods_multiselect",
                on_change=lambda: [
                    st.session_state.update({"ig_moods": st.session_state.ig_moods_multiselect}),
                    auto_save()
                ]
            )
            st.session_state.ig_moods = selected_moods

            male_tags = [
                "개천용", "거만남", "계략남", "군인", "그리스인", "까칠남", "나쁜남자", "냉혹남",
                "뇌섹남", "능글남", "다정남", "대형견남", "동정남", "라틴남", "러시아인",
                "마피아/범죄자", "목장주", "미소년", "바람둥이", "법조인", "병약남", "보디가드",
                "사기꾼", "사별남", "사이다남", "상처남", "소방관", "순정남", "시크남",
                "아랍인(세이크)", "애교남", "언론인", "연예인남", "연하남", "영국인", "오만남",
                "촌사람", "왕족/귀족", "외국인남", "요섹남", "운동선수", "의료업", "이탈리아인",
                "이혼남", "인기남", "재벌남", "전남친", "절륜남", "존댓말남", "직진남", "진중남",
                "짐승남", "짝사랑남", "차도남", "천재", "철벽남", "초식남", "카리스마남",
                "카우보이", "평범", "프랑스인", "후계자", "후회남", "훈남"
            ]
            selected_male = st.multiselect(
                "Male Lead (남자주인공)", male_tags,
                default=st.session_state.ig_male,
                key="ig_male_multiselect",
                on_change=lambda: [
                    st.session_state.update({"ig_male": st.session_state.ig_male_multiselect}),
                    auto_save()
                ]
            )
            st.session_state.ig_male = selected_male

            female_tags = [
                "4차원/엉뚱녀", "가정부/메이드", "건어물녀", "걸크러시", "결혼식들러리", "계략녀",
                "귀여운여인", "금지옥엽", "기자", "까칠녀", "꽃미녀", "남장여자", "뇌섹녀",
                "능글녀", "당당/당찬녀", "도도녀/무심녀", "동정녀", "디자이너", "라틴녀", "모델",
                "몰락재벌집 딸", "미망인", "미혼모", "백치미(둔녀)", "베이비시터", "병약녀", "비서",
                "사이다녀", "상처녀", "생활고여주", "서비스업", "신비녀", "악녀", "애교녀",
                "연예인녀", "영국인", "왕족/귀족", "외유내강녀", "웨딩플래너", "의료업", "이혼녀",
                "자상녀", "재벌녀/상속녀", "전부인", "절륜녀", "직진녀", "짝사랑녀", "차도녀",
                "철벽녀", "청순녀/순진녀", "친절녀", "카우걸", "캔디", "커리어우먼", "터프녀",
                "털털녀", "파티셰", "평범", "환골탈태녀", "후회녀"
            ]
            selected_female = st.multiselect(
                "Female Lead (여자주인공)", female_tags,
                default=st.session_state.ig_female,
                key="ig_female_multiselect",
                on_change=lambda: [
                    st.session_state.update({"ig_female": st.session_state.ig_female_multiselect}),
                    auto_save()
                ]
            )
            st.session_state.ig_female = selected_female

            char_arc = st.text_input(
                "Character Arc (관계 변화)",
                value=st.session_state.ig_arc,
                placeholder="예: 혐관 -> 찐사랑",
                key="ig_arc_input",
                on_change=lambda: [
                    st.session_state.update({"ig_arc": st.session_state.ig_arc_input}),
                    auto_save()
                ]
            )
            st.session_state.ig_arc = char_arc

            style_guide = st.text_area(
                "스타일 가이드:",
                value=st.session_state.ig_style_guide,
                height=120,
                key="ig_style_guide_textarea",
                on_change=lambda: [
                    st.session_state.update({"ig_style_guide": st.session_state.ig_style_guide_textarea}),
                    auto_save()
                ]
            )
            st.session_state.ig_style_guide = style_guide
            
            if st.button("🎲 Generate Random Story Concept"):
                with st.spinner("Brainstorming creative ideas..."):
                    try:
                        res = requests.post(
                            f"{BACKEND_URL}/generate/idea",
                            json={
                                "genre": selected_genre,
                                "spice_level": selected_spice,
                                "model": selected_model_assistant,
                                "apply_trends": True,
                                "moods": selected_moods,
                                "male_tags": selected_male,
                                "female_tags": selected_female,
                                "arc": char_arc,
                                "char_sheet": st.session_state.get("char_sheet", ""),
                                "world_setting": st.session_state.get("world_setting", "")
                            }, timeout=120,
                        )
                        if res.status_code == 200:
                            st.session_state.idea_suggestion = res.json().get("idea")
                            auto_save()
                            st.rerun()
                        else:
                            st.error(f"Error: {res.text}")
                    except Exception as e:
                        st.error(f"Connection Error: {e}")

            if st.session_state.get("idea_suggestion"):
                st.text_area(
                    "Suggested Setup (Editable):",
                    value=st.session_state.idea_suggestion,
                    height=200,
                    key="idea_suggestion_input_tab",
                    on_change=lambda: [st.session_state.update({"idea_suggestion": st.session_state.idea_suggestion_input_tab}), auto_save()]
                )

            # Marketing Metadata Packaging
            st.markdown("---")
            st.markdown("### 🏷️ Book Packaging (마케팅 타이틀 & 책 소개 생성)")
            st.caption("집필된 전체 소설 원고를 기반으로 플랫폼에 출판할 때 적합한 제목 후보군, 강렬한 소개글(Blurb) 및 키워드 해시태그를 자동 생성합니다.")
            if st.button("✨ Generate Metadata"):
                # Determine best source text (current_story, plot_outline, or idea_suggestion)
                source_text = ""
                source_type = ""
                
                if st.session_state.get("current_story") and len(st.session_state.current_story.strip()) >= 500:
                    source_text = st.session_state.current_story.strip()
                    source_type = "소설 본문"
                elif st.session_state.get("plot_outline") and len(st.session_state.plot_outline.strip()) >= 100:
                    source_text = st.session_state.plot_outline.strip()
                    source_type = "플롯 개요(전체 뼈대)"
                elif st.session_state.get("idea_suggestion") and len(st.session_state.idea_suggestion.strip()) >= 100:
                    source_text = st.session_state.idea_suggestion.strip()
                    source_type = "시놉시스(아이디어)"
                
                if not source_text:
                    st.error("마케팅 메타데이터를 분석 생성하기 위해 본문(최소 500자), 플롯 개요(전체 뼈대), 또는 시놉시스(아이디어) 중 하나가 입력되어 있어야 합니다.")
                else:
                    with st.spinner(f"{source_type} 분석 및 북 패키징(책 소개 & 제목) 생성 중..."):
                        settings_payload = {
                            "genre": st.session_state.get("ig_genre", ""),
                            "mood": ", ".join(st.session_state.get("ig_moods", [])),
                            "trends": st.session_state.get("ig_style_guide", ""),
                            "characters": st.session_state.get("char_sheet", ""),
                            "world": st.session_state.get("world_setting", "")
                        }
                        try:
                            res = requests.post(
                                f"{BACKEND_URL}/generate/packaging", 
                                json={
                                    "settings": settings_payload, 
                                    "outline": source_text,
                                    "model": selected_model_assistant
                                }, timeout=120
                            )
                            if res.status_code == 200:
                                pkg = res.json()
                                st.session_state.pkg_titles = pkg.get("titles", [])
                                st.session_state.pkg_blurb = pkg.get("blurb", "")
                                st.session_state.pkg_keywords = pkg.get("keywords", [])
                                auto_save()
                                st.rerun()
                            else:
                                st.error(f"Error: {res.text}")
                        except Exception as e:
                            st.error(f"Analysis Error: {e}")

            if st.session_state.get("pkg_titles"):
                st.write("#### 📢 추천 제목 (Suggested Titles)")
                for t in st.session_state.pkg_titles:
                    st.write(f"- {t}")
                st.write("#### 📝 책 소개글 (Blurb)")
                st.info(st.session_state.pkg_blurb)
                st.write("#### 🏷️ 검색 해시태그 키워드")
                st.caption(" ".join(st.session_state.pkg_keywords))

            # Export & Publish downloads (Moved to tab 3 to declutter editor)
        # ==========================================
        # SUB-TAB 2: PLOT & SYNC (뼈대 기획 / 설정 동기화)
        # ==========================================
        with engine_tab_plot:
            st.markdown('<div class="section-title">📝 전체 뼈대 기획 및 설정 동기화 (Plot & Sync)</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-subtitle">소설 집필의 토대가 되는 50화 아웃라인을 생성하고, 인물/세계관 설정에 모순이 없도록 동기화 검증 및 흥행 요소를 다듬는 공간입니다.</div>', unsafe_allow_html=True)

            
            st.markdown("#### 1. Plot Generator (플롯 기획)")
            apply_styles_to_plot = st.checkbox(
                "⚙️ Apply Writer Settings (Style Preset, Humor, Persona) to Plot", 
                value=st.session_state.apply_plot_style, 
                key="apply_plot_style_checkbox", 
                on_change=lambda: [st.session_state.update({"apply_plot_style": st.session_state.apply_plot_style_checkbox}), auto_save()]
            )
            
            if st.button("Generate Plot Outline"):
                if not st.session_state.get("ig_genre"):
                    st.error("Please select a Genre in the Idea Generator first (📚 Story Bible & Inspiration 탭에서 설정).")
                else:
                    with st.spinner("Designing plot..."):
                        settings_payload = {
                            "genre": st.session_state.get("ig_genre", ""),
                            "spice": st.session_state.get("ig_spice", ""),
                            "mood": ", ".join(st.session_state.get("ig_moods", [])),
                            "chars": st.session_state.get("char_sheet", ""),
                            "world": st.session_state.get("world_setting", ""),
                            "arc": st.session_state.get("ig_arc", ""),
                            "trends": st.session_state.get("ig_style_guide", ""),
                            "idea_premise": st.session_state.get("idea_suggestion", ""),
                            "target_chapters": st.session_state.get("setting_target_chapters", 50)
                        }
                        
                        if apply_styles_to_plot:
                            settings_payload["style"] = selected_style
                            settings_payload["persona"] = st.session_state.custom_persona_input
                            settings_payload["humor_level"] = st.session_state.setting_humor
                        
                        try:
                            res = requests.post(f"{BACKEND_URL}/analyze/plot", json={"settings": settings_payload, "model": selected_model_assistant}, timeout=300)
                            if res.status_code == 200:
                                plot_text = res.json().get("plot", "")
                                st.session_state.plot_outline = plot_text
                                
                                # Populate chapters structure
                                matches = re.findall(r'(?:제\s*(\d+)\s*화|Chapter\s*(\d+))', plot_text, re.IGNORECASE)
                                found_indices = []
                                for m in matches:
                                    idx_str = m[0] or m[1]
                                    if idx_str:
                                        found_indices.append(int(idx_str))
                                
                                target_chaps_limit = int(st.session_state.get("setting_target_chapters", 50))
                                if not found_indices or len(found_indices) < 2:
                                    found_indices = list(range(1, target_chaps_limit + 1))
                                else:
                                    found_indices = sorted(list(set(found_indices)))
                                    if len(found_indices) < target_chaps_limit:
                                        for idx in range(1, target_chaps_limit + 1):
                                            if idx not in found_indices:
                                                found_indices.append(idx)
                                        found_indices = sorted(found_indices)
                                
                                # Preserve existing chapter text if any to prevent data loss
                                old_chapters = st.session_state.get("chapters", {})
                                st.session_state.chapters = {str(i): old_chapters.get(str(i), "") for i in found_indices}
                                st.session_state.current_chapter_idx = found_indices[0] if found_indices else 1
                                st.session_state.current_story = st.session_state.chapters.get(str(st.session_state.current_chapter_idx), "")
                                st.session_state.story_editor_ui = st.session_state.current_story
                                auto_save()
                                st.success("플롯 기반 화차가 성공적으로 빌드되었습니다!")
                                st.rerun()
                            else:
                                st.error(f"Error: {res.text}")
                        except Exception as e:
                            st.error(f"Connection Error: {e}")

            st.session_state.plot_outline = st.text_area(
                "Plot Outline (전체 뼈대)", 
                value=st.session_state.get("plot_outline", ""),
                height=300,
                key="plot_outline_input_tab",
                on_change=lambda: [st.session_state.update({"plot_outline": st.session_state.plot_outline_input_tab}), auto_save()]
            )

            # --- 🔄 Sync Outline & Bible ---
            st.markdown("#### 2. Settings Synchronization (설정 동기화 검증)")
            st.caption("인물 캐릭터 시트, 세계관 설정과 줄거리 아웃라인을 교차 검토하여 설정 모순이 있는 사건 전개를 자동으로 동기화합니다.")
            if st.button("🔄 바뀐 설정에 맞게 줄거리/시놉시스 동기화 검증", use_container_width=True):
                with st.spinner("설정 모순 분석 및 동기화된 플롯 제안서 생성 중..."):
                    try:
                        res_sync = requests.post(
                            f"{BACKEND_URL}/analyze/consistency/sync-outline",
                            json={
                                "char_sheet": st.session_state.get("char_sheet", ""),
                                "world_setting": st.session_state.get("world_setting", ""),
                                "blurb": st.session_state.get("pkg_blurb", ""),
                                "plot_outline": st.session_state.get("plot_outline", ""),
                                "model": selected_model_assistant
                            }, timeout=300
                        )
                        if res_sync.status_code == 200:
                            st.session_state.proposed_sync_resolution = res_sync.json()
                            st.rerun()
                        else:
                            try:
                                err_msg = res_sync.json().get("detail", "동기화 분석 실패")
                            except Exception:
                                err_msg = res_sync.text
                            st.error(f"⚠️ **설정 동기화 검증 실패**\n\n{err_msg}\n\n*기존 줄거리 데이터는 변경되거나 삭제되지 않고 안전하게 보존되었습니다.*")
                    except Exception as e:
                        st.error(f"⚠️ **서버 통신 에러**: {str(e)}\n\n*기존 줄거리 데이터는 변경되거나 삭제되지 않고 안전하게 보존되었습니다.*")

            if st.session_state.get("proposed_sync_resolution"):
                prop_sync = st.session_state.proposed_sync_resolution
                st.info("💡 **인물/세계관 설정 정렬된 동기화 제안** (사건 전개는 완전 보존)")
                
                sync_blurb = prop_sync.get("blurb_synced", "").strip()
                sync_outline = prop_sync.get("plot_outline_synced", "").strip()
                
                col_sync_view1, col_sync_view2 = st.columns(2)
                with col_sync_view1:
                    st.subheader("📖 동기화된 책 소개 (Blurb)")
                    st.text_area("Synced Blurb Preview", value=sync_blurb, height=180, disabled=True)
                with col_sync_view2:
                    st.subheader("📝 동기화된 플롯 개요 (Plot Outline)")
                    st.text_area("Synced Outline Preview", value=sync_outline, height=180, disabled=True)
                    
                col_sync_apply1, col_sync_apply2 = st.columns(2)
                with col_sync_apply1:
                    if st.button("💾 동기화 제안을 블러브/플롯에 최종 반영", type="primary", use_container_width=True):
                        st.session_state.pkg_blurb = sync_blurb
                        st.session_state.plot_outline = sync_outline
                        auto_save()
                        st.success("줄거리와 시놉시스 동기화 반영 및 저장이 완료되었습니다! 💾")
                        st.session_state.pop("proposed_sync_resolution", None)
                        st.rerun()
                with col_sync_apply2:
                    if st.button("❌ 동기화 제안 취소", key="cancel_sync_res", use_container_width=True):
                        st.session_state.pop("proposed_sync_resolution", None)
                        st.rerun()

            # --- 🔮 Plot Evaluator ---
            st.markdown("#### 3. Commercial Success Prediction (흥행성 분석)")
            st.caption("현재 기획한 50화 뼈대가 시장에서 갖는 상업적 잠재력과 독자 흡입력을 검증합니다.")
            if st.button("Analyze Plot Potential"):
                if not st.session_state.get("plot_outline"):
                     st.error("Please generate or write a Plot Outline first.")
                else:
                     with st.spinner("Analyzing Market Trends & Reader Psychology..."):
                         settings_payload = {
                            "genre": st.session_state.get("ig_genre", ""),
                            "mood": ", ".join(st.session_state.get("ig_moods", [])),
                            "characters": st.session_state.get("char_sheet", ""),
                            "world": st.session_state.get("world_setting", ""),
                            "trends": st.session_state.get("ig_style_guide", "")
                         }
                         try:
                             res = requests.post(
                                 f"{BACKEND_URL}/analyze/prediction", 
                                 json={
                                     "settings": settings_payload, 
                                     "outline": st.session_state.plot_outline,
                                     "model": selected_model_assistant
                                 }, timeout=120
                             )
                             if res.status_code == 200:
                                 st.session_state.prediction_report = res.json().get("report")
                                 auto_save()
                             else:
                                 st.error(f"Error: {res.text}")
                         except Exception as e:
                             st.error(f"Connection Error: {e}")

            if "prediction_report" in st.session_state:
                report = st.session_state.prediction_report
                if report and "error" not in report:
                    c1, c2 = st.columns(2)
                    with c1:
                        st.metric("💰 Commercial Success", f"{report.get('commercial_score', 0)}/100")
                        st.progress(report.get('commercial_score', 0) / 100)
                    with c2:
                        st.metric("🔥 Binge-Reading Factor", f"{report.get('binge_score', 0)}/100")
                        st.progress(report.get('binge_score', 0) / 100)
                    
                    st.info(f"🎯 **Target Audience**: {report.get('target_audience', {}).get('gender', 'Unknown')} in {report.get('target_audience', {}).get('age', 'Unknown')} (Buying Power: {report.get('target_audience', {}).get('buying_power', 'Unknown')})")
                    
                    st.markdown("#### 📊 SWOT Analysis")
                    swot = report.get('swot', {})
                    c1, c2 = st.columns(2)
                    c1.success(f"**Strengths**:\n" + "\n".join([f"- {x}" for x in swot.get('strengths', [])]))
                    c2.error(f"**Weaknesses**:\n" + "\n".join([f"- {x}" for x in swot.get('weaknesses', [])]))
                    c1.info(f"**Opportunities**:\n" + "\n".join([f"- {x}" for x in swot.get('opportunities', [])]))
                    c2.warning(f"**Threats**:\n" + "\n".join([f"- {x}" for x in swot.get('threats', [])]))

                    st.write("### 💡 Improvement Advice")
                    initial_advice = report.get('improvement_advice', 'No advice generated.')
                    user_advice = st.text_area(
                        "Review & Edit Improvement Instructions:", 
                        value=initial_advice,
                        height=150
                    )
                    
                    if st.button("✨ Auto-Improve Plot with these Instructions"):
                        with st.spinner("Refining plot structure..."):
                            try:
                                settings_payload = {
                                    "genre": st.session_state.get("ig_genre", ""),
                                    "mood": ", ".join(st.session_state.get("ig_moods", [])),
                                    "characters": st.session_state.get("char_sheet", ""),
                                    "world": st.session_state.get("world_setting", "")
                                }
                                res = requests.post(
                                    f"{BACKEND_URL}/analyze/improve_plot", 
                                    json={
                                        "settings": settings_payload, 
                                        "outline": st.session_state.plot_outline,
                                        "advice": user_advice,
                                        "model": selected_model_assistant
                                    }, timeout=300
                                )
                                if res.status_code == 200:
                                    st.session_state.plot_outline = res.json().get("plot", "")
                                    auto_save()
                                    st.success("Plot updated based on advice!")
                                    st.rerun()
                                else:
                                    try:
                                        err_msg = res.json().get("detail", res.text)
                                    except Exception:
                                        err_msg = res.text
                                    st.error(f"⚠️ **줄거리 개선 및 업데이트 실패**\n\n{err_msg}\n\n*기존 줄거리 데이터는 변경되거나 삭제되지 않고 안전하게 보존되었습니다.*")
                            except Exception as e:
                                st.error(f"⚠️ **서버 통신 에러**: {str(e)}\n\n*기존 줄거리 데이터는 변경되거나 삭제되지 않고 안전하게 보존되었습니다.*")

                    st.caption(f"Overall Review: {report.get('overall_review', '')}")

        # ==========================================
        # SUB-TAB 1: CHAPTER EDITOR
        # ==========================================
        with engine_tab_editor:
            col1, col2 = st.columns([2, 1])

            # ----- Left column – Chapter Editor CORE -----
            with col1:
                st.markdown('<div class="section-title">📖 Chapter Management</div>', unsafe_allow_html=True)

                
                # Get available chapters (numeric sorted)
                available_chapters = sorted([int(k) for k in st.session_state.chapters.keys()]) if st.session_state.get("chapters") else [1]
                chapter_options = [f"제 {i}화" for i in available_chapters] + ["📖 전체 본문 보기 (Full Draft)"]
                
                # Find current index
                curr_idx = st.session_state.get("current_chapter_idx", 1)
                selected_label = f"제 {curr_idx}화"
                if selected_label not in chapter_options:
                    default_idx = 0
                else:
                    default_idx = chapter_options.index(selected_label)
                
                def on_chapter_change():
                    sel = st.session_state.chapter_picker
                    if "전체 본문 보기" in sel:
                        st.session_state.view_all_mode = True
                    else:
                        st.session_state.view_all_mode = False
                        new_idx = int(re.search(r'\d+', sel).group())
                        auto_save()
                        st.session_state.current_chapter_idx = new_idx
                        st.session_state.current_story = st.session_state.chapters.get(str(new_idx), "")
                        st.session_state.story_editor_ui = st.session_state.current_story
                        st.session_state.last_prompt = ""
                        if "next_prompt_options" in st.session_state:
                            del st.session_state.next_prompt_options
                        auto_save()

                col_picker, col_del, col_reset = st.columns([2, 1, 1])
                with col_picker:
                    selected_chapter_ui = st.selectbox(
                        "편집할 화차를 선택하세요",
                        chapter_options,
                        index=default_idx,
                        key="chapter_picker",
                        on_change=on_chapter_change
                    )
                with col_del:
                    st.write("") 
                    st.write("")
                    if st.button("🗑️ 현재 화차 삭제", use_container_width=True, help="현재 선택된 화차를 삭제하고 챕터 순서를 당깁니다."):
                        curr_idx = st.session_state.current_chapter_idx
                        if "chapters" in st.session_state and str(curr_idx) in st.session_state.chapters:
                            del st.session_state.chapters[str(curr_idx)]
                            
                            # Re-index remaining chapters
                            old_chaps = st.session_state.chapters
                            new_chaps = {}
                            for new_i, old_k in enumerate(sorted([int(k) for k in old_chaps.keys()]), start=1):
                                new_chaps[str(new_i)] = old_chaps[str(old_k)]
                            st.session_state.chapters = new_chaps
                            
                            # ── 장기기억장치 (memory_chain) 수정 및 재색인 ────────────────
                            if "memory_chain" in st.session_state and st.session_state.memory_chain:
                                updated_chain = []
                                for m in st.session_state.memory_chain:
                                    if not isinstance(m, dict):
                                        continue
                                    try:
                                        ch_num = int(m.get("chapter", 0))
                                    except (ValueError, TypeError):
                                        ch_num = 0
                                    
                                    # 삭제 대상 화차는 제외
                                    if ch_num == curr_idx:
                                        continue
                                    # 삭제 대상 화차보다 뒤의 화차들은 번호를 1씩 감소
                                    elif ch_num > curr_idx:
                                        m["chapter"] = ch_num - 1
                                    updated_chain.append(m)
                                st.session_state.memory_chain = updated_chain
                                st.session_state.memory_chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                
                            # ── 연속성 원장 (continuity_ledger) 수정 및 재색인 ─────────────
                            if "continuity_ledger" in st.session_state and st.session_state.continuity_ledger:
                                updated_ledger = []
                                for l in st.session_state.continuity_ledger:
                                    if not isinstance(l, dict):
                                        continue
                                    try:
                                        ch_num = int(l.get("chapter", 0))
                                    except (ValueError, TypeError):
                                        ch_num = 0
                                        
                                    # 삭제 대상 화차는 제외
                                    if ch_num == curr_idx:
                                        continue
                                    # 삭제 대상 화차보다 뒤의 화차들은 번호를 1씩 감소
                                    elif ch_num > curr_idx:
                                        l["chapter"] = ch_num - 1
                                    updated_ledger.append(l)
                                st.session_state.continuity_ledger = updated_ledger
                                st.session_state.continuity_ledger.sort(key=lambda x: int(x.get("chapter", 0)))
                            
                            new_idx = max(1, curr_idx - 1)
                            if new_chaps and str(new_idx) not in new_chaps:
                                new_idx = max([int(k) for k in new_chaps.keys()])
                            st.session_state.current_chapter_idx = new_idx
                            st.session_state.current_story = new_chaps.get(str(new_idx), "")
                            
                            # 지침(Brief) 캐시 날림
                            if "chapter_brief" in st.session_state:
                                st.session_state.chapter_brief = ""
                                
                            auto_save()
                            st.success("화차 및 장기기억장치가 성공적으로 삭제 및 정렬되었습니다.")
                            st.rerun()
                with col_reset:
                    st.write("") 
                    st.write("")
                    if st.button("🧹 전체 초기화", use_container_width=True, help="전체 챕터 본문을 삭제하고 1화부터 다시 작성합니다."):
                        st.session_state.chapters = {}
                        st.session_state.current_chapter_idx = 1
                        st.session_state.current_story = ""
                        
                        # 장기기억장치 및 연속성 원장, 지침 캐시 전체 초기화
                        st.session_state.memory_chain = []
                        st.session_state.continuity_ledger = []
                        st.session_state.chapter_brief = ""
                        if "next_prompt_options" in st.session_state:
                            del st.session_state.next_prompt_options
                            
                        auto_save()
                        st.success("전체 집필 챕터 및 장기기억장치가 초기화되었습니다.")
                        st.rerun()

                # Build structure from plot outline button
                if st.session_state.get("plot_outline"):
                    if st.button("🛠️ 현재 Plot Outline 기반으로 화차 구조 생성/복원", use_container_width=True, help="기존 생성된 Plot Outline 또는 목표 화수 설정에 맞추어 전체 화차 구조를 즉시 자동 생성합니다."):
                        plot_text = st.session_state.plot_outline
                        matches = re.findall(r'(?:제\s*(\d+)\s*화|Chapter\s*(\d+))', plot_text, re.IGNORECASE)
                        found_indices = []
                        for m in matches:
                            idx_str = m[0] or m[1]
                            if idx_str:
                                found_indices.append(int(idx_str))
                        
                        target_chaps_limit = int(st.session_state.get("setting_target_chapters", 50))
                        
                        if not found_indices or len(found_indices) < 2:
                            found_indices = list(range(1, target_chaps_limit + 1))
                        else:
                            found_indices = sorted(list(set(found_indices)))
                            if len(found_indices) < target_chaps_limit:
                                for idx in range(1, target_chaps_limit + 1):
                                    if idx not in found_indices:
                                        found_indices.append(idx)
                                found_indices = sorted(found_indices)
                        
                        if "chapters" not in st.session_state or not st.session_state.chapters:
                            st.session_state.chapters = {}
                        for i in found_indices:
                            if str(i) not in st.session_state.chapters:
                                st.session_state.chapters[str(i)] = ""
                                
                        st.session_state.current_chapter_idx = found_indices[0] if found_indices else 1
                        st.session_state.current_story = st.session_state.chapters.get(str(st.session_state.current_chapter_idx), "")
                        st.session_state.story_editor_ui = st.session_state.current_story
                        auto_save()
                        st.success(f"총 {len(found_indices)}화 분량의 화차 구조가 성공적으로 구성되었습니다!")
                        st.rerun()

                if st.session_state.get("view_all_mode"):
                    st.markdown("---")
                    st.markdown("### 📜 전체 원고 미리보기")
                    all_text = ""
                    for i in sorted([int(k) for k in st.session_state.chapters.keys()]):
                        all_text += f"## Chapter {i}\n\n" + st.session_state.chapters.get(str(i), "") + "\n\n"
                    st.text_area("Read-only Full Story", value=all_text, height=600, disabled=True)
                    st.info("💡 개별 회차를 수정하려면 위 선택박스에서 해당 회차를 선택해 주세요.")
                else:
                    # Main Story Editor
                    col_editor_title, col_save_btn = st.columns([3, 1])
                    with col_editor_title:
                        st.markdown(f"### ✍️ 제 {st.session_state.current_chapter_idx}화 집필 중")
                    with col_save_btn:
                        if st.button("💾 프로젝트 저장", use_container_width=True, help="현재 수정한 텍스트 및 모든 설정을 즉시 디스크에 저장합니다."):
                            if "story_editor_ui" in st.session_state:
                                st.session_state.current_story = st.session_state.story_editor_ui
                            auto_save()
                            
                            # Automatically update memory chain summary for the edited chapter
                            curr_ch = st.session_state.get("current_chapter_idx", 1)
                            curr_text = st.session_state.get("current_story", "")
                            if curr_text.strip():
                                with st.spinner(f"제 {curr_ch}화 요약 정보 자동 갱신 중..."):
                                    try:
                                        r_ext = requests.post(
                                            f"{BACKEND_URL}/summarize",
                                            json={
                                                "text": curr_text,
                                                "chapter_num": curr_ch,
                                                "model": selected_model_assistant
                                            }, timeout=120
                                        )
                                        if r_ext.status_code == 200:
                                            res_data = r_ext.json()
                                            chain = st.session_state.get("memory_chain", [])
                                            # Remove duplicates
                                            chain = [m for m in chain if int(m.get("chapter", 0)) != int(curr_ch)]
                                            chain.append(res_data)
                                            chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                            st.session_state.memory_chain = chain
                                            # Register as applied fix if this chapter was in recommended list
                                            if "review_result" in st.session_state and isinstance(st.session_state.review_result, dict):
                                                if "applied_fixes" not in st.session_state.review_result or not isinstance(st.session_state.review_result["applied_fixes"], list):
                                                    st.session_state.review_result["applied_fixes"] = []
                                                if int(curr_ch) not in st.session_state.review_result["applied_fixes"]:
                                                    raw_recs = st.session_state.review_result.get("recommended_chapters", [])
                                                    is_rec = False
                                                    for item in raw_recs:
                                                        if isinstance(item, dict) and int(item.get("chapter", 0)) == int(curr_ch):
                                                            is_rec = True
                                                            break
                                                        elif isinstance(item, (int, str)) and str(item).isdigit() and int(item) == int(curr_ch):
                                                            is_rec = True
                                                            break
                                                    if is_rec:
                                                        st.session_state.review_result["applied_fixes"].append(int(curr_ch))
                                            auto_save()
                                            st.toast("프로젝트 저장 및 기억 요약 자동 갱신이 완료되었습니다! 💾")
                                        else:
                                            st.toast("프로젝트는 저장되었으나, 기억 요약 갱신에 실패했습니다.")
                                    except Exception as ext_err:
                                        print(f"Memory update failed: {ext_err}")
                                        st.toast("프로젝트는 저장되었으나, 네트워크 지연으로 요약 갱신을 생략합니다.")
                            else:
                                st.toast("프로젝트가 성공적으로 저장되었습니다! 💾")
                    
                    st.session_state.current_story = st.text_area(
                        "Story Text",
                        value=st.session_state.current_story,
                        height=500,
                        placeholder="여기에 소설 내용을 작성하거나 AI로 생성하세요...",
                        key="story_editor_ui",
                        on_change=auto_save
                    )
                    st.caption("💡 작성 중인 내용을 저장하려면 입력창 바깥을 누르거나 `Ctrl + Enter`를 누르세요. 혹은 우측 상단의 **[💾 프로젝트 저장]** 버튼을 클릭하세요.")

                    # --- Interactive Next Prompt Choices ---
                    st.markdown("### 🧭 향후 전개 제안받기")
                    current_chapter_idx = st.session_state.get("current_chapter_idx", 1)
                    
                    # Capture chapter plot
                    chapter_plot = ""
                    if st.session_state.get("plot_outline"):
                        outline_text = st.session_state.plot_outline
                        lines = outline_text.split('\n')
                        capture = False
                        for line in lines:
                            line_clean = line.replace(" ", "").lower()
                            if f"제{current_chapter_idx}화" in line_clean or f"chapter{current_chapter_idx}" in line_clean:
                                capture = True
                            elif capture and (f"제{current_chapter_idx+1}화" in line_clean or f"chapter{current_chapter_idx+1}" in line_clean):
                                capture = False
                                break
                            if capture:
                                chapter_plot += line + "\n"
                    
                    if not chapter_plot.strip():
                        chapter_plot = "현재 회차의 구체적인 플롯을 찾을 수 없습니다. 전체 흐름에 맞게 추천해 주세요."

                    if st.button(f"✨ [제{current_chapter_idx}화] 다음 전개 방향 3가지 추천받기"):
                        if not st.session_state.get("plot_outline"):
                            st.warning("먼저 Plot & Sync 탭에서 아웃라인(Plot Outline)을 생성해 주세요.")
                        else:
                            with st.spinner(f"제{current_chapter_idx}화 플롯을 분석하여 전개 방향을 고민 중입니다..."):
                                recent_mem_list = []
                                for m in st.session_state.get("memory_chain", [])[-3:]:
                                    ch_sum = m.get("chunk_summary", m.get("summary", ""))
                                    ch_updates = m.get("entity_changes", m.get("entity_updates", {}))
                                    ch_cliff = m.get("cliffhanger_point", "")
                                    recent_mem_list.append(
                                        f"제{m.get('chapter', '?')}화 요약: {ch_sum} | 인물변화: {ch_updates.get('characters', '')} | 설정변화: {ch_updates.get('settings', '')} | 클리프행어: {ch_cliff}"
                                    )
                                recent_mem = "\n".join(recent_mem_list)
                                
                                try:
                                    res = requests.post(
                                        f"{BACKEND_URL}/generate/next_prompts",
                                        json={
                                            "chapter_num": current_chapter_idx,
                                            "chapter_outline": chapter_plot,
                                            "recent_memory": recent_mem
                                        },
                                        timeout=80
                                    )
                                    if res.status_code == 200:
                                        st.session_state.next_prompt_options = res.json().get("options", {})
                                    else:
                                        st.error(f"Error: {res.text}")
                                except Exception as e:
                                    st.error(f"Connection Error: {e}")
                    
                    if st.session_state.get("next_prompt_options"):
                        st.info("💡 **원하는 전개 방향을 선택하세요:**")
                        for key, val in st.session_state.next_prompt_options.items():
                            with st.container():
                                st.markdown(f"**👉 {key}**")
                                st.write(val)
                                if st.button(f"'{key}' 선택하기", key=f"btn_opt_{key}"):
                                    st.session_state.last_prompt = val
                                    del st.session_state.next_prompt_options
                                    # ─────────────────────────────────────────────
                                    # [Proactive STEP A] Pre-Write Brief 자동 생성
                                    # 선택된 전개 방향을 기반으로 이번화 집필 지침 생성
                                    # ─────────────────────────────────────────────
                                    try:
                                        with st.spinner(f"📋 [STEP A] 제{current_chapter_idx}화 집필 지침 생성 중..."):
                                            brief_res = requests.post(
                                                f"{BACKEND_URL}/generate/chapter_brief",
                                                json={
                                                    "chapter_num": current_chapter_idx,
                                                    "selected_choice": val,
                                                    "char_sheet": st.session_state.get("char_sheet", ""),
                                                    "world_setting": st.session_state.get("world_setting", ""),
                                                    "plot_outline": st.session_state.get("plot_outline", ""),
                                                    "continuity_ledger": st.session_state.get("continuity_ledger", []),
                                                    "memory_chain": st.session_state.get("memory_chain", [])[-5:],
                                                },
                                                timeout=30
                                            )
                                            if brief_res.status_code == 200:
                                                brief_text = brief_res.json().get("brief", "")
                                                st.session_state.chapter_brief = brief_text
                                    except Exception as brief_err:
                                        # Fail-safe: Brief 실패해도 집필 계속
                                        print(f"[STEP A] Brief 생성 실패 (무시): {brief_err}")
                                        st.session_state.chapter_brief = ""
                                    auto_save()
                                    st.rerun()
                        st.divider()

                    # [Proactive] 현재 화의 집필 지침 표시 (STEP A 결과)
                    if st.session_state.get("chapter_brief"):
                        with st.expander(f"📋 [집필 지침] 제{current_chapter_idx}화 — AI 지침 (자동 생성됨)", expanded=False):
                            st.markdown(st.session_state.chapter_brief)
                            if st.button("🗑️ 지침 초기화", key="clear_brief_btn"):
                                st.session_state.chapter_brief = ""
                                st.rerun()

                    # Individual Chapter Customization
                    curr_idx_str = str(current_chapter_idx)
                    if "chapters_settings" not in st.session_state:
                        st.session_state.chapters_settings = {}
                    
                    chap_meta = st.session_state.chapters_settings.get(curr_idx_str, {})
                    default_temp = chap_meta.get("temperature", st.session_state.get("setting_temperature", 0.7))
                    default_humor = chap_meta.get("humor_level", st.session_state.get("setting_humor", 0))
                    
                    st.markdown(f"⚙️ **제 {current_chapter_idx}화 개별 집필 설정 (Chapter Style Customization)**")
                    col_slide1, col_slide2 = st.columns(2)
                    with col_slide1:
                        chap_temp = st.slider(
                            "창의성 (Creativity Temperature)",
                            0.1, 1.0,
                            value=float(default_temp),
                            key=f"chap_temp_{curr_idx_str}",
                            on_change=auto_save
                        )
                    with col_slide2:
                        chap_humor = st.slider(
                            "유머 감각 (Humor Level)",
                            0, 10,
                            value=int(default_humor),
                            key=f"chap_humor_{curr_idx_str}",
                            on_change=auto_save
                        )
                    
                    st.session_state.chapters_settings[curr_idx_str] = {
                        "temperature": chap_temp,
                        "humor_level": chap_humor
                    }

                    user_input = st.text_input("Next Prompt (What happens next?):", key="last_prompt", on_change=auto_save)
                    
                    c1_gen, c2_gen = st.columns([1, 1])
                    with c1_gen:
                        if st.button("Generate / Continue"):
                            if user_input:
                                with st.spinner("Writing..."):
                                    if "story_editor_ui" in st.session_state:
                                        del st.session_state.story_editor_ui
                                    new_text = generate_story(
                                        prompt=user_input, 
                                        temperature=chap_temp, 
                                        model=selected_model_writer, 
                                        style=selected_style, 
                                        persona=st.session_state.custom_persona_input,
                                        humor_level=chap_humor,
                                        chapter_num=current_chapter_idx,
                                        ch_focus=chapter_plot,
                                        plot_summary=st.session_state.plot_outline,
                                        writer_memo=st.session_state.get("writer_memo", ""),
                                        chapter_brief=st.session_state.get("chapter_brief", ""),
                                        continuity_ledger=st.session_state.get("continuity_ledger", [])
                                    )
                                    if new_text and not new_text.startswith("[Error"):
                                        if "chapters" not in st.session_state:
                                            st.session_state.chapters = {}

                                        # ──────────────────────────────────────────
                                        # [Proactive STEP B] Post-Write QC + Self-Heal
                                        # ──────────────────────────────────────────
                                        final_text = new_text
                                        try:
                                            with st.spinner("🔍 [STEP B] 연속성 QC 검사 중..."):
                                                qc_res = requests.post(
                                                    f"{BACKEND_URL}/generate/chapter_qc",
                                                    json={
                                                        "chapter_num": current_chapter_idx,
                                                        "chapter_text": new_text,
                                                        "chapter_brief": st.session_state.get("chapter_brief", ""),
                                                        "continuity_ledger": st.session_state.get("continuity_ledger", []),
                                                        "char_sheet": st.session_state.get("char_sheet", ""),
                                                    },
                                                    timeout=90
                                                )
                                                if qc_res.status_code == 200:
                                                    qc_data = qc_res.json()
                                                    if not qc_data.get("passed", True):
                                                        healed = qc_data.get("healed_text", "")
                                                        issues = qc_data.get("issues", [])
                                                        severity = qc_data.get("severity", "low")
                                                        if healed and len(healed) > 100:
                                                            final_text = healed
                                                            st.toast(
                                                                f"⚕️ [STEP B] QC 교정 적용 (severity: {severity}) — "
                                                                + ", ".join(issues[:2]),
                                                                icon="⚕️"
                                                            )
                                                        else:
                                                            # Heal 결과 없으면 원본 유지 + 경고
                                                            st.toast(
                                                                f"⚠️ [STEP B] 연속성 이슈 감지 (원본 유지): " + ", ".join(issues[:2]),
                                                                icon="⚠️"
                                                            )
                                        except Exception as qc_err:
                                            # Fail-safe: QC 실패해도 집필 계속
                                            print(f"[STEP B] QC 실패 (원본 유지): {qc_err}")

                                        st.session_state.chapters[str(current_chapter_idx)] = final_text
                                        st.session_state.current_story = final_text
                                        
                                        # Auto-summarize & update memory
                                        with st.spinner("이전 화차 완료: 요약을 자동으로 생성하여 장기 기억에 추가 중..."):
                                            try:
                                                res_sum = requests.post(
                                                    f"{BACKEND_URL}/summarize",
                                                    json={
                                                        "text": final_text,
                                                        "chapter_num": current_chapter_idx,
                                                        "model": selected_model_assistant
                                                    },
                                                    timeout=60
                                                )
                                                if res_sum.status_code == 200:
                                                    if "memory_chain" not in st.session_state:
                                                        st.session_state.memory_chain = []
                                                    st.session_state.memory_chain = [m for m in st.session_state.memory_chain if m.get("chapter") != current_chapter_idx]
                                                    st.session_state.memory_chain.append(res_sum.json())
                                                    st.session_state.memory_chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                            except:
                                                pass

                                        # ──────────────────────────────────────────
                                        # [Proactive STEP C] Continuity Ledger 업데이트
                                        # ──────────────────────────────────────────
                                        try:
                                            with st.spinner("📋 [STEP C] 연속성 원장 업데이트 중..."):
                                                ledger_res = requests.post(
                                                    f"{BACKEND_URL}/generate/chapter_ledger",
                                                    json={
                                                        "chapter_num": current_chapter_idx,
                                                        "chapter_text": final_text,
                                                        "existing_ledger": st.session_state.get("continuity_ledger", []),
                                                    },
                                                    timeout=30
                                                )
                                                if ledger_res.status_code == 200:
                                                    new_ledger_item = ledger_res.json().get("ledger_item", {})
                                                    if "continuity_ledger" not in st.session_state:
                                                        st.session_state.continuity_ledger = []
                                                    # 동일 화 기존 항목 교체
                                                    st.session_state.continuity_ledger = [
                                                        e for e in st.session_state.continuity_ledger
                                                        if e.get("chapter") != current_chapter_idx
                                                    ]
                                                    st.session_state.continuity_ledger.append(new_ledger_item)
                                                    # 화번호 순 정렬
                                                    st.session_state.continuity_ledger.sort(
                                                        key=lambda x: int(x.get("chapter", 0))
                                                    )
                                                    st.toast("📋 [STEP C] 연속성 원장 업데이트 완료", icon="📋")
                                        except Exception as ledger_err:
                                            # Fail-safe: Ledger 실패해도 저장 계속
                                            print(f"[STEP C] Ledger 업데이트 실패 (무시): {ledger_err}")

                                        # Brief 초기화 (이번화 집필 완료)
                                        st.session_state.chapter_brief = ""

                                        # Advance chapter index
                                        next_idx = current_chapter_idx + 1
                                        st.session_state.current_chapter_idx = next_idx
                                        st.session_state.current_story = st.session_state.chapters.get(str(next_idx), "")
                                        if "story_editor_ui" in st.session_state:
                                            del st.session_state.story_editor_ui
                                        st.session_state.clear_last_prompt_flag = True
                                        if "next_prompt_options" in st.session_state:
                                            del st.session_state.next_prompt_options
                                        
                                        auto_save()
                                        st.rerun()
                                    else:
                                        if new_text:
                                            st.error(new_text)
                                        else:
                                            st.error("오류: AI가 소설을 생성하지 못했습니다.")
                            else:
                                st.warning("⚠️ 다음 전개 방향을 입력해 주세요.")
                    with c2_gen:
                        if st.button("Clear Story"):
                            st.session_state.current_story = ""
                            auto_save()
                            st.rerun()

            # ----- Right column – AI tools -----
            with col2:
                st.subheader("AI Editor Tools")

                # RAG Polish
                with st.expander("💎 RAG 기반 스마트 문단 윤색 (Gemini Polish)", expanded=True):
                    st.caption("소설 본문의 각 문단을 DB 로맨스 스타일과 실시간 매칭하여 윤색합니다. ✨강조 문장✨이 있으면 부분 교정합니다.")
                    
                    story_content = st.session_state.current_story
                    if not story_content.strip():
                        st.info("먼저 본문을 생성해 주세요.")
                    else:
                        paragraphs = [p.strip() for p in story_content.split("\n\n") if p.strip()]
                        
                        for idx, para in enumerate(paragraphs):
                            if "✨" in para:
                                parts = para.split("✨")
                                if len(parts) >= 3:
                                    para = f"{parts[0]}<STYLE>{parts[1]}</STYLE>{''.join(parts[2:])}"
                            
                            is_highlighted = "<STYLE>" in para or "<style>" in para
                            st.markdown(f"**문단 {idx + 1}**")
                            clean_para = para.replace("<STYLE>", "✨").replace("</STYLE>", "✨")
                            
                            if is_highlighted:
                                st.warning("💡 **스타일 교정 추천 구간:**\n\n" + clean_para)
                            else:
                                st.info(clean_para)
                            
                            col_btn1, col_btn2 = st.columns([1, 1])
                            with col_btn1:
                                if st.button(f"✨ 교정안 생성", key=f"polish_btn_{idx}"):
                                    with st.spinner("RAG 스타일 분석 및 교정안 생성 중..."):
                                        target_text_para = para
                                        surgical_data = None
                                        
                                        if "<STYLE>" in para and "</STYLE>" in para:
                                            try:
                                                match = re.search(r'(.*?)<STYLE>(.*?)</STYLE>(.*)', para, re.DOTALL)
                                                if match:
                                                    prefix, marked, suffix = match.groups()
                                                    target_text_para = marked.strip()
                                                    surgical_data = {"prefix": prefix, "suffix": suffix}
                                                    st.session_state[f"surgical_info_{idx}"] = surgical_data
                                            except:
                                                pass
                                        
                                        rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                                        payload = {
                                            "paragraph": target_text_para,
                                            "text": target_text_para,
                                            "model": selected_model_assistant,
                                            "rag_enabled": rag_enabled,
                                            "rag_category_id": rag_category_id,
                                            "rag_series_id": rag_series_id,
                                            "rag_keyword": rag_keyword
                                        }
                                        payload = {k: v for k, v in payload.items() if v is not None}
                                        
                                        try:
                                            res = requests.post(f"{BACKEND_URL}/generate/polish", json=payload, timeout=120)
                                            if res.status_code == 200:
                                                st.session_state[f"polish_options_{idx}"] = res.json().get("options", {})
                                            else:
                                                st.error(f"서버 오류: {res.text}")
                                        except Exception as e:
                                            st.error(f"연결 오류: {e}")
                            
                            opt_key = f"polish_options_{idx}"
                            if opt_key in st.session_state:
                                st.markdown("---")
                                opts = st.session_state[opt_key]
                                for label, text in opts.items():
                                    with st.container():
                                        st.markdown(f"**{label}**")
                                        st.write(text)
                                        if st.button("이 버전으로 교체하기", key=f"apply_{idx}_{label}"):
                                            s_info = st.session_state.get(f"surgical_info_{idx}")
                                            if s_info:
                                                new_para = f"{s_info['prefix']}<STYLE>{text}</STYLE>{s_info['suffix']}"
                                                paragraphs[idx] = new_para
                                                del st.session_state[f"surgical_info_{idx}"]
                                            else:
                                                paragraphs[idx] = text
                                                
                                            st.session_state.current_story = "\n\n".join(paragraphs)
                                            del st.session_state[opt_key]
                                            auto_save()
                                            st.success("교체 완료!")
                                            st.rerun()
                                if st.button("닫기", key=f"close_opt_{idx}"):
                                    del st.session_state[opt_key]
                                    st.rerun()
                            st.divider()

                # Critique & Auto-Fix
                st.markdown("### 🕵️‍♂️ Auto-Editor")
                current_num = st.session_state.get("current_chapter_idx", 1)
                st.markdown(f"📍 **분석 대상**: 제 {current_num}화 (현재 집필 중인 본문)")
                target_text = st.session_state.get("current_story", "")
                if not target_text.strip() and "chapters" in st.session_state:
                    target_text = st.session_state.chapters.get(str(current_num), "")

                # Spell Checker
                st.markdown("### ✍️ Spell Checker (맞춤법 검사)")
                if st.button("Run Spell Check", key="run_spell_check_story_engine"):
                    if not target_text.strip():
                        st.warning("검사할 본문 내용이 비어 있습니다.")
                    else:
                        with st.spinner("맞춤법 및 문장 검사 중..."):
                            try:
                                res = requests.post(
                                    f"{BACKEND_URL}/publisher/check-spell",
                                    json={"text": target_text, "model": selected_model_assistant},
                                    timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state.story_engine_spell_report = res.json().get("report", "")
                                    auto_save()
                                    st.rerun()
                                else:
                                    st.error(res.text)
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                if st.session_state.get("story_engine_spell_report"):
                    with st.expander("맞춤법 및 문장 교정 결과", expanded=True):
                        st.markdown(st.session_state.story_engine_spell_report)
                        if st.button("결과 닫기", key="close_spell_report"):
                            del st.session_state.story_engine_spell_report
                            st.rerun()
                st.markdown("---")

                # Sliding Window Context
                prev_summaries = [m for m in st.session_state.get("memory_chain", []) if int(m.get("chapter", 0)) < current_num][-2:]
                sliding_context_list = []
                for m in prev_summaries:
                    ch_sum = m.get('chunk_summary', m.get('summary', ''))
                    ch_updates = m.get('entity_changes', m.get('entity_updates', {}))
                    ch_cliff = m.get('cliffhanger_point', '')
                    sliding_context_list.append(
                        f"[제{m.get('chapter')}화 요약]: {ch_sum}\n"
                        f"[인물 변동]: {ch_updates.get('characters', '')}\n"
                        f"[설정 변동]: {ch_updates.get('settings', '')}\n"
                        f"[클리프행어]: {ch_cliff}"
                    )
                sliding_context = "\n\n".join(sliding_context_list)

                if st.button("Run Deep Analysis", key="run_deep_analysis_auto_editor"):
                    if not target_text.strip():
                        st.warning("분석할 본문 내용이 비어 있습니다.")
                    else:
                        st.session_state.analysis_target_text = target_text
                        with st.spinner(f"제{current_num}화 분석 중 ({len(target_text)}자)..."):
                            try:
                                analysis_prompt = target_text
                                if sliding_context:
                                    analysis_prompt = f"--- [이전 화들 정보] ---\n{sliding_context}\n\n--- [분석 대상 본문] ---\n{target_text}"
                                
                                res = requests.post(
                                    f"{BACKEND_URL}/analyze/review",
                                    json={"text": analysis_prompt, "model": selected_model_assistant}, timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state.critique_json = res.json()
                                    scores = st.session_state.critique_json.get("scores", {})
                                    feedback = st.session_state.critique_json.get("feedback", {})
                                    display_text = f"**Scores**: Consistency {scores.get('consistency')}, Creativity {scores.get('creativity')}\n\n"
                                    display_text += f"**Feedback**: {feedback}\n\n"
                                    display_text += f"**Overall**: {st.session_state.critique_json.get('overall_critique')}"
                                    st.session_state.critique = display_text
                                    auto_save()
                                else:
                                    st.error(res.text)
                            except Exception as e:
                                st.error(f"Analysis Error: {e}")

                if st.session_state.get("critique"):
                    with st.expander("View Analysis Report", expanded=True):
                        st.write(st.session_state.critique)
                        
                        if st.session_state.get("critique_json"):
                            st.markdown("---")
                            st.caption("✨ AI can rewrite this chapter based on the report above.")
                            if st.button("✨ Auto-Fix Chapter (Self-Healing)"):
                                 target_text = st.session_state.get("analysis_target_text", "")
                                 if not target_text:
                                     st.error("No analysis target found.")
                                 else:
                                     with st.spinner("Healing the story..."):
                                        try:
                                            rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                                            res = requests.post(
                                                f"{BACKEND_URL}/generate/rewrite",
                                                json={
                                                    "text": target_text, 
                                                    "critique": st.session_state.critique_json,
                                                    "model": selected_model_assistant,
                                                    "chars": st.session_state.get("char_sheet", "기본 인물"),
                                                    "world": st.session_state.get("world_setting", "기본 세계관"),
                                                    "style_guide": st.session_state.get("ig_style_guide", ""),
                                                    "rag_enabled": rag_enabled,
                                                    "rag_category_id": rag_category_id,
                                                    "rag_series_id": rag_series_id,
                                                    "rag_keyword": rag_keyword
                                                }, timeout=3600
                                            )
                                            if res.status_code == 200:
                                                new_text = res.json().get("rewritten_text")
                                                if new_text and len(new_text) > 100:
                                                    clean_new = new_text.strip()
                                                    terminal_punctuations = ('.', '!', '?', '”', '"', '…', '~', '*', '>', '}')
                                                    is_truncated = False
                                                    if clean_new and not clean_new.endswith(terminal_punctuations):
                                                        is_truncated = True
                                                    if len(target_text) > 500 and len(clean_new) < len(target_text) * 0.4:
                                                        is_truncated = True
                                                        
                                                    if is_truncated:
                                                        st.error("⚠️ AI가 생성한 문장이 중간에 잘린 것으로 감지되었습니다. 원본 보존을 위해 반영하지 않았습니다.")
                                                    else:
                                                        st.session_state.auto_editor_proposed_fix = new_text
                                                        st.success("✨ 교정안 생성 완료! 아래 검토 패널에서 확인하세요.")
                                                        st.rerun()
                                            else:
                                                st.error(res.text)
                                        except Exception as e:
                                            st.error(f"Rewrite Error: {e}")

                # Preview Proposed Auto-Fix
                if st.session_state.get("auto_editor_proposed_fix"):
                    st.markdown("---")
                    st.markdown("### 🔄 교정안 비교 및 최종 검토")
                    st.info("수정본을 확인 및 편집하신 뒤 [교정안 반영] 버튼을 눌러 적용해 주세요.")
                    
                    with st.expander("원래 작성한 본문 보기", expanded=False):
                        st.text_area("원본", value=st.session_state.get("analysis_target_text", ""), height=200, disabled=True, key="auto_editor_orig_preview")
                    
                    proposed_text = st.text_area("AI 수정 제안 (본문 덮어쓰기 전 편집 가능)", value=st.session_state.auto_editor_proposed_fix, height=350, key="auto_editor_fix_preview")
                    
                    col_fix_1, col_fix_2 = st.columns(2)
                    with col_fix_1:
                        if st.button("✅ 교정안 반영", key="auto_editor_accept"):
                            st.session_state.chapters[str(current_num)] = proposed_text
                            st.session_state.current_story = proposed_text
                            
                            # Summarize and update Memory Chain
                            with st.spinner("교정된 본문의 초정밀 요약을 자동 갱신 중..."):
                                try:
                                    sum_res = requests.post(
                                        f"{BACKEND_URL}/summarize",
                                        json={
                                            "text": proposed_text,
                                            "chapter_num": current_num,
                                            "model": selected_model_assistant
                                        }, timeout=60
                                    )
                                    if sum_res.status_code == 200:
                                        if "memory_chain" not in st.session_state:
                                            st.session_state.memory_chain = []
                                        st.session_state.memory_chain = [m for m in st.session_state.memory_chain if m.get("chapter") != current_num]
                                        st.session_state.memory_chain.append(sum_res.json())
                                        st.session_state.memory_chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                except Exception as e:
                                    st.warning(f"장기 기억 요약 갱신 실패: {e}")
                                    
                            del st.session_state.auto_editor_proposed_fix
                            if "critique" in st.session_state:
                                del st.session_state.critique
                            auto_save()
                            st.success("교정이 본문에 성공적으로 반영되었습니다! 💾")

                            # ──────────────────────────────────────────
                            # [Proactive STEP C] Auto-Editor 반영 후 Ledger 업데이트
                            # ──────────────────────────────────────────
                            try:
                                ledger_res = requests.post(
                                    f"{BACKEND_URL}/generate/chapter_ledger",
                                    json={
                                        "chapter_num": current_num,
                                        "chapter_text": proposed_text,
                                        "existing_ledger": st.session_state.get("continuity_ledger", []),
                                    },
                                    timeout=30
                                )
                                if ledger_res.status_code == 200:
                                    new_ledger_item = ledger_res.json().get("ledger_item", {})
                                    if "continuity_ledger" not in st.session_state:
                                        st.session_state.continuity_ledger = []
                                    st.session_state.continuity_ledger = [
                                        e for e in st.session_state.continuity_ledger
                                        if e.get("chapter") != current_num
                                    ]
                                    st.session_state.continuity_ledger.append(new_ledger_item)
                                    st.session_state.continuity_ledger.sort(
                                        key=lambda x: int(x.get("chapter", 0))
                                    )
                                    auto_save()
                            except Exception as ledger_err:
                                print(f"[STEP C/Auto-Editor] Ledger 업데이트 실패 (무시): {ledger_err}")

                            st.rerun()
                    with col_fix_2:
                        if st.button("❌ 취소", key="auto_editor_cancel"):
                            del st.session_state.auto_editor_proposed_fix
                            st.rerun()

                # Consistency check
                if st.button("🕵️ Consistency Check"):
                    if st.session_state.current_story:
                        with st.spinner("Checking..."):
                            try:
                                enriched_chars = st.session_state.char_sheet
                                enriched_world = st.session_state.world_setting
                                current_num = st.session_state.get("current_chapter_idx", 1)
                                prev_memories = [m for m in st.session_state.get("memory_chain", []) if int(m.get("chapter", 0)) < current_num]
                                
                                if prev_memories:
                                    enriched_chars += "\n\n[이전 화차들 인물 변동]"
                                    enriched_world += "\n\n[이전 화차들 설정/배경 변동]"
                                    for m in prev_memories:
                                        ch_num = m.get("chapter", 0)
                                        ch_sum = m.get("chunk_summary", "")
                                        ch_updates = m.get("entity_changes", {})
                                        ch_cliff = m.get("cliffhanger_point", "")
                                        enriched_chars += f"\n- 제{ch_num}화: {ch_updates.get('characters', '')} (요약: {ch_sum})"
                                        enriched_world += f"\n- 제{ch_num}화: {ch_updates.get('settings', '')} (클리프행어: {ch_cliff})"

                                res = requests.post(
                                    f"{BACKEND_URL}/analyze/consistency",
                                    json={
                                        "text": st.session_state.current_story,
                                        "char_sheet": enriched_chars,
                                        "world_setting": enriched_world,
                                        "model": selected_model_assistant,
                                    }, timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state.consistency_report = res.json().get("report", {})
                                    st.session_state.pop("proposed_settings_resolution", None)
                                    st.session_state.pop("proposed_story_resolution", None)
                                    st.rerun()
                                else:
                                    st.error(res.text)
                            except Exception as e:
                                st.error(f"Conn Error: {e}")

                if st.session_state.get("consistency_report"):
                    report = st.session_state.consistency_report
                    if report.get("name_errors"):
                        st.error(f"⚠️ 인물명/호칭 오류: {report['name_errors']}")
                    else:
                        st.success("✅ 인물명 및 호칭 일관성 확인 완료")
                        
                    if report.get("plot_errors"):
                        st.warning(f"⚠️ 설정/스토리 모순 감지: {report['plot_errors']}")
                        st.markdown("### 🛠️ 모순 자동 조치")
                        col_act1, col_act2 = st.columns(2)
                        with col_act1:
                            if st.button("⚙️ Option A: 설정 자동 업데이트", use_container_width=True):
                                with st.spinner("설정 보완 제안 생성 중..."):
                                    try:
                                        res_set = requests.post(
                                            f"{BACKEND_URL}/analyze/consistency/resolve-setting",
                                            json={
                                                "char_sheet": st.session_state.char_sheet,
                                                "world_setting": st.session_state.world_setting,
                                                "plot_errors": report.get("plot_errors", []),
                                                "model": selected_model_assistant
                                            }, timeout=120
                                        )
                                        if res_set.status_code == 200:
                                            st.session_state.proposed_settings_resolution = res_set.json()
                                            st.session_state.pop("proposed_story_resolution", None)
                                            st.rerun()
                                        else:
                                            st.error("설정 제안 생성 실패")
                                    except Exception as e:
                                        st.error(f"통신 에러: {e}")
                                        
                        with col_act2:
                            if st.button("✍️ Option B: 本문 자동 교정", use_container_width=True):
                                with st.spinner("본문 교정안 생성 중..."):
                                    try:
                                        rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                                        style_guide = st.session_state.get("ig_style_guide", "")
                                        res_story = requests.post(
                                            f"{BACKEND_URL}/analyze/consistency/resolve-story",
                                            json={
                                                "text": st.session_state.current_story,
                                                "char_sheet": st.session_state.char_sheet,
                                                "world_setting": st.session_state.world_setting,
                                                "plot_errors": report.get("plot_errors", []),
                                                "style_guide": style_guide,
                                                "rag_context": rag_keyword,
                                                "model": selected_model_writer
                                            }, timeout=150
                                        )
                                        if res_story.status_code == 200:
                                            st.session_state.proposed_story_resolution = res_story.json().get("revised_text", "")
                                            st.session_state.pop("proposed_settings_resolution", None)
                                            st.rerun()
                                        else:
                                            st.error("본문 교정안 생성 실패")
                                    except Exception as e:
                                        st.error(f"통신 에러: {e}")

                    if st.session_state.get("proposed_settings_resolution"):
                        prop = st.session_state.proposed_settings_resolution
                        st.info("💡 **인물/세계관 설정 보완 제안**")
                        c_merged = prop.get("char_sheet_merged", "").strip()
                        w_merged = prop.get("world_setting_merged", "").strip()
                        
                        if c_merged:
                            st.text_area("보완 통합된 캐릭터 시트", value=c_merged, height=200, disabled=True)
                        if w_merged:
                            st.text_area("보완 통합된 세계관 설정", value=w_merged, height=200, disabled=True)
                        
                        col_set_apply1, col_set_apply2 = st.columns(2)
                        with col_set_apply1:
                            if st.button("💾 설정을 시트에 자동 반영", type="primary"):
                                if c_merged:
                                    st.session_state.char_sheet = c_merged
                                if w_merged:
                                    st.session_state.world_setting = w_merged
                                auto_save()
                                st.success("설정이 완전히 반영되었습니다.")
                                st.session_state.pop("proposed_settings_resolution", None)
                                st.session_state.pop("consistency_report", None)
                                st.rerun()
                        with col_set_apply2:
                            if st.button("❌ 제안 취소", key="cancel_set_res"):
                                st.session_state.pop("proposed_settings_resolution", None)
                                st.rerun()

                    if st.session_state.get("proposed_story_resolution"):
                        revised = st.session_state.proposed_story_resolution
                        st.info("💡 **본문 모순 해결 교정안**")
                        st.text_area("수정된 본문 프리뷰", value=revised, height=300)
                        
                        col_story_apply1, col_story_apply2 = st.columns(2)
                        with col_story_apply1:
                            if st.button("💾 교정안을 본문에 자동 반영", type="primary"):
                                st.session_state.current_story = revised
                                auto_save()
                                st.success("본문이 교정되었습니다.")
                                st.session_state.pop("proposed_story_resolution", None)
                                st.session_state.pop("consistency_report", None)
                                st.rerun()
                        with col_story_apply2:
                            if st.button("❌ 제안 취소", key="cancel_story_res"):
                                st.session_state.pop("proposed_story_resolution", None)
                                st.rerun()
                    st.divider()

        # ==========================================
        # SUB-TAB 4: EXPORT & PUBLISH (원고 내려받기)
        # ==========================================
        with engine_tab_export:
            st.markdown('<div class="section-title">📦 Export & Publish (원고 내려받기)</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-subtitle">완성된 회차들을 하나의 통합 원고 또는 EPUB/TXT 파일로 조립하여 안전하게 내려받을 수 있는 도구입니다.</div>', unsafe_allow_html=True)

            # Define smart split settings path for persistence
            smart_dir = os.path.join(BASE_DATA_DIR, username, st.session_state.current_project)
            smart_settings_path = os.path.join(smart_dir, "smart_split_settings.json")

            def load_smart_split_settings(p_path):
                import json
                if os.path.exists(p_path):
                    try:
                        with open(p_path, "r", encoding="utf-8") as f:
                            return json.load(f)
                    except Exception:
                        pass
                return None

            def save_smart_split_settings(p_path, data):
                import json
                try:
                    os.makedirs(os.path.dirname(p_path), exist_ok=True)
                    with open(p_path, "w", encoding="utf-8") as f:
                        json.dump(data, f, ensure_ascii=False, indent=2)
                except Exception:
                    pass

            # Restore recommendations if not loaded for current project
            if st.session_state.get("smart_split_project") != st.session_state.current_project:
                saved_recs = load_smart_split_settings(smart_settings_path)
                if saved_recs:
                    st.session_state.smart_split_recommendations = saved_recs
                else:
                    st.session_state.pop("smart_split_recommendations", None)
                st.session_state.smart_split_project = st.session_state.current_project

            if st.session_state.get("chapters") or st.session_state.get("current_story"):
                export_title = st.text_input("Book Title", value=st.session_state.current_project, key="export_title_input")
                col_meta1, col_meta2 = st.columns(2)
                with col_meta1:
                    export_author = st.text_input("Author Name", value=st.session_state.user, key="export_author_input")
                with col_meta2:
                    export_publisher = st.text_input("Publisher (출판사)", value="", placeholder="e.g., My Romance Books", key="export_publisher_input")

                cover_file_path = os.path.join(BASE_DATA_DIR, username, st.session_state.current_project, "cover.png")
                cover_exists = os.path.exists(cover_file_path)
                if not cover_exists:
                    st.warning("⚠️ 등록된 책 표지 이미지(cover.png)가 없습니다. 기본 텍스트 표지로 대체됩니다. 표지는 '🎨 Art Studio' 탭을 이용하십시오.")
                else:
                    st.success("🟢 등록된 표지 이미지가 확인되었습니다. EPUB 전자책 생성 시 자동으로 탑재됩니다.")

                # --- EPUB 세부 출판 설정 옵션 ---
                st.markdown("#### ⚙️ EPUB 출판 옵션 설정")
                show_title_body = st.checkbox(
                    "📖 챕터 제목 본문 표시",
                    value=st.session_state.get("show_chapter_title_in_body", True),
                    key="show_chapter_title_in_body_chk",
                    help="해당 옵션을 해제하면 각 챕터의 첫 줄 혹은 생성되는 챕터 제목(예: 제1화)이 본문 상단에 표시되지 않고 바로 본문만 시작됩니다."
                )
                st.session_state.show_chapter_title_in_body = show_title_body

                full_content = ""
                if st.session_state.get("chapters"):
                    sorted_ch_nums = sorted([int(k) for k in st.session_state.chapters.keys()])
                    for num in sorted_ch_nums:
                        ch_text = st.session_state.chapters.get(str(num), "").strip()
                        if ch_text:
                            clean_ch = ch_text.strip()
                            if not re.match(r'^\[Chapter\s+\d+\]', clean_ch, re.IGNORECASE):
                                full_content += f"\n\n[Chapter {num}]\n\n{clean_ch}"
                            else:
                                full_content += f"\n\n{clean_ch}"
                if not full_content.strip():
                    full_content = st.session_state.get("current_story", "")

                col_exp1, col_exp2, col_exp3 = st.columns(3)
                with col_exp1:
                    if st.button("📄 Download Full TXT"):
                        st.session_state.pop("export_download_data", None)
                        settings_payload = {
                             "title": export_title,
                             "author": export_author,
                             "publisher": export_publisher,
                             "content": full_content,
                             "export_type": "txt"
                        }
                        try:
                            res = requests.post(f"{BACKEND_URL}/export/download", json=settings_payload, timeout=300)
                            if res.status_code == 200:
                                 st.session_state.export_download_data = {
                                     "data": res.content,
                                     "file_name": f"{export_title}.txt",
                                     "mime": "text/plain",
                                     "label": "⬇️ Click to Save TXT"
                                 }
                                 st.rerun()
                        except:
                            pass
                with col_exp2:
                    if st.button("✂️ Download Serial TXT (ZIP)"):
                        st.session_state.pop("export_download_data", None)
                        settings_payload = {
                             "title": export_title,
                             "author": export_author,
                             "publisher": export_publisher,
                             "content": full_content,
                             "export_type": "txt_zip"
                        }
                        try:
                            res = requests.post(f"{BACKEND_URL}/export/download", json=settings_payload, timeout=300)
                            if res.status_code == 200:
                                 st.session_state.export_download_data = {
                                     "data": res.content,
                                     "file_name": f"{st.session_state.current_project}_serial_txt.zip",
                                     "mime": "application/zip",
                                     "label": "⬇️ Click to Save ZIP (TXT)"
                                 }
                                 st.rerun()
                        except:
                            pass
                with col_exp3:
                    if st.button("📚 Download Serial EPUB (ZIP)"):
                        st.session_state.pop("export_download_data", None)
                        settings_payload = {
                             "title": export_title,
                             "author": export_author,
                             "publisher": export_publisher,
                             "content": full_content,
                             "export_type": "epub_zip",
                             "cover_image_path": cover_file_path if cover_exists else None,
                             "show_chapter_title_in_body": st.session_state.get("show_chapter_title_in_body", True),
                             "add_chapter_title_page": st.session_state.get("add_chapter_title_page", False)
                        }
                        try:
                            res = requests.post(f"{BACKEND_URL}/export/download", json=settings_payload, timeout=300)
                            if res.status_code == 200:
                                 st.session_state.export_download_data = {
                                     "data": res.content,
                                     "file_name": f"{export_title}_serial_epub.zip",
                                     "mime": "application/zip",
                                     "label": "⬇️ Click to Save ZIP (EPUB)"
                                 }
                                 st.rerun()
                        except:
                            pass

                # Render download button if ready
                if "export_download_data" in st.session_state:
                    st.markdown("---")
                    down = st.session_state.export_download_data
                    downloaded = st.download_button(
                        label=down["label"],
                        data=down["data"],
                        file_name=down["file_name"],
                        mime=down["mime"],
                        key="export_download_btn_normal"
                    )
                    if downloaded:
                        del st.session_state.export_download_data
                        st.success("다운로드가 완료되어 캐시가 자동으로 정리되었습니다.")
                        st.rerun()

                # --- 스마트 단행본 빌더 (Smart Book Packer) ---
                st.markdown("---")
                st.markdown("### 📚 스마트 단행본 빌더 (Smart Book Packer)")
                st.caption("AI가 전체 화차의 줄거리와 갈등 전개를 분석하여 단행본 규격에 어울리는 분할 지점(부/장)과 소제목을 추천합니다.")
                
                if st.button("🔍 AI 스마트 분할 및 제목 분석 실행", key="run_smart_book_packer"):
                    if not st.session_state.get("memory_chain"):
                        st.warning("분석을 위한 장기 기억(Memory Chain) 데이터가 없습니다. 본문을 먼저 요약하여 저장해 주세요.")
                    else:
                        with st.spinner("스토리 뼈대 및 변곡점 분석 중..."):
                            try:
                                res = requests.post(
                                    f"{BACKEND_URL}/export/smart-split-recommendation",
                                    json={"memory_chain": st.session_state.memory_chain},
                                    timeout=120
                                )
                                if res.status_code == 200:
                                    data = res.json()
                                    recs = data.get("recommendations", [])
                                    st.session_state.smart_split_recommendations = recs
                                    save_smart_split_settings(smart_settings_path, recs)
                                    st.success("AI 분석 완료! 아래에서 추천 구성을 확인하고 편집할 수 있습니다.")
                                    st.rerun()
                                else:
                                    st.error(f"분석 실패: {res.text}")
                            except Exception as e:
                                st.error(f"통신 오류: {e}")

                # 만약 추천 리스트가 세션 상태에 있다면 에디터 리스트 표시
                if st.session_state.get("smart_split_recommendations"):
                    add_title_page = st.checkbox(
                        "📄 챕터 간지(제목만 있는 단독 페이지) 추가",
                        value=st.session_state.get("add_chapter_title_page", False),
                        key="add_chapter_title_page_chk",
                        help="각 챕터(장)가 시작되기 전에 제목만 크게 들어간 독립적인 페이지를 1장 추가합니다."
                    )
                    st.session_state.add_chapter_title_page = add_title_page
                    
                    st.info("💡 각 부(장)의 제목을 필요에 따라 수정한 뒤 EPUB를 다운로드하세요.")
                    
                    updated_recs = []
                    recs_changed = False
                    for idx, rec in enumerate(st.session_state.smart_split_recommendations):
                        st.markdown(f"**📖 장(Volume) {rec['volume_num']}** ({rec['start_chap']}화 ~ {rec['end_chap']}화)")
                        if rec.get("rationale"):
                            st.caption(f"💡 AI 추천 근거: {rec['rationale']}")
                        
                        # Input for Title editing
                        input_title = st.text_input(
                            f"장 제목 (수정 가능)", 
                            value=rec["title"], 
                            key=f"smart_vol_title_{idx}"
                        )
                        
                        if input_title != rec["title"]:
                            rec["title"] = input_title
                            recs_changed = True
                            
                        updated_recs.append({
                            "volume_num": rec["volume_num"],
                            "start_chap": rec["start_chap"],
                            "end_chap": rec["end_chap"],
                            "title": input_title,
                            "rationale": rec.get("rationale", "")
                        })
                        st.markdown("---")
                        
                    if recs_changed:
                        save_smart_split_settings(smart_settings_path, st.session_state.smart_split_recommendations)
                        
                    # Button to download EPUB using this custom split
                    if st.button("📘 단행본 규격 EPUB 다운로드", key="download_smart_epub"):
                        st.session_state.pop("smart_export_download_data", None)
                        settings_payload = {
                             "title": export_title,
                             "author": export_author,
                             "publisher": export_publisher,
                             "content": full_content,
                             "export_type": "epub",
                             "cover_image_path": cover_file_path if cover_exists else None,
                             "volumes": updated_recs,
                             "show_chapter_title_in_body": st.session_state.get("show_chapter_title_in_body", True),
                             "add_chapter_title_page": st.session_state.get("add_chapter_title_page", False)
                        }
                        with st.spinner("단행본 병합 및 이북 빌드 중..."):
                            try:
                                res = requests.post(f"{BACKEND_URL}/export/download", json=settings_payload, timeout=300)
                                if res.status_code == 200:
                                     st.session_state.smart_export_download_data = {
                                         "data": res.content,
                                         "file_name": f"{export_title}_단행본.epub",
                                         "mime": "application/epub+zip",
                                         "label": "⬇️ Click to Save EPUB (단행본)"
                                     }
                                     st.rerun()
                                else:
                                    st.error(f"Export Failed: {res.text}")
                            except Exception as e:
                                st.error(f"Conn Error: {e}")

                    if "smart_export_download_data" in st.session_state:
                        st.markdown("---")
                        st.success("🎉 단행본 EPUB 빌드가 완료되었습니다! 아래 다운로드 버튼을 클릭하세요.")
                        down_smart = st.session_state.smart_export_download_data
                        downloaded_smart = st.download_button(
                            label=down_smart["label"],
                            data=down_smart["data"],
                            file_name=down_smart["file_name"],
                            mime=down_smart["mime"],
                            key="export_download_btn_smart",
                            type="primary"
                        )
                        if downloaded_smart:
                            del st.session_state.smart_export_download_data
                            st.success("단행본 다운로드가 완료되어 캐시가 자동으로 정리되었습니다.")
                            st.rerun()

                    # Reset all smart split recommendations
                    st.markdown(" ")
                    if st.button("🗑️ 스마트 단행본 추천 구성 초기화 (Clear)", key="clear_smart_split_cache"):
                        if os.path.exists(smart_settings_path):
                            try:
                                os.remove(smart_settings_path)
                            except:
                                pass
                        st.session_state.pop("smart_split_recommendations", None)
                        st.session_state.pop("smart_export_download_data", None)
                        st.success("단행본 추천 구성 데이터가 성공적으로 초기화되었습니다.")
                        st.rerun()

    # ==========================================
    # TAB 2: NOVEL FACTORY (Batch)
    # ==========================================
    elif st.session_state.current_page == "🧐 Editor's Desk":
        st.header("🧐 Comprehensive Review & Auto-Fix")
        
        # Select target focus for review
        scope_option = st.radio(
            "Analysis Target",
            ["Entire Story (Map-Reduce Summaries)", "Current Chapter Only"],
            index=0,
            horizontal=True,
            help="Entire Story uses saved hyper-precision chapter summaries. Current Chapter reads target chapter raw text."
        )

        chapter_list = sorted(map(int, st.session_state.chapters.keys())) if st.session_state.get("chapters") else [1]
        
        target_analyze_ch_num = st.session_state.get("current_chapter_idx", 1)
        if scope_option == "Current Chapter Only":
            default_analyze_idx = 0
            if target_analyze_ch_num in chapter_list:
                default_analyze_idx = chapter_list.index(target_analyze_ch_num)
            target_analyze_ch_num = st.selectbox("Select Chapter to Analyze", chapter_list, index=default_analyze_idx)

        # 1. Review Section
        if st.button("Run Deep Analysis", key="run_deep_analysis_review"):
             with st.spinner("Analyzing story..."):
                  try:
                      payload = {"model": selected_model_assistant}
                      if scope_option.startswith("Entire"):
                          payload["memory_chain"] = st.session_state.get("memory_chain", [])
                      else:
                          payload["text"] = st.session_state.chapters.get(str(target_analyze_ch_num), "")
                      
                      # 기존에 이미 사용자가 완료했던 교정 완료 목록(applied_fixes)을 가져와 백엔드로 전송

                      
                      old_review = st.session_state.get("review_result")

                      
                      old_fixes = []

                      
                      if isinstance(old_review, dict):

                      
                          old_fixes = old_review.get("applied_fixes", [])

                      
                          if not isinstance(old_fixes, list):

                      
                              old_fixes = []


                      
                      res = requests.post(

                      
                          f"{BACKEND_URL}/analyze/review_comprehensive",

                      
                          json={

                      
                              "text": payload.get("text"),

                      
                              "memory_chain": payload.get("memory_chain"),

                      
                              "applied_fixes": old_fixes,

                      
                              "model": payload["model"]

                      
                          }, timeout=120

                      
                      )

                      
                      if res.status_code == 200:
                           new_review = res.json().get("review", {})
                           if not isinstance(new_review, dict):
                               new_review = {}
                           # applied_fixes 영구 보존: 백엔드 반환값 + 기존 로컬값 병합
                           server_fixes = new_review.get("applied_fixes", [])
                           if not isinstance(server_fixes, list):
                               server_fixes = []
                           merged_fixes = set()
                           for x in server_fixes:
                               try: merged_fixes.add(int(x))
                               except (ValueError, TypeError): pass
                           for fix_ch in old_fixes:
                               try: merged_fixes.add(int(fix_ch))
                               except (ValueError, TypeError): pass
                           new_review["applied_fixes"] = sorted(list(merged_fixes))
                           st.session_state.review_result = new_review
                           auto_save()
                      else:
                           st.error(res.text)
                  except Exception as e:
                      st.error(f"Connection Error: {e}")

        if "review_result" in st.session_state and st.session_state.review_result:
            review = st.session_state.review_result
            if not isinstance(review, dict):
                st.info("비평 원본 데이터:")
                st.write(review)
            else:
                scores_val = review.get('scores', {})
                if isinstance(scores_val, dict):
                    st.markdown(f"### 📊 Score: 일관성 **{scores_val.get('consistency', 'N/A')}** | 가독성 **{scores_val.get('grammar_flow', 'N/A')}** | 창의성 **{scores_val.get('creativity', 'N/A')}**")
                else:
                    st.subheader(f"Score: {scores_val}")
            
            # Show AI Recommended Chapters to Fix
            raw_recs = review.get("recommended_chapters", 
                                  review.get("recommended_chapter", 
                                  review.get("suggested_chapters", 
                                  review.get("chapters_to_fix", []))))
            rec_chaps = []
            if raw_recs:
                for item in raw_recs:
                    if isinstance(item, dict):
                        ch_id = int(item.get("chapter", 0))
                        reason = item.get("reason", "교정 필요")
                        rec_chaps.append({"chapter": ch_id, "reason": reason})
                    elif isinstance(item, (int, str)):
                        try:
                            rec_chaps.append({"chapter": int(item), "reason": "교정 필요"})
                        except ValueError:
                            pass
                try:
                    rec_chaps = sorted(rec_chaps, key=lambda x: x["chapter"])
                except Exception:
                    pass

            applied_fixes = review.get("applied_fixes", []) if isinstance(review.get("applied_fixes"), list) else []
            if rec_chaps:
                st.success("🎯 **AI 비평가 선정 추천 교정 대상 화차**")
                for item in rec_chaps:
                    ch_id = item["chapter"]
                    reason = item["reason"]
                    if ch_id in applied_fixes:
                        st.markdown(f"- 🟢 **제 {ch_id}화 (교정 완료)**: {reason}")
                    else:
                        st.markdown(f"- 🔴 **제 {ch_id}화 (미반영)**: {reason}")
            else:
                st.info("ℹ️ **추천 교정 대상 화차 없음**: AI 비평가 분석 결과, 흐름상 심각한 결함이 있는 화차가 없거나 모든 화차의 완성도가 매우 뛰어납니다.")
            
            with st.expander("Detailed Critique", expanded=True):
                feedback_val = review.get("feedback")
                if isinstance(feedback_val, dict):
                    st.markdown(f"**🧐 일관성 및 개연성 (Consistency)**: {feedback_val.get('consistency', '')}")
                    st.markdown(f"**✍️ 문장 및 서사 전개 (Grammar & Flow)**: {feedback_val.get('grammar_flow', '')}")
                    st.markdown(f"**💡 창의성 및 텐션 (Creativity)**: {feedback_val.get('creativity', '')}")
                else:
                    st.markdown(str(feedback_val))
                
            with st.expander("Improvement Suggestions"):
                suggestions_val = review.get("improvement_suggestions")
                if isinstance(suggestions_val, list):
                    for sug in suggestions_val:
                        st.markdown(f"- {sug}")
                else:
                    st.markdown(str(suggestions_val))

            # 2. Auto-Fix (Rewriting)
            st.markdown("---")
            st.markdown("### 🛠️ Auto-Fix Assistant")

            # ── 세션 상태 초기화 ──
            if "batch_fix_result" not in st.session_state:
                st.session_state.batch_fix_result = None
            if "batch_fix_selected" not in st.session_state:
                st.session_state.batch_fix_selected = []

            fix_tab_batch, fix_tab_single = st.tabs(["🚀 Holistic Batch Fix (권장)", "🔧 단일 화차 Auto-Fix"])

            # ── TAB A: HOLISTIC BATCH FIX ──
            with fix_tab_batch:
                st.caption(
                    "미반영 화차들을 **하나의 서사 문제**로 인식합니다. "
                    "AI가 전체 맥락에서 수술 계획서를 먼저 설계한 뒤, "
                    "각 화차가 서로를 인식하며 일괄 수정됩니다."
                )

                if not rec_chaps:
                    st.info("ℹ️ Run Deep Analysis를 먼저 실행하여 교정 대상 화차를 찾아주세요.")
                else:
                    unfixed_chaps = [item for item in rec_chaps if item["chapter"] not in applied_fixes]
                    fixed_chaps_list = [item for item in rec_chaps if item["chapter"] in applied_fixes]

                    if fixed_chaps_list:
                        st.success(f"✅ 교정 완료: {', '.join(['제' + str(x.get('chapter', '')) + '화' for x in fixed_chaps_list])}")

                    if not unfixed_chaps:
                        st.success("🎉 모든 추천 화차가 교정 완료되었습니다!")
                    else:
                        st.markdown(f"**📋 미반영 화차 {len(unfixed_chaps)}개** — 수정할 화차를 선택하세요:")

                        col_sel_all, col_desel_all = st.columns([1, 1])
                        with col_sel_all:
                            if st.button("✅ 전체 선택", key="batch_select_all"):
                                st.session_state.batch_fix_selected = [item["chapter"] for item in unfixed_chaps]
                                st.rerun()
                        with col_desel_all:
                            if st.button("❌ 전체 해제", key="batch_desel_all"):
                                st.session_state.batch_fix_selected = []
                                st.rerun()

                        selected_chapters = []
                        for item in unfixed_chaps:
                            ch_id = item["chapter"]
                            ch_reason = item["reason"]
                            is_checked = ch_id in st.session_state.batch_fix_selected
                            checked = st.checkbox(
                                f"제{ch_id}화 — {ch_reason[:80]}{'...' if len(ch_reason) > 80 else ''}",
                                value=is_checked,
                                key=f"batch_chk_{ch_id}"
                            )
                            if checked:
                                selected_chapters.append(item)

                        st.session_state.batch_fix_selected = [item["chapter"] for item in selected_chapters]

                        if selected_chapters:
                            nums_str = ', '.join([f'제{x["chapter"]}화' for x in selected_chapters])
                            st.info(f"💡 선택됨: **{nums_str}** ({len(selected_chapters)}개) — AI가 이 화차들을 하나의 서사 문제로 인식합니다.")

                            if st.button(
                                f"🔬 Holistic Batch Fix 실행 ({len(selected_chapters)}개 화차)",
                                key="run_batch_fix",
                                type="primary",
                                use_container_width=True
                            ):
                                rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                                chapters_to_fix_payload = []
                                for item in selected_chapters:
                                    ch_num = item["chapter"]
                                    ch_text = st.session_state.chapters.get(str(ch_num), "")
                                    if ch_text.strip():
                                        chapters_to_fix_payload.append({
                                            "chapter_num": ch_num,
                                            "reason": item["reason"],
                                            "text": ch_text
                                        })

                                review_obj = st.session_state.get("review_result", {})
                                overall_critique_text = ""
                                if isinstance(review_obj.get("feedback"), dict):
                                    overall_critique_text = "\n".join([f"{k}: {v}" for k, v in review_obj["feedback"].items()])
                                elif isinstance(review_obj.get("feedback"), str):
                                    overall_critique_text = review_obj["feedback"]
                                sug = review_obj.get("improvement_suggestions", [])
                                if sug:
                                    if isinstance(sug, list):
                                        overall_critique_text += "\n개선 제안:\n" + "\n".join(f"- {s}" for s in sug)

                                with st.spinner(f"🔬 Pass 1: 수술 계획서 작성 → Pass 2: {len(selected_chapters)}개 화차 수정 → Pass 3: 검증 (예상 {len(selected_chapters)*50}초+)"):
                                    try:
                                        res = requests.post(
                                            f"{BACKEND_URL}/analyze/batch_fix",
                                            json={
                                                "chapters_to_fix": chapters_to_fix_payload,
                                                "all_memory_chain": st.session_state.get("memory_chain", []),
                                                "plot_outline": st.session_state.get("plot_outline", ""),
                                                "char_sheet": st.session_state.get("char_sheet", ""),
                                                "world_setting": st.session_state.get("world_setting", ""),
                                                "overall_critique": overall_critique_text,
                                                "style_guide": st.session_state.get("ig_style_guide", ""),
                                                "model": selected_model_assistant,
                                                "rag_enabled": rag_enabled,
                                                "rag_category_id": rag_category_id,
                                                "rag_series_id": rag_series_id,
                                                "rag_keyword": rag_keyword
                                            },
                                            timeout=600
                                        )
                                        if res.status_code == 200:
                                            st.session_state.batch_fix_result = res.json()
                                            st.success("✅ Holistic Batch Fix 완료! 아래 결과를 검토하세요.")
                                        else:
                                            try:
                                                err = res.json()
                                                err_msg = err.get('detail', res.text)
                                                st.warning(f"⚠️ Batch Fix 오류: {err_msg}\n\n기존 데이터는 변경되지 않았습니다.")
                                                log_error_to_backend("Batch_Fix_Error", f"HTTP {res.status_code}: {err_msg}", context={"chapters": [c["chapter"] for c in selected_chapters]})
                                            except Exception as inner_e:
                                                st.error(res.text)
                                                log_error_to_backend("Batch_Fix_Error_Raw", f"HTTP {res.status_code}: {res.text}", detail=str(inner_e))
                                    except Exception as e:
                                        st.warning(f"⚠️ 연결 오류: {e}\n\n기존 데이터는 안전하게 보존되어 있습니다.")
                                        log_error_to_backend("Batch_Fix_Connection_Error", str(e))

                # ── Batch Fix 결과 표시 ──
                if st.session_state.batch_fix_result:
                    result = st.session_state.batch_fix_result
                    fixed_chapters_data = result.get("fixed_chapters", [])

                    st.markdown("---")
                    st.markdown("### 📊 Holistic Batch Fix 결과")

                    with st.expander("📋 Pass 1: 수술 계획서 (Surgical Plan)", expanded=True):
                        st.markdown(result.get("surgical_plan", "(수술 계획서 없음)"))

                    st.markdown("#### 📝 Pass 2: 화차별 수정 결과")
                    for ch_result in fixed_chapters_data:
                        ch_num = ch_result["chapter_num"]
                        with st.expander(f"제{ch_num}화 수정 결과 비교", expanded=False):
                            col_orig, col_fixed = st.columns(2)
                            with col_orig:
                                st.caption("🔴 원본")
                                st.text_area(f"원본_{ch_num}", value=ch_result.get("original_text", ""), height=300, label_visibility="collapsed")
                            with col_fixed:
                                st.caption("🟢 수정본")
                                st.text_area(f"수정본_{ch_num}", value=ch_result.get("fixed_text", ""), height=300, label_visibility="collapsed")

                    with st.expander("✅ Pass 3: AI 검증 리포트", expanded=True):
                        st.markdown(result.get("verification_report", "(검증 리포트 없음)"))

                    st.markdown("---")
                    col_save, col_cancel = st.columns(2)
                    with col_save:
                        if st.button("💾 전체 일괄 저장 & memory 업데이트", type="primary", use_container_width=True):
                            saved_count = 0
                            for ch_result in fixed_chapters_data:
                                ch_num = ch_result["chapter_num"]
                                fixed_text = ch_result.get("fixed_text", "")
                                if fixed_text and fixed_text.strip():
                                    st.session_state.chapters[str(ch_num)] = fixed_text
                                    if "applied_fixes" not in st.session_state.review_result or \
                                       not isinstance(st.session_state.review_result.get("applied_fixes"), list):
                                        st.session_state.review_result["applied_fixes"] = []
                                    if int(ch_num) not in st.session_state.review_result["applied_fixes"]:
                                        st.session_state.review_result["applied_fixes"].append(int(ch_num))
                                    if st.session_state.get("current_chapter_idx") == ch_num:
                                        st.session_state.current_story = fixed_text
                                    saved_count += 1
                                    
                            # ── 병렬 처리 (ThreadPoolExecutor)로 대량 저장 속도 최적화 ──
                            from concurrent.futures import ThreadPoolExecutor
                            tasks = [
                                (r["chapter_num"], r.get("fixed_text", ""))
                                for r in fixed_chapters_data
                                if r.get("fixed_text", "").strip()
                            ]
                            
                            if tasks:
                                with st.spinner(f"💾 {len(tasks)}개 화차의 장기기억 및 맥락원장 병렬 업데이트 중..."):
                                    char_sheet = st.session_state.get("char_sheet", "")
                                    existing_ledger = st.session_state.get("continuity_ledger", [])
                                    
                                    def _worker(task):
                                        c_num, text = task
                                        s_data = None
                                        l_data = None
                                        # 1. Summarize 호출
                                        try:
                                            r_ext = requests.post(
                                                f"{BACKEND_URL}/summarize",
                                                json={"text": text, "chapter_num": c_num, "model": selected_model_assistant},
                                                timeout=60
                                            )
                                            if r_ext.status_code == 200:
                                                s_data = r_ext.json()
                                        except Exception:
                                            pass
                                        # 2. Ledger 호출
                                        try:
                                            ledger_res = requests.post(
                                                f"{BACKEND_URL}/generate/chapter_ledger",
                                                json={
                                                    "chapter_num": int(c_num),
                                                    "chapter_text": text,
                                                    "char_sheet": char_sheet,
                                                    "existing_ledger": existing_ledger,
                                                    "model": selected_model_assistant
                                                },
                                                timeout=60
                                            )
                                            if ledger_res.status_code == 200:
                                                l_data = ledger_res.json().get("ledger_item")
                                        except Exception:
                                            pass
                                        return c_num, s_data, l_data
                                    
                                    with ThreadPoolExecutor(max_workers=5) as executor:
                                        results = list(executor.map(_worker, tasks))
                                    
                                    # 경쟁 조건을 방지하기 위해 메인 스레드에서 결과 수합
                                    chain = st.session_state.get("memory_chain", [])
                                    ledger = st.session_state.get("continuity_ledger", [])
                                    
                                    for c_num, s_data, l_data in results:
                                        if s_data:
                                            chain = [m for m in chain if int(m.get("chapter", 0)) != int(c_num)]
                                            chain.append(s_data)
                                        if l_data:
                                            ledger = [e for e in ledger if int(e.get("chapter", 0)) != int(c_num)]
                                            ledger.append(l_data)
                                            
                                    chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                    ledger.sort(key=lambda x: int(x.get("chapter", 0)))
                                    
                                    st.session_state.memory_chain = chain
                                    st.session_state.continuity_ledger = ledger

                            auto_save()
                            st.success(f"✅ {saved_count}개 화차 저장 및 장기기억/맥락장부 일괄 업데이트 완료!")
                            st.session_state.batch_fix_result = None
                            st.session_state.batch_fix_selected = []
                            st.rerun()
                    with col_cancel:
                        if st.button("❌ 결과 취소 (저장 안 함)", use_container_width=True):
                            st.session_state.batch_fix_result = None
                            st.rerun()

            # ── TAB B: 단일 화차 Auto-Fix ──
            with fix_tab_single:
                st.caption("특정 1개 화차만 빠르게 수정할 때 사용합니다.")
                default_fix_idx = len(chapter_list) - 1
                if rec_chaps:
                    first_rec = int(rec_chaps[0].get("chapter", 0))
                    if first_rec in chapter_list:
                        default_fix_idx = chapter_list.index(first_rec)
                target_ch_num = st.selectbox("Target Chapter to Auto-Fix", chapter_list, index=default_fix_idx, key="target_ch_num_selection")
                st.caption(f"✨ Pinpoint rewriting targeting **Chapter {target_ch_num}** based on overall critique.")

                if st.button("Generate Fix based on Critique", key="single_fix_btn"):
                    target_text = st.session_state.chapters.get(str(target_ch_num), "")
                    with st.spinner(f"Rewriting Chapter {target_ch_num}..."):
                        try:
                            chapter_specific_reason = ""
                            for item in rec_chaps:
                                if int(item.get("chapter", 0)) == int(target_ch_num):
                                    chapter_specific_reason = item.get("reason", "")
                                    break
                            memory_chain_s = st.session_state.get("memory_chain", [])
                            context_summaries = []
                            for m in memory_chain_s:
                                ch_no = int(m.get("chapter", 0))
                                if abs(ch_no - int(target_ch_num)) <= 2 and ch_no != int(target_ch_num):
                                    summary = m.get("chunk_summary", m.get("summary", ""))
                                    context_summaries.append(f"[제{ch_no}화 요약]: {summary}")
                            context_block = "\n".join(context_summaries)
                            critique_summary = ""
                            if chapter_specific_reason:
                                critique_summary += f"[🎯 제{target_ch_num}화 핀포인트 교정]\n{chapter_specific_reason}\n\n"
                            if context_block:
                                critique_summary += f"[전후 화차 맥락]\n{context_block}\n\n"
                            critique_summary += f"[전체 피드백]\nFeedback: {review.get('feedback')}\nSuggestions: {review.get('improvement_suggestions')}"
                            rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                            res = requests.post(
                                f"{BACKEND_URL}/analyze/rewrite",
                                json={
                                    "text": target_text, "critique": critique_summary,
                                    "char_sheet": st.session_state.char_sheet, "world_setting": st.session_state.world_setting,
                                    "model": selected_model_assistant, "style_guide": st.session_state.get("ig_style_guide", ""),
                                    "rag_enabled": rag_enabled, "rag_category_id": rag_category_id,
                                    "rag_series_id": rag_series_id, "rag_keyword": rag_keyword
                                }, timeout=180
                            )
                            if res.status_code == 200:
                                                                st.session_state.rewritten_text = res.json().get("rewritten")
                            else:
                                st.error(res.text)
                                log_error_to_backend("Single_AutoFix_Error", f"HTTP {res.status_code}: {res.text}", context={"target_chapter": target_ch_num})
                        except Exception as e:
                            st.error(str(e))
                            log_error_to_backend("Single_AutoFix_Exception", str(e), context={"target_chapter": target_ch_num})

                if "rewritten_text" in st.session_state:
                    st.subheader("Proposed Fix Preview")
                    proposed = st.session_state.rewritten_text
                    st.text_area("Proposed Edit", value=proposed, height=300)
                    col_btn1, col_btn2 = st.columns(2)
                    with col_btn1:
                        if st.button("💾 Apply Auto-Fix and Overwrite", key="single_apply_btn"):
                            st.session_state.chapters[str(target_ch_num)] = proposed
                            if "applied_fixes" not in st.session_state.review_result or not isinstance(st.session_state.review_result["applied_fixes"], list):
                                st.session_state.review_result["applied_fixes"] = []
                            if int(target_ch_num) not in st.session_state.review_result["applied_fixes"]:
                                st.session_state.review_result["applied_fixes"].append(int(target_ch_num))
                            if st.session_state.get("current_chapter_idx") == target_ch_num:
                                st.session_state.current_story = proposed
                            auto_save()
                            try:
                                r_ext = requests.post(f"{BACKEND_URL}/summarize", json={"text": proposed, "chapter_num": target_ch_num, "model": selected_model_assistant}, timeout=120)
                                if r_ext.status_code == 200:
                                    res_data = r_ext.json()
                                    chain = st.session_state.get("memory_chain", [])
                                    chain = [m for m in chain if int(m.get("chapter", 0)) != int(target_ch_num)]
                                    chain.append(res_data)
                                    chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                    st.session_state.memory_chain = chain
                                    auto_save()
                            except Exception as e:
                                print(f"Failed to update metadata: {e}")

                            # [Proactive STEP C] Continuity Ledger 업데이트
                            try:
                                ledger_res = requests.post(
                                    f"{BACKEND_URL}/generate/chapter_ledger",
                                    json={
                                        "chapter_num": int(target_ch_num),
                                        "chapter_text": proposed,
                                        "char_sheet": st.session_state.get("char_sheet", ""),
                                        "existing_ledger": st.session_state.get("continuity_ledger", []),
                                        "model": selected_model_assistant
                                    },
                                    timeout=60
                                )
                                if ledger_res.status_code == 200:
                                    new_ledger_item = ledger_res.json().get("ledger_item", {})
                                    if "continuity_ledger" not in st.session_state:
                                        st.session_state.continuity_ledger = []
                                    st.session_state.continuity_ledger = [
                                        e for e in st.session_state.continuity_ledger
                                        if int(e.get("chapter", 0)) != int(target_ch_num)
                                    ]
                                    st.session_state.continuity_ledger.append(new_ledger_item)
                                    st.session_state.continuity_ledger.sort(key=lambda x: int(x.get("chapter", 0)))
                                    auto_save()
                            except Exception as ledger_err:
                                print(f"[STEP C/Single Fix] Ledger 업데이트 실패 (무시): {ledger_err}")

                            st.success("Auto-Fix applied successfully! 💾")
                            del st.session_state.rewritten_text
                            st.rerun()
                    with col_btn2:
                        if st.button("❌ Discard Proposed Fix", key="single_discard_btn"):
                            del st.session_state.rewritten_text
                            st.rerun()

            # ─── 🔧 Arc Repair: 다화차 감정 아크 진단 및 수술 패널 ───
            st.markdown("---")
            st.markdown("### 🎭 Arc Repair — 감정 아크 다화차 진단 & 집필 수정 지침")
            st.caption(
                "단일 화차 Auto-Fix로 해결이 안 되는 **다화차에 걸친 감정선·개연성 구조 문제**를 진단합니다. "
                "비평가가 지적한 화차를 기준으로 전후 3화의 감정 흐름을 Pro AI가 종합 분석하여, "
                "각 화차별 구체적인 **집필 수정 지침**을 생성해 드립니다."
            )

            arc_center_ch = st.number_input(
                "진단 기준 화차 (비평가가 지적한 화차 번호)",
                min_value=1,
                max_value=max(chapter_list) if chapter_list else 999,
                value=int(target_ch_num) if target_ch_num else 1,
                step=1,
                key="arc_repair_center"
            )

            if st.button("🔬 감정 아크 진단 및 화차별 집필 지침 생성", key="arc_repair_btn", type="primary"):
                with st.spinner(f"제{arc_center_ch}화 중심으로 전후 3화 감정 아크 분석 중... (약 20~40초 소요)"):
                    try:
                        # 분석 범위: 기준 화차 ±3화
                        arc_range = range(max(1, arc_center_ch - 3), min(max(chapter_list) + 1 if chapter_list else arc_center_ch + 4, arc_center_ch + 4))
                        arc_memory = []
                        arc_texts = {}
                        mc = st.session_state.get("memory_chain", [])
                        chaps = st.session_state.get("chapters", {})

                        for ch_no in arc_range:
                            ch_str_k = str(ch_no)
                            mem_item = next((m for m in mc if int(m.get("chapter", 0)) == ch_no), None)
                            if mem_item:
                                arc_memory.append({
                                    "chapter": ch_no,
                                    "summary": mem_item.get("chunk_summary", mem_item.get("summary", "")),
                                    "characters": mem_item.get("entity_changes", {}).get("characters", "") if isinstance(mem_item.get("entity_changes"), dict) else ""
                                })
                            # 본문 앞부분 500자만 샘플링
                            if ch_str_k in chaps and chaps[ch_str_k].strip():
                                arc_texts[ch_no] = chaps[ch_str_k][:500]

                        # 비평가의 지적 사유 수집
                        flagged_reason = ""
                        for item in rec_chaps:
                            if int(item.get("chapter", 0)) == arc_center_ch:
                                flagged_reason = item.get("reason", "")
                                break

                        arc_prompt = f"""당신은 대한민국 최고의 웹소설 전문 편집장입니다.
다음 정보를 바탕으로 제{arc_center_ch}화를 중심으로 한 감정 아크 문제를 정밀 진단하고,
각 화차별 집필 수정 지침을 생성하십시오.

[비평가 지적 사유]
{flagged_reason or "감정 변화의 개연성 부족 또는 급작스러운 전환"}

[제{arc_center_ch}화 전후 스토리 메타데이터]
{chr(10).join(f"제{m['chapter']}화: {m['summary']} | 감정상태: {m['characters']}" for m in arc_memory)}

[각 화 본문 샘플 (앞부분 500자)]
{chr(10).join(f"제{k}화 샘플: {v}" for k, v in arc_texts.items())}

[전체 인물 설정]
{st.session_state.get('char_sheet', '')}

[지시사항]
1. 위 {len(arc_range)}개 화차의 감정 아크를 분석하여 **어느 화에서 어떤 감정 징검다리가 빠져 있는지** 구체적으로 진단하십시오.
2. 각 화차별로 "이 화에 추가/수정해야 할 내용"을 구체적인 집필 지침으로 작성하십시오.
3. 각 화 지침은 실제로 작가가 그대로 반영할 수 있도록 **구체적인 장면/대사/내면 묘사 방향**까지 포함하십시오.
4. 전체 아크가 자연스럽게 흐르도록 각 화의 감정 강도(1~10)를 제안하십시오.

반드시 한국어로, 다음 형식으로 출력하십시오:

## 🔍 감정 아크 진단 결과
(전체 진단 요약 2~3문장)

## 📋 화차별 집필 수정 지침
### 제XX화 [현재 감정강도 → 목표 감정강도]
- **문제**: ...
- **추가/수정 지침**: ...
- **구체적 장면 제안**: ...

(각 화차별 반복)

## ✅ 수정 후 예상 감정 아크 흐름
(수정 후의 자연스러운 감정 흐름을 한눈에 보여주는 요약)
"""
                        res_arc = requests.post(
                            f"{BACKEND_URL}/analyze/arc_repair",
                            json={
                                "prompt": arc_prompt,
                                "model": selected_model_assistant
                            },
                            timeout=180
                        )
                        if res_arc.status_code == 200:
                            arc_data = res_arc.json().get("review", {})
                            # overall_critique에 Arc Repair 결과 담겨 옴
                            arc_guide = arc_data.get("overall_critique", "")
                            if not arc_guide:
                                arc_guide = str(arc_data)
                            st.session_state["arc_repair_result"] = {
                                "center": arc_center_ch,
                                "guide": arc_guide
                            }
                        else:
                            st.error(f"Arc Repair 분석 실패: {res_arc.text}")
                    except Exception as e:
                        st.error(f"Arc Repair 오류: {e}")

            if st.session_state.get("arc_repair_result"):
                arc_res = st.session_state["arc_repair_result"]
                st.markdown(f"#### 📊 제{arc_res['center']}화 중심 감정 아크 진단 결과")
                st.markdown(arc_res["guide"])

                # ── Arc Repair → Batch Fix 원스톱 연결 ──
                st.markdown("---")
                st.info(
                    "💡 **Arc Repair 진단 결과를 Batch Fix에 바로 적용할 수 있습니다.**\n\n"
                    "아래 버튼을 누르면 이 진단 결과를 수술 계획서의 기반으로 활용하여, "
                    "비평가가 지적한 화차들을 Holistic Batch Fix로 일괄 수정합니다."
                )
                if st.button(
                    "🔧 이 Arc Repair 진단으로 Holistic Batch Fix 실행",
                    key="arc_to_batch_fix",
                    type="primary",
                    use_container_width=True
                ):
                    # Arc Repair 진단 결과를 overall_critique로, 비평가 추천 화차를 선택으로 전달
                    arc_overall_critique = arc_res.get("guide", "")
                    # 현재 비평가 추천 화차 + 아직 미반영된 것들을 자동 선택
                    review_obj_arc = st.session_state.get("review_result", {})
                    raw_recs_arc = review_obj_arc.get("recommended_chapters",
                                    review_obj_arc.get("chapters_to_fix", []))
                    applied_arc = review_obj_arc.get("applied_fixes", []) if isinstance(review_obj_arc.get("applied_fixes"), list) else []
                    arc_unfixed = []
                    for item in raw_recs_arc:
                        if isinstance(item, dict):
                            ch_id = int(item.get("chapter", 0))
                            if ch_id not in applied_arc:
                                arc_unfixed.append({"chapter": ch_id, "reason": item.get("reason", "Arc Repair 진단 기반 수정")})
                        elif isinstance(item, (int, str)):
                            try:
                                ch_id = int(item)
                                if ch_id not in applied_arc:
                                    arc_unfixed.append({"chapter": ch_id, "reason": "Arc Repair 진단 기반 수정"})
                            except Exception:
                                pass

                    if not arc_unfixed:
                        st.warning("미반영 화차가 없거나 비평 결과가 없습니다. Run Deep Analysis를 먼저 실행하세요.")
                    else:
                        rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                        chapters_to_fix_arc = []
                        for item in arc_unfixed:
                            ch_num = item["chapter"]
                            ch_text = st.session_state.chapters.get(str(ch_num), "")
                            if ch_text.strip():
                                chapters_to_fix_arc.append({
                                    "chapter_num": ch_num,
                                    "reason": item["reason"],
                                    "text": ch_text
                                })
                        with st.spinner(f"🔬 Arc Repair 기반 Batch Fix 실행 중... (예상 {len(chapters_to_fix_arc)*50}초+)"):
                            try:
                                res_bf = requests.post(
                                    f"{BACKEND_URL}/analyze/batch_fix",
                                    json={
                                        "chapters_to_fix": chapters_to_fix_arc,
                                        "all_memory_chain": st.session_state.get("memory_chain", []),
                                        "plot_outline": st.session_state.get("plot_outline", ""),
                                        "char_sheet": st.session_state.get("char_sheet", ""),
                                        "world_setting": st.session_state.get("world_setting", ""),
                                        "overall_critique": arc_overall_critique,
                                        "style_guide": st.session_state.get("ig_style_guide", ""),
                                        "model": selected_model_assistant,
                                        "rag_enabled": rag_enabled,
                                        "rag_category_id": rag_category_id,
                                        "rag_series_id": rag_series_id,
                                        "rag_keyword": rag_keyword
                                    },
                                    timeout=600
                                )
                                if res_bf.status_code == 200:
                                    st.session_state.batch_fix_result = res_bf.json()
                                    st.success("✅ Arc Repair 기반 Batch Fix 완료! Auto-Fix 탭에서 결과를 확인하세요.")
                                else:
                                    st.warning(f"⚠️ 오류: {res_bf.text}\n\n기존 데이터는 안전하게 보존되어 있습니다.")
                            except Exception as e:
                                st.warning(f"⚠️ 연결 오류: {e}")

                col_arc_clear, _ = st.columns([1, 2])
                with col_arc_clear:
                    if st.button("🗑️ Arc Repair 결과 지우기", key="clear_arc_repair"):
                        del st.session_state["arc_repair_result"]
                        st.rerun()


    # ==========================================
    # TAB 4: ART STUDIO (Production)
    # ==========================================
    elif st.session_state.current_page == "🎨 Art Studio":
        st.header("🎨 표지 제작 디자인 스튜디오")
        st.caption("Imagen 3 엔진을 활용하여 고품질 웹소설 표지를 생성하고 책의 공식 커버로 등록합니다.")
        
        # Define local cover image path
        cover_dir = os.path.join(BASE_DATA_DIR, username, st.session_state.current_project)
        cover_path = os.path.join(cover_dir, "cover.png")
        settings_path = os.path.join(cover_dir, "cover_settings.json")
        
        # Helpers for persistence
        def load_cover_settings(p_path, def_title, def_author):
            import json
            if os.path.exists(p_path):
                try:
                    with open(p_path, "r", encoding="utf-8") as f:
                        data = json.load(f)
                        return {
                            "cover_prompt": data.get("cover_prompt", ""),
                            "cover_style": data.get("cover_style", "기본(AI 추천)"),
                            "cover_focus": data.get("cover_focus", "기본(AI 추천)"),
                            "cover_title": data.get("cover_title", def_title),
                            "cover_author": data.get("cover_author", def_author),
                            "cover_include_typography": data.get("cover_include_typography", False)
                        }
                except Exception:
                    pass
            return {
                "cover_prompt": "",
                "cover_style": "기본(AI 추천)",
                "cover_focus": "기본(AI 추천)",
                "cover_title": def_title,
                "cover_author": def_author,
                "cover_include_typography": False
            }

        def save_cover_settings(p_path, s_dict):
            import json
            try:
                os.makedirs(os.path.dirname(p_path), exist_ok=True)
                with open(p_path, "w", encoding="utf-8") as f:
                    json.dump(s_dict, f, ensure_ascii=False, indent=2)
            except Exception:
                pass

        default_title = st.session_state.current_project.split("/")[-1] if st.session_state.get("current_project") else ""
        default_author = st.session_state.get("username", "작가")

        if st.session_state.get("cover_settings_project") != st.session_state.current_project:
            settings = load_cover_settings(settings_path, default_title, default_author)
            st.session_state.cover_prompt = settings["cover_prompt"]
            st.session_state.cover_style = settings["cover_style"]
            st.session_state.cover_focus = settings["cover_focus"]
            st.session_state.cover_title = settings["cover_title"]
            st.session_state.cover_author = settings["cover_author"]
            st.session_state.cover_include_typography = settings["cover_include_typography"]
            st.session_state.cover_settings_project = st.session_state.current_project

        # If cover exists, show it permanently
        if os.path.exists(cover_path):
            st.info("📖 현재 저장된 최종 책 표지")
            try:
                with open(cover_path, "rb") as img_file:
                    cover_bytes = img_file.read()
                st.image(cover_bytes, width=300, caption="등록된 책 표지 이미지 (EPUB 내보내기 시 자동 탑재)")
            except Exception:
                st.image(cover_path, width=300, caption="등록된 책 표지 이미지 (EPUB 내보내기 시 자동 탑재)")
            if st.button("🗑️ 표지 이미지 삭제", key="delete_cover_btn"):
                try:
                    os.remove(cover_path)
                    st.success("표지 이미지가 삭제되었습니다.")
                    st.rerun()
                except Exception as e:
                    st.error(f"삭제 오류: {e}")
        else:
            st.warning("⚠️ 등록된 표지 이미지가 없습니다. 아래에서 이미지를 생성하거나 올려주세요.")
        
        st.markdown("---")
        
        # 1. Prompt Generation Settings
        st.subheader("💡 프롬프트 추출 옵션 설정")
        col_style, col_focus = st.columns(2)
        with col_style:
            style_options = ["기본(AI 추천)", "웹툰/만화 일러스트", "실사 사진", "수채화/유화 손그림", "연필 스케치/소묘", "패브릭/펠트 공예", "미니멀 그래픽 디자인", "독창적인 판타지/초현실"]
            try:
                style_index = style_options.index(st.session_state.cover_style)
            except ValueError:
                style_index = 0
            cover_style = st.selectbox(
                "🎨 표지 화풍 (Style)",
                style_options,
                index=style_index,
                key="cover_style_selection_widget"
            )
        with col_focus:
            focus_options = ["기본(AI 추천)", "인물 위주", "남주 인물 위주", "여주 인물 위주", "배경 위주"]
            try:
                focus_index = focus_options.index(st.session_state.cover_focus)
            except ValueError:
                focus_index = 0
            cover_focus = st.selectbox(
                "🎯 구도 및 포커스 (Focus)",
                focus_options,
                index=focus_index,
                key="cover_focus_selection_widget"
            )
        # 타이포그래피(제목/작가명) 삽입 설정 추가
        col_title, col_author = st.columns(2)
        with col_title:
            cover_title = st.text_input("📖 소설 제목 (Title)", value=st.session_state.cover_title, key="cover_title_input_widget")
        with col_author:
            cover_author = st.text_input("✍️ 작가명 (Author)", value=st.session_state.cover_author, key="cover_author_input_widget")
            
        include_typography = st.checkbox("✒️ 표지에 소설 제목 및 작가명 글자(Typography) 포함하기", value=st.session_state.cover_include_typography, key="cover_include_typography_widget")
        
        if st.button("📖 소설 본문에서 이미지 프롬프트 자동 추출"):
            with st.spinner("본문에서 연출 장면 분석 및 이미지 프롬프트 추출 중..."):
                try:
                    res = requests.post(
                        f"{BACKEND_URL}/analyze/cover_prompt",
                        json={
                            "text": st.session_state.current_story,
                            "style": cover_style,
                            "focus": cover_focus,
                            "include_typography": include_typography,
                            "title": cover_title,
                            "author": cover_author,
                            "model": st.session_state.get("setting_model_assistant", "models/gemini-2.5-flash")
                        },
                        timeout=120
                    )
                    if res.status_code == 200:
                        extracted_prompt = res.json().get("cover_prompt")
                        st.session_state.cover_prompt = extracted_prompt
                        save_cover_settings(settings_path, {
                            "cover_prompt": extracted_prompt,
                            "cover_style": cover_style,
                            "cover_focus": cover_focus,
                            "cover_title": cover_title,
                            "cover_author": cover_author,
                            "cover_include_typography": include_typography
                        })
                        st.success("프롬프트가 성공적으로 추출되었습니다! (아래 에디터에서 수정 가능)")
                        st.rerun()
                    else:
                        st.error(res.text)
                except Exception as e:
                    st.error(str(e))
        
        prompt_input = st.text_area("🎨 표지 이미지 묘사 (프롬프트 입력 - 영어 입력 권장)", value=st.session_state.cover_prompt, height=120, key="cover_prompt_input_widget")
        
        # Save state changes locally
        settings_changed = (
            cover_style != st.session_state.cover_style or
            cover_focus != st.session_state.cover_focus or
            cover_title != st.session_state.cover_title or
            cover_author != st.session_state.cover_author or
            include_typography != st.session_state.cover_include_typography or
            prompt_input != st.session_state.cover_prompt
        )

        if settings_changed:
            st.session_state.cover_style = cover_style
            st.session_state.cover_focus = cover_focus
            st.session_state.cover_title = cover_title
            st.session_state.cover_author = cover_author
            st.session_state.cover_include_typography = include_typography
            st.session_state.cover_prompt = prompt_input
            
            save_cover_settings(settings_path, {
                "cover_prompt": prompt_input,
                "cover_style": cover_style,
                "cover_focus": cover_focus,
                "cover_title": cover_title,
                "cover_author": cover_author,
                "cover_include_typography": include_typography
            })
        
        # 2. Image Generation
        if st.button("✨ 표지 이미지 생성하기 (Imagen 3)", use_container_width=True):
            if prompt_input:
                # ── 안전한 영문 매핑을 활용한 한글 인코딩 문제 차단 ──
                STYLE_MAP = {
                    "기본(AI 추천)": "default",
                    "웹툰/만화 일러스트": "webtoon",
                    "실사 사진": "photo",
                    "수채화/유화 손그림": "painting",
                    "연필 스케치/소묘": "sketch",
                    "패브릭/펠트 공예": "fabric",
                    "미니멀 그래픽 디자인": "minimal",
                    "독창적인 판타지/초현실": "surreal"
                }
                style_key = STYLE_MAP.get(cover_style, "default")
                
                FOCUS_MAP = {
                    "기본(AI 추천)": "default",
                    "인물 위주": "pair",
                    "남주 인물 위주": "male",
                    "여주 인물 위주": "female",
                    "배경 위주": "background"
                }
                focus_key = FOCUS_MAP.get(cover_focus, "default")
                
                style_prefix = ""
                style_suffix = ""
                
                if style_key == "webtoon":
                    style_prefix = "A 2D Korean webtoon manga illustration, clean fine line art, anime cel shading style, vibrant colors, hand-drawn vector art. "
                    style_suffix = " -- Absolutely NO 3D rendering, NO photorealism, NO computer graphics, NO oil painting textures, 2D anime manga art only."
                elif style_key == "photo":
                    style_prefix = "A hyper-realistic professional photography portrait, shot on 35mm film, shallow depth of field, sharp focus, natural skin texture, cinematic lighting, 8k resolution. "
                    style_suffix = " -- Absolutely NO anime, NO 2D illustration, NO painting feel, NO drawing outlines, realistic photo only."
                elif style_key == "painting":
                    style_prefix = "A classic hand-painted fine art, thick oil painting with heavy impasto brushstrokes, or dreamy watercolor washes on textured paper. "
                    style_suffix = " -- Absolutely NO smooth vector art, NO glossy digital rendering, NO computer graphics, hand-drawn art only."
                elif style_key == "sketch":
                    style_prefix = "A monochrome hand-drawn pencil sketch, fine charcoal and graphite drawing, cross-hatching, highly detailed shading, on textured sketch paper. "
                    style_suffix = " -- Monochrome only, Absolutely NO color, NO digital painting, NO digital vectors."
                elif style_key == "fabric":
                    style_prefix = "A cozy 3D fabric art, stitched textile, cute felt craft illustration, yarn thread textures, patchwork, needle felt. "
                    style_suffix = " -- Absolutely NO smooth digital paint, NO realistic photography, NO drawing lines."
                elif style_key == "minimal":
                    style_prefix = "A minimalist modern flat vector graphic illustration, clean geometric shapes, bold solid color blocks, Swiss design poster style. "
                    style_suffix = " -- Absolutely NO complex 3D shadows, NO realistic textures, NO hand-drawn outlines."
                elif style_key == "surreal":
                    style_prefix = "Highly creative surrealism fantasy artwork, abstract dreamlike painting, magical realism, ethereal and symbolic style. "
                    style_suffix = " -- Avoid generic glossy digital romance art."
                else:
                    style_prefix = "A high-quality web novel cover style digital painting, beautiful illustration, romantic atmosphere. "
                    style_suffix = " -- Beautiful web novel illustration."

                # 2) 구도/포커스 접두사 결정
                focus_prefix = ""
                if focus_key == "pair":
                    focus_prefix = "A close-up cover featuring the main male and female characters standing close together, highlighting their romantic chemistry and emotional tension. "
                elif focus_key == "male":
                    focus_prefix = "A portrait featuring ONLY the handsome Korean male main character, solo handsome man, showing his charismatic and gentle face, solo male look. Absolutely NO women, NO female characters. "
                elif focus_key == "female":
                    focus_prefix = "A portrait featuring ONLY the beautiful Korean female main character, solo beautiful woman, showing her detailed eyes and elegant clothing. Absolutely NO men, NO male characters. "
                elif focus_key == "background":
                    focus_prefix = "A wide shot landscape scenery emphasizing the background environment, weather, and magical atmosphere. The characters are small, silhouettes, or shown from behind. "
                else:
                    focus_prefix = "A balanced composition showing both characters and background. "

                # 최종 조합 프롬프트
                final_prompt = f"{style_prefix}{focus_prefix}{prompt_input}{style_suffix}"
                
                with st.spinner("AI가 고해상도 표지를 그리는 중... (약 10초 소요)"):
                    try:
                        res = requests.post(f"{BACKEND_URL}/generate/imagen3", json={"prompt": final_prompt, "style": style_key}, timeout=120)
                        if res.status_code == 200 and "image_base64" in res.json():
                            img_bytes = base64.b64decode(res.json()["image_base64"])
                            
                            # Save to local path immediately to make it persistent
                            os.makedirs(cover_dir, exist_ok=True)
                            with open(cover_path, "wb") as f:
                                f.write(img_bytes)
                                
                            st.success("🎉 표지 이미지가 성공적으로 생성되어 프로젝트에 공식 등록되었습니다!")
                            st.rerun()
                        else:
                            st.error(res.text)
                    except Exception as e:
                         st.error(str(e))
            else:
                st.error("이미지 생성을 위해 프롬프트를 먼저 입력하거나 소설 본문에서 추출해 주십시오.")


    # ==========================================
    # TAB 5: NOVEL FACTORY (Batch Production)
    # ==========================================
    elif st.session_state.current_page == "🏭 Novel Factory (Batch)":
        st.header("🏭 Novel Factory (Batch Generation)")
        
        # 1. Job Control
        c1, c2 = st.columns([1, 1])
        with c1:
            st.subheader("🚀 Start Production")
            # Model Selection Defaults
            w_model = st.session_state.get("setting_model_writer", "gemini-2.5-pro")
            default_w_idx = 0
            if w_model in writer_model_options:
                default_w_idx = writer_model_options.index(w_model)

            p_model = st.session_state.get("setting_model_assistant", "gemini-2.5-flash")
            default_p_idx = 0
            if p_model in assistant_model_options:
                default_p_idx = assistant_model_options.index(p_model)

            batch_writer_model = st.selectbox("Writer Model", writer_model_options, key="batch_writer", index=default_w_idx)
            batch_planner_model = st.selectbox("Planner Model", assistant_model_options, key="batch_planner", index=default_p_idx)
            
            # Additional Settings
            st.write("Current Settings:")
            st.info(f"Creativity: {temperature} | Humor: {st.session_state.setting_humor}/10 | Style: {selected_style}")
            
            # Auto-Merge Option
            auto_merge = st.checkbox("☑️ 집필 완료 후 Story Engine에 개별 화차 자동 병합 저장", value=st.session_state.auto_merge_enabled, key="auto_merge_chk")
            if auto_merge != st.session_state.auto_merge_enabled:
                st.session_state.auto_merge_enabled = auto_merge
                auto_save()

            # Use Existing Outline Option
            use_existing_outline_directly = st.checkbox("Plot & Sync 탭의 아웃라인을 그대로 사용하여 집필 (AI 기획 단계 건너뛰기)", value=True, key="use_existing_outline_directly", on_change=auto_save)

            # Self-Healing Option
            self_healing = st.toggle("🩺 Enable Self-Healing (Quality Control)", value=st.session_state.self_healing_enabled, key="self_healing_chk", help="If enabled, AI will automatically rewrite chapters with low review scores (<70). Increases generation time.")
            if self_healing != st.session_state.self_healing_enabled:
                st.session_state.self_healing_enabled = self_healing
                auto_save()

            # Calculate Completed Chapters & Estimates
            existing_ch_dict = st.session_state.get("chapters", {})
            completed_ch_nums = [int(k) for k, v in existing_ch_dict.items() if v.strip()]
            completed_ch_nums.sort()
            num_completed = len(completed_ch_nums)
            target_chapters_total = int(st.session_state.get("setting_target_vols", 1)) * int(st.session_state.get("setting_target_chapters", 50))
            num_to_generate = max(0, target_chapters_total - num_completed)

            st.markdown("---")
            st.markdown("### 📊 집필 작업 및 비용/시간 예상")
            st.markdown(f"* **이미 완료된 회차 (건너뜀 및 가져오기):** `{num_completed}`개 화차 {f'(화차 번호: {completed_ch_nums})' if completed_ch_nums else ''}")
            st.markdown(f"* **새로 생성할 회차:** `{num_to_generate}`개 화차")

            # Cost / Time Estimation
            is_pro_writer = "pro" in batch_writer_model.lower() or "pro" in batch_planner_model.lower()
            if is_pro_writer:
                est_sec_per_ch = 35
                est_cost_per_ch = 0.025
                model_tier = "Pro 등급"
            else:
                est_sec_per_ch = 15
                est_cost_per_ch = 0.0015
                model_tier = "Flash 등급"

            total_est_seconds = num_to_generate * est_sec_per_ch
            est_minutes = total_est_seconds // 60
            est_seconds = total_est_seconds % 60
            total_est_cost = num_to_generate * est_cost_per_ch

            st.info(
                f"💡 **예상 비용 및 소요 시간 ({model_tier} 기준)**\n\n"
                f"- **예상 소요 시간:** 약 `{est_minutes}분 {est_seconds}초` (화당 약 {est_sec_per_ch}초)\n"
                f"- **예상 API 비용:** 약 `${total_est_cost:.4f} USD` (화당 약 ${est_cost_per_ch:.4f})\n\n"
                f"*이 값은 예상치이며 API 트래픽 및 네트워크 지연에 따라 실소요 시간은 변동될 수 있습니다. 본문은 작성 완료 후 'Comprehensive Review & Auto-Fix' 탭 등에서 정밀 검토 및 교정이 가능합니다.*"
            )

            st.markdown("### 📋 대량 집필 실행 승인")
            approve_start = st.checkbox("위의 예상 시간, 비용 및 가져오기 규칙을 확인했으며 자동 집필 시작을 승인합니다.", value=True, key="approve_batch_production")

            if st.button(f"Start Production ({target_chapters_total} Chapters)", disabled=not approve_start):
                # Prepare Settings Payload
                rag_enabled, rag_category_id, rag_series_id, rag_keyword = get_clean_rag_params()
                batch_settings = {
                    "genre": st.session_state.get("ig_genre", "Romance"),
                    "spice": st.session_state.get("ig_spice", "Unknown"),
                    "spice_level": st.session_state.get("ig_spice", "Unknown"),
                    "mood": ", ".join(st.session_state.get("ig_moods", [])),
                    "moods": st.session_state.get("ig_moods", []),
                    "chars": st.session_state.get("char_sheet", ""),
                    "characters": st.session_state.get("char_sheet", ""),
                    "char_sheet": st.session_state.get("char_sheet", ""),
                    "world": st.session_state.get("world_setting", ""),
                    "world_setting": st.session_state.get("world_setting", ""),
                    "arc": st.session_state.get("ig_arc", ""),
                    "style": selected_style,
                    "persona": st.session_state.custom_persona_input,
                    "humor_level": st.session_state.setting_humor,
                    "idea_premise": st.session_state.get("idea_suggestion", ""),
                    "creativity": temperature,
                    "rag_enabled": rag_enabled,
                    "rag_category_id": rag_category_id,
                    "rag_series_id": rag_series_id,
                    "rag_keyword": rag_keyword,
                    "style_guide": st.session_state.get("ig_style_guide", ""),
                    "trends": st.session_state.get("ig_style_guide", ""),
                    "male_tags": st.session_state.get("ig_male", []),
                    "female_tags": st.session_state.get("ig_female", [])
                }
                batch_settings = {k: v for k, v in batch_settings.items() if v is not None}
                
                with st.spinner("Initializing Factory..."):
                    try:
                        batch_payload = {
                            "settings": batch_settings,
                            "target_vols": int(st.session_state.get("setting_target_vols", 1)),
                            "chapters_per_volume": int(st.session_state.get("setting_target_chapters", 50)),
                            "model_writer": batch_writer_model,
                            "model_planner": batch_planner_model,
                            "reference_outline": st.session_state.get("plot_outline", ""),
                            "self_healing": self_healing,
                            "existing_chapters": {str(k): v for k, v in existing_ch_dict.items() if v.strip()},
                            "use_existing_outline": st.session_state.get("use_existing_outline_directly", True),
                            "chapters_settings": st.session_state.get("chapters_settings", {})
                        }
                        batch_payload = {k: v for k, v in batch_payload.items() if v is not None}
                        res = requests.post(
                            f"{BACKEND_URL}/generate/batch_start",
                            json=batch_payload, timeout=120
                        )
                        if res.status_code == 200:
                            job_id = res.json().get("job_id")
                            st.session_state.batch_job_id = job_id
                            st.success(f"Job Started! ID: {job_id}")
                        else:
                            st.error(f"Failed to start: {res.text}")
                    except Exception as e:
                        st.error(f"Connection Error: {e}")

        with c2:
            st.subheader("📺 Monitor Job")
            if st.session_state.batch_job_id:
                st.write(f"Active Job: `{st.session_state.batch_job_id}`")
                if st.button("🔄 Refresh Status"):
                    try:
                        res = requests.get(f"{BACKEND_URL}/generate/batch_status/{st.session_state.batch_job_id}")
                        if res.status_code == 200:
                            status = res.json()
                            st.session_state.batch_status = status
                            
                            # Merge results back into st.session_state.chapters in real-time if enabled
                            if st.session_state.auto_merge_enabled and "results" in status:
                                merged_any = False
                                for ch in status.get("results", []):
                                    ch_num = str(ch["chapter_num"])
                                    
                                    # Merge text
                                    if st.session_state.chapters.get(ch_num) != ch["text"]:
                                        st.session_state.chapters[ch_num] = ch["text"]
                                        merged_any = True
                                        
                                    # Merge long-term memory (metadata)
                                    if "metadata" in ch and ch["metadata"]:
                                        ch_meta = ch["metadata"]
                                        if "memory_chain" not in st.session_state or st.session_state.memory_chain is None:
                                            st.session_state.memory_chain = []
                                            
                                        # Check if chapter memory is already present
                                        existing_ids = [int(m.get("chapter", 0)) for m in st.session_state.memory_chain]
                                        if int(ch_num) not in existing_ids:
                                            st.session_state.memory_chain.append(ch_meta)
                                            # Sort memory chain by chapter number
                                            st.session_state.memory_chain.sort(key=lambda x: int(x.get("chapter", 0)))
                                            merged_any = True
                                
                                if merged_any:
                                    # Sync active story view
                                    curr_idx = str(st.session_state.get("current_chapter_idx", 1))
                                    st.session_state.current_story = st.session_state.chapters.get(curr_idx, "")
                                    auto_save()
                            
                            # Check for completion (simple check)
                            if status.get("status") == "completed":
                                st.success("Job Done!")
                                st.success("✅ 대량 집필 완료! 전체 챕터 데이터가 Story Engine의 개별 화차로 성공적으로 병합되었습니다.")

                        else:
                            st.error("Job not found or error.")
                    except Exception as e:
                       st.error(str(e))
                
                # Display Status
                if "batch_status" in st.session_state:
                    stat = st.session_state.batch_status
                    st.progress(stat.get("progress", 0) / 100)
                    st.write(f"Status: **{stat.get('status')}**")
                    
                    with st.expander("Logs", expanded=False):
                        for log in stat.get("log", []):
                            st.text(log)
                    
                    st.subheader("Generated Chapters")
                    for ch in stat.get("results", []):
                        with st.expander(f"Chapter {ch['chapter_num']}: {ch['title']}"):
                            st.write(ch['text'])
                            st.info(f"Review Score: {ch.get('review', {}).get('scores', 'N/A')}")
            else:
                st.info("No active job.")


    # ==========================================
    # TAB 5: PUBLISHER HUB
    # ==========================================
    elif st.session_state.current_page == "📦 Publisher Hub":
        st.markdown("### 📦 출판 허브 (Publisher Hub)")
        st.caption("외부 원고를 가져와 분할하거나, 편집·교정 작업을 수행합니다.")

        pub_splitter, pub_editor = st.tabs(["📄 연재 분할기", "✏️ 편집 작업대"])

        # ── Session State Init ──
        if "pub_raw_text" not in st.session_state:
            st.session_state.pub_raw_text = ""
        if "pub_editor_filename" not in st.session_state:
            st.session_state.pub_editor_filename = ""
        if "pub_episodes" not in st.session_state:
            st.session_state.pub_episodes = []
        if "pub_split_metadata" not in st.session_state:
            st.session_state.pub_split_metadata = {}
        if "pub_editor_current_ep" not in st.session_state:
            st.session_state.pub_editor_current_ep = 0
        if "pub_editor_edits" not in st.session_state:
            st.session_state.pub_editor_edits = {}
        if "pub_local_char_sheet" not in st.session_state:
            st.session_state.pub_local_char_sheet = ""
        if "pub_local_world_setting" not in st.session_state:
            st.session_state.pub_local_world_setting = ""

        # ──────────────────────────────────────────
        # SUB-TAB 1: 연재 분할기 (Serial Splitter)
        # ──────────────────────────────────────────
        with pub_splitter:
            st.markdown("#### 📄 연재 분할기")
            st.info("📁 TXT 또는 EPUB 파일을 업로드하면, 원하는 회차 수 또는 글자 수 기준으로 스마트 분할합니다.")

            uploaded_file = st.file_uploader(
                "📤 파일 업로드 (TXT, EPUB)",
                type=["txt", "epub"],
                key="pub_split_uploader",
                help="TXT 또는 EPUB 파일을 선택하세요.",
            )

            if uploaded_file is not None:
                if st.button("📥 파일 읽기", key="pub_read_file"):
                    with st.spinner("파일을 읽는 중..."):
                        try:
                            import requests as req
                            files = {"file": (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type or "application/octet-stream")}
                            res = req.post(f"{BACKEND_URL}/publisher/upload", files=files, timeout=60)
                            if res.status_code == 200:
                                data = res.json()
                                st.session_state.pub_raw_text = data["text"]
                                st.session_state.pub_editor_filename = data["filename"]
                                st.session_state.pub_episodes = []
                                st.session_state.pub_split_metadata = {}
                                
                                # 자동 원고 제목 생성 및 사이드바 목록 등록
                                clean_title = data["filename"].rsplit(".", 1)[0]
                                clean_title = "".join(ch for ch in clean_title if ch.isalnum() or ch in "_- ").strip().replace(" ", "_")
                                st.session_state.pub_current_doc = clean_title
                                auto_save_publisher_doc()
                                st.success(f"✅ '{data['filename']}' 읽기 완료! (총 {data['char_count']:,}자)")
                                st.rerun()
                            else:
                                st.error(f"오류: {res.text}")
                        except Exception as e:
                            st.error(f"통신 에러: {e}")

            if st.session_state.pub_raw_text:
                total_chars = len(st.session_state.pub_raw_text)
                st.markdown(f"**📋 로드된 파일:** `{st.session_state.pub_editor_filename}` | **총 글자 수:** {total_chars:,}자")

                with st.expander("📖 원문 미리보기 (처음 2000자)", expanded=False):
                    st.text_area(
                        "원문",
                        value=st.session_state.pub_raw_text[:2000] + ("..." if total_chars > 2000 else ""),
                        height=300,
                        disabled=True,
                        key="pub_preview_raw",
                    )

                st.markdown("---")
                st.markdown("#### ⚙️ 분할 설정")
                col_mode, col_val = st.columns(2)

                with col_mode:
                    split_mode = st.radio(
                        "분할 모드",
                        ["📚 회차 수 기준", "📝 글자 수 기준"],
                        key="pub_split_mode",
                        horizontal=True,
                    )

                with col_val:
                    if "회차" in split_mode:
                        split_value = st.number_input(
                            "총 회차 수", min_value=2, max_value=500, value=10, step=1, key="pub_split_value_ch"
                        )
                        mode_key = "chapter"
                    else:
                        split_value = st.number_input(
                            "회차당 목표 글자 수", min_value=500, max_value=50000, value=3000, step=500, key="pub_split_value_tok"
                        )
                        mode_key = "token"

                if st.button("✂️ 분할 실행", key="pub_do_split", type="primary"):
                    with st.spinner("스마트 분할 중... (문장 경계를 분석하고 있습니다)"):
                        try:
                            import requests as req
                            res = req.post(
                                f"{BACKEND_URL}/publisher/split",
                                json={
                                    "text": st.session_state.pub_raw_text,
                                    "mode": mode_key,
                                    "value": split_value,
                                },
                                timeout=60,
                            )
                            if res.status_code == 200:
                                data = res.json()
                                st.session_state.pub_episodes = data["episodes"]
                                st.session_state.pub_split_metadata = data["metadata"]
                                auto_save_publisher_doc()
                                st.success(f"✅ 분할 완료! 총 {data['metadata']['total_episodes']}화")
                                st.rerun()
                            else:
                                st.error(f"오류: {res.text}")
                        except Exception as e:
                            st.error(f"통신 에러: {e}")

                if st.session_state.pub_episodes:
                    meta = st.session_state.pub_split_metadata
                    episodes = st.session_state.pub_episodes

                    st.markdown("---")
                    st.markdown("#### 📊 분할 결과")

                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric("총 회차", f"{meta.get('total_episodes', 0)}화")
                    c2.metric("총 글자 수", f"{meta.get('total_chars', 0):,}자")
                    c3.metric("평균 글자 수", f"{meta.get('avg_chars', 0):,}자")
                    c4.metric("편차", f"{meta.get('min_chars', 0):,} ~ {meta.get('max_chars', 0):,}자")

                    st.markdown("##### 📋 회차별 미리보기")
                    char_counts = meta.get("char_counts", [])

                    for idx, ep in enumerate(episodes):
                        ep_chars = char_counts[idx] if idx < len(char_counts) else len(ep)
                        with st.expander(f"📖 {idx + 1}화 ({ep_chars:,}자)", expanded=False):
                            st.text_area(
                                f"내용 ({idx + 1}화)",
                                value=ep,
                                height=200,
                                disabled=True,
                                key=f"pub_ep_preview_{idx}",
                            )

                    st.markdown("---")
                    st.markdown("#### 💾 내보내기")
                    col_title, col_author = st.columns(2)
                    with col_title:
                        export_title = st.text_input("작품 제목", value=st.session_state.pub_editor_filename.rsplit(".", 1)[0] if st.session_state.pub_editor_filename else "작품", key="pub_export_title")
                    with col_author:
                        export_author = st.text_input("저자", value="작가", key="pub_export_author")

                    export_format = st.radio("내보내기 형식", ["TXT", "EPUB"], horizontal=True, key="pub_export_format")

                    if st.button("📦 ZIP 다운로드", key="pub_download_zip", type="primary"):
                        with st.spinner("ZIP 파일 생성 중..."):
                            try:
                                import requests as req
                                res = req.post(
                                    f"{BACKEND_URL}/publisher/export-split",
                                    json={
                                        "episodes": episodes,
                                        "title": export_title,
                                        "author": export_author,
                                        "format_type": export_format.lower(),
                                    },
                                    timeout=120,
                                )
                                if res.status_code == 200:
                                    st.download_button(
                                        label="⬇️ 다운로드",
                                        data=res.content,
                                        file_name=f"{export_title}_분할.zip",
                                        mime="application/zip",
                                        key="pub_zip_dl_btn",
                                    )
                                else:
                                    st.error(f"오류: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                    st.markdown("---")
                    if st.button("✏️ 편집 작업대로 보내기", key="pub_send_to_editor"):
                        st.session_state.pub_episodes = list(episodes)
                        st.session_state.pub_editor_current_ep = 0
                        st.session_state.pub_editor_edits = {}
                        auto_save_publisher_doc()
                        st.success("✅ 편집 작업대로 전송 완료! '✏️ 편집 작업대' 탭으로 이동하세요.")

        # ──────────────────────────────────────────
        # SUB-TAB 2: 편집 작업대 (Editorial Workbench)
        # ──────────────────────────────────────────
        with pub_editor:
            st.markdown("#### ✏️ 편집 작업대")
            st.info("📖 외부 원고를 가져와 회차별로 편집·교정하고, AI 분석 도구를 활용합니다.")

            editor_upload = st.file_uploader(
                "📤 파일 업로드 (TXT, EPUB) — 또는 분할기에서 가져오기",
                type=["txt", "epub"],
                key="pub_editor_uploader",
            )

            if editor_upload is not None:
                if st.button("📥 파일 읽기 (편집용)", key="pub_editor_read"):
                    with st.spinner("파일을 읽는 중..."):
                        try:
                            import requests as req
                            files = {"file": (editor_upload.name, editor_upload.getvalue(), editor_upload.type or "application/octet-stream")}
                            res = req.post(f"{BACKEND_URL}/publisher/upload", files=files, timeout=60)
                            if res.status_code == 200:
                                data = res.json()
                                st.session_state.pub_episodes = [data["text"]]
                                st.session_state.pub_raw_text = data["text"]
                                st.session_state.pub_editor_filename = data["filename"]
                                st.session_state.pub_editor_current_ep = 0
                                st.session_state.pub_editor_edits = {}
                                
                                # 자동 원고 제목 생성 및 사이드바 목록 등록
                                clean_title = data["filename"].rsplit(".", 1)[0]
                                clean_title = "".join(ch for ch in clean_title if ch.isalnum() or ch in "_- ").strip().replace(" ", "_")
                                st.session_state.pub_current_doc = clean_title
                                auto_save_publisher_doc()
                                st.success(f"✅ '{data['filename']}' 로드 완료! ({data['char_count']:,}자)")
                                st.rerun()
                            else:
                                st.error(f"오류: {res.text}")
                        except Exception as e:
                            st.error(f"통신 에러: {e}")

            if st.session_state.pub_episodes:
                episodes = st.session_state.pub_episodes
                total_eps = len(episodes)
                cur_ep = st.session_state.pub_editor_current_ep

                st.markdown(f"**📋 파일:** `{st.session_state.pub_editor_filename}` | **총 {total_eps}화**")

                st.markdown("---")
                nav_col1, nav_col2, nav_col3 = st.columns([1, 3, 1])

                with nav_col1:
                    if st.button("⬅️ 이전 화", key="pub_nav_prev", disabled=(cur_ep <= 0)):
                        st.session_state.pub_editor_current_ep = max(0, cur_ep - 1)
                        st.rerun()

                with nav_col2:
                    selected_ep = st.selectbox(
                        "회차 선택",
                        range(total_eps),
                        index=cur_ep,
                        format_func=lambda x: f"{x + 1}화 ({len(episodes[x]):,}자)",
                        key="pub_ep_selector",
                    )
                    if selected_ep != cur_ep:
                        st.session_state.pub_editor_current_ep = selected_ep
                        st.rerun()

                with nav_col3:
                    if st.button("➡️ 다음 화", key="pub_nav_next", disabled=(cur_ep >= total_eps - 1)):
                        st.session_state.pub_editor_current_ep = min(total_eps - 1, cur_ep + 1)
                        st.rerun()

                cur_ep = st.session_state.pub_editor_current_ep
                current_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])

                # ── 원고 전용 로컬 설정 사전 ──
                if "pub_local_char_sheet" not in st.session_state:
                    st.session_state.pub_local_char_sheet = ""
                if "pub_local_world_setting" not in st.session_state:
                    st.session_state.pub_local_world_setting = ""

                with st.expander("⚙️ 이 원고 전용 설정 사전 (스토리 엔진과 분리)", expanded=False):
                    set_col1, set_col2 = st.columns(2)
                    with set_col1:
                        st.session_state.pub_local_char_sheet = st.text_area(
                            "👤 캐릭터 시트",
                            value=st.session_state.pub_local_char_sheet,
                            placeholder="예: 주인공 김도진 (남자, 28세, 흑발, 까칠한 성격)...",
                            height=150,
                            key="pub_input_local_char_sheet"
                        )
                    with set_col2:
                        st.session_state.pub_local_world_setting = st.text_area(
                            "🌍 세계관 및 배경 설정",
                            value=st.session_state.pub_local_world_setting,
                            placeholder="예: 현대 서울, 마법이 존재하지만 숨겨진 가상 현대 판타지...",
                            height=150,
                            key="pub_input_local_world_setting"
                        )

                st.markdown(f"##### 📝 {cur_ep + 1}화 편집")
                edit_col1, edit_col2 = st.columns(2)

                with edit_col1:
                    st.markdown("**📖 원본 (읽기 전용)**")
                    st.text_area(
                        "원본",
                        value=episodes[cur_ep],
                        height=500,
                        disabled=True,
                        key=f"pub_editor_orig_{cur_ep}",
                        label_visibility="collapsed",
                    )

                with edit_col2:
                    st.markdown("**✏️ 편집본**")
                    edited_text = st.text_area(
                        "편집",
                        value=current_text,
                        height=500,
                        key=f"pub_editor_edit_{cur_ep}",
                        label_visibility="collapsed",
                    )

                if edited_text != episodes[cur_ep]:
                    st.session_state.pub_editor_edits[cur_ep] = edited_text
                    diff_chars = len(edited_text) - len(episodes[cur_ep])
                    st.caption(f"📝 수정됨 (원본 대비 {'+' if diff_chars >= 0 else ''}{diff_chars}자)")
                elif cur_ep in st.session_state.pub_editor_edits:
                    del st.session_state.pub_editor_edits[cur_ep]

                st.markdown("---")
                st.markdown("##### 🤖 AI 도구")

                ai_col1, ai_col2, ai_col3, ai_col4, ai_col5 = st.columns(5)

                with ai_col1:
                    if st.button("📊 분석", key="pub_ai_analyze", help="현재 회차 텍스트 분석"):
                        with st.spinner("AI 분석 중..."):
                            try:
                                import requests as req
                                target_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                                res = req.post(
                                    f"{BACKEND_URL}/analyze/review",
                                    json={"text": target_text, "model": "models/gemini-2.5-flash"},
                                    timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state[f"pub_ai_result_{cur_ep}"] = {"type": "분석", "data": res.json()}
                                    st.rerun()
                                else:
                                    st.error(f"오류: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                with ai_col2:
                    if st.button("🔍 일관성", key="pub_ai_consistency", help="설정과의 일관성 검사"):
                        with st.spinner("일관성 분석 중..."):
                            try:
                                import requests as req
                                target_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                                res = req.post(
                                    f"{BACKEND_URL}/analyze/consistency",
                                    json={
                                        "text": target_text,
                                        "char_sheet": st.session_state.get("pub_local_char_sheet", ""),
                                        "world_setting": st.session_state.get("pub_local_world_setting", ""),
                                        "model": "models/gemini-2.5-flash",
                                    },
                                    timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state[f"pub_ai_result_{cur_ep}"] = {"type": "일관성", "data": res.json()}
                                    st.rerun()
                                else:
                                    st.error(f"오류: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                with ai_col3:
                    if st.button("📝 종합 리뷰", key="pub_ai_review", help="종합적인 리뷰 및 점수"):
                        with st.spinner("종합 리뷰 중..."):
                            try:
                                import requests as req
                                target_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                                res = req.post(
                                    f"{BACKEND_URL}/analyze/review_comprehensive",
                                    json={"text": target_text, "model": "models/gemini-2.5-flash"},
                                    timeout=180,
                                )
                                if res.status_code == 200:
                                    st.session_state[f"pub_ai_result_{cur_ep}"] = {"type": "종합리뷰", "data": res.json()}
                                    st.rerun()
                                else:
                                    st.error(f"오류: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                with ai_col4:
                    if st.button("✍️ 맞춤법 검사", key="pub_ai_spellcheck", help="한국어 맞춤법 및 소설 교정"):
                        with st.spinner("맞춤법 및 문장 검사 중..."):
                            try:
                                import requests as req
                                target_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                                res = req.post(
                                    f"{BACKEND_URL}/publisher/check-spell",
                                    json={"text": target_text, "model": "models/gemini-2.5-flash"},
                                    timeout=120,
                                )
                                if res.status_code == 200:
                                    st.session_state[f"pub_ai_result_{cur_ep}"] = {"type": "맞춤법", "data": res.json()}
                                    st.rerun()
                                else:
                                    st.error(f"오류: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                with ai_col5:
                    if st.button("🧠 장기기억 저장", key="pub_ai_memory", help="현재 회차를 장기기억에 추가"):
                        target_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                        if "memory_chain" not in st.session_state:
                            st.session_state.memory_chain = []
                        with st.spinner("요약 생성 중..."):
                            try:
                                import requests as req
                                res = req.post(
                                    f"{BACKEND_URL}/summarize",
                                    json={
                                        "text": target_text,
                                        "chapter_num": cur_ep + 1,
                                        "chars": st.session_state.get("pub_local_char_sheet", ""),
                                    },
                                    timeout=60,
                                )
                                if res.status_code == 200:
                                    summary = res.json().get("summary", target_text[:200])
                                    st.session_state.memory_chain.append({
                                        "chapter": cur_ep + 1,
                                        "summary": summary,
                                        "source": f"[편집 작업대] {st.session_state.pub_editor_filename}",
                                    })
                                    auto_save()
                                    st.success(f"✅ {cur_ep + 1}화 요약이 장기기억에 저장되었습니다!")
                                else:
                                    st.error(f"요약 실패: {res.text}")
                            except Exception as e:
                                st.error(f"통신 에러: {e}")

                # ── AI Results Display ──
                result_key = f"pub_ai_result_{cur_ep}"
                if result_key in st.session_state:
                    result = st.session_state[result_key]
                    result_type = result.get("type", "결과")
                    result_data = result.get("data", {})

                    st.markdown("---")
                    st.markdown(f"##### 🤖 AI {result_type} 결과")

                    if result_type == "분석":
                        review = result_data.get("review", result_data)
                        if isinstance(review, dict):
                            scores = review.get("scores", {})
                            if isinstance(scores, dict):
                                score_cols = st.columns(len(scores))
                                for i, (k, v) in enumerate(scores.items()):
                                    score_cols[i].metric(k, f"{v}/100" if isinstance(v, (int, float)) else str(v))
                            feedback = review.get("feedback", review.get("critique", ""))
                            if feedback:
                                if isinstance(feedback, dict):
                                    st.markdown(f"**🧐 일관성 및 개연성 (Consistency)**: {feedback.get('consistency', '')}")
                                    st.markdown(f"**✍️ 문장 및 서사 전개 (Grammar & Flow)**: {feedback.get('grammar_flow', '')}")
                                    st.markdown(f"**💡 창의성 및 텐션 (Creativity)**: {feedback.get('creativity', '')}")
                                else:
                                    st.markdown(str(feedback))
                        else:
                            st.write(review)

                    elif result_type == "일관성":
                        report = result_data.get("report", result_data)
                        if isinstance(report, dict):
                            errors = report.get("plot_errors", [])
                            if errors:
                                st.warning(f"⚠️ {len(errors)}개의 불일치 발견")
                                for err in errors:
                                    st.markdown(f"- {err}")
                            else:
                                st.success("✅ 일관성 문제 없음!")
                            notes = report.get("notes", "")
                            if notes:
                                st.info(notes)
                        else:
                            st.write(report)

                    elif result_type == "맞춤법":
                        report = result_data.get("report", result_data)
                        if report:
                            st.markdown(report)
                        else:
                            st.info("맞춤법 및 문장 교정 결과가 비어 있습니다.")

                    elif result_type == "종합리뷰":
                        review = result_data.get("review", result_data)
                        if isinstance(review, dict):
                            scores = review.get("scores", {})
                            if isinstance(scores, dict):
                                score_cols = st.columns(min(len(scores), 5))
                                for i, (k, v) in enumerate(scores.items()):
                                    if i < 5:
                                        score_cols[i].metric(k, f"{v}/100" if isinstance(v, (int, float)) else str(v))
                            critique = review.get("critique", review.get("feedback", ""))
                            if critique:
                                with st.expander("📝 상세 비평", expanded=True):
                                    if isinstance(critique, dict):
                                        st.markdown(f"**🧐 일관성 및 개연성 (Consistency)**: {critique.get('consistency', '')}")
                                        st.markdown(f"**✍️ 문장 및 서사 전개 (Grammar & Flow)**: {critique.get('grammar_flow', '')}")
                                        st.markdown(f"**💡 창의성 및 텐션 (Creativity)**: {critique.get('creativity', '')}")
                                    else:
                                        st.markdown(str(critique))
                            suggestions = review.get("suggestions", [])
                            if suggestions:
                                with st.expander("💡 개선 제안"):
                                    for s in suggestions:
                                        st.markdown(f"- {s}")
                        else:
                            st.write(review)

                st.markdown("---")
                st.markdown("##### 💾 저장 및 내보내기")

                dl_col1, dl_col2, dl_col3 = st.columns(3)

                with dl_col1:
                    current_download_text = st.session_state.pub_editor_edits.get(cur_ep, episodes[cur_ep])
                    safe_fname = st.session_state.pub_editor_filename.rsplit('.', 1)[0] if '.' in st.session_state.pub_editor_filename else st.session_state.pub_editor_filename
                    st.download_button(
                        "⬇️ 현재 회차 다운로드",
                        data=current_download_text.encode("utf-8"),
                        file_name=f"{safe_fname}_{cur_ep + 1}화.txt",
                        mime="text/plain",
                        key="pub_dl_current",
                    )

                with dl_col2:
                    if st.button("📦 전체 수정본 ZIP", key="pub_dl_all_zip"):
                        import zipfile as zf_mod
                        zip_buf = BytesIO()
                        with zf_mod.ZipFile(zip_buf, "w", zf_mod.ZIP_DEFLATED) as zf:
                            for i, ep in enumerate(episodes):
                                edited = st.session_state.pub_editor_edits.get(i, ep)
                                zf.writestr(f"{i + 1:03d}화.txt", edited.encode("utf-8"))
                        zip_buf.seek(0)
                        st.download_button(
                            "⬇️ 다운로드",
                            data=zip_buf.getvalue(),
                            file_name=f"{safe_fname}_편집본.zip",
                            mime="application/zip",
                            key="pub_dl_zip_btn",
                        )

                with dl_col3:
                    if st.button("📄 분할기로 보내기", key="pub_send_to_splitter"):
                        merged = "\n\n".join(
                            st.session_state.pub_editor_edits.get(i, ep)
                            for i, ep in enumerate(episodes)
                        )
                        st.session_state.pub_raw_text = merged
                        safe_fn = st.session_state.pub_editor_filename.rsplit('.', 1)[0] if '.' in st.session_state.pub_editor_filename else st.session_state.pub_editor_filename
                        st.session_state.pub_editor_filename = f"{safe_fn}_편집본.txt"
                        st.session_state.pub_episodes = []
                        st.session_state.pub_split_metadata = {}
                        auto_save_publisher_doc()
                        st.success("✅ 분할기로 전송 완료! '📄 연재 분할기' 탭으로 이동하세요.")

                edited_count = len(st.session_state.pub_editor_edits)
                if edited_count > 0:
                    st.caption(f"📝 수정된 회차: {edited_count}/{total_eps}화")
            else:
                st.info("📁 파일을 업로드하거나, '📄 연재 분할기' 탭에서 '편집 작업대로 보내기'를 사용하세요.")

        # 원고 자동 저장 실행
        auto_save_publisher_doc()


if __name__ == "__main__":
    main()
